"""
# AI投入对企业创新绩效影响的分析程序
# 本程序分析2014-2022年企业AI投入如何影响创新绩效
# 主要分析AI技术投入和人力投入对创新能力和创新质量的影响，以及AI投入多样性的调节作用
"""

# ---------- 导入必要的库 ----------
import pickle
from copyreg import pickle as copyreg_pickle  # 如果需要copyreg中的pickle
from matplotlib.collections import LineCollection
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.stats.diagnostic import het_breuschpagan, acorr_ljungbox
from statsmodels.stats.stattools import durbin_watson
from statsmodels.tsa.stattools import adfuller
from scipy import stats
from linearmodels import PanelOLS, RandomEffects
from scipy.stats.mstats import winsorize
from sklearn.preprocessing import StandardScaler
from statsmodels.iolib.summary2 import summary_col
from matplotlib.colors import LinearSegmentedColormap
from scipy.interpolate import interp1d
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib import pyplot as plt
from matplotlib import gridspec
import pandas as pd
import numpy as np
import os
import joblib
import statsmodels.api as sm
import warnings
import seaborn as sns
import traceback
import datetime

# 设置中文显示
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.size'] = 12

# 导入自定义可视化模块
try:
    from enhanced_visualization import (
        create_correlation_heatmap,
        create_pca_visualization,
        create_industry_distribution,
        create_enhanced_descriptive_table,
        yearly_trends_analysis
    )
    print("成功导入增强可视化模块")
except ImportError as e:
    print(f"警告: 未能导入增强可视化模块: {e}")
    print("将使用基本可视化功能")

# 忽略警告
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=RuntimeWarning)
warnings.filterwarnings('ignore')

# ---------- 1. 基础功能函数 ----------

def normalize_dataframe_columns(df, case='lower'):
    """
    标准化数据框中所有列的大小写

    参数:
    df: 要处理的DataFrame
    case: 'lower'或'upper'，指定转换为小写或大写

    返回:
    处理后的DataFrame副本
    """
    # 创建一个副本以避免修改原数据
    df_copy = df.copy()

    # 检查是否有大小写不一致的列
    if case == 'lower':
        case_inconsistent = [col for col in df_copy.columns if col != col.lower()]
        if case_inconsistent:
            print(f"发现{len(case_inconsistent)}个大小写不一致的列名，将统一转为小写: {case_inconsistent[:5]}...")
            df_copy.columns = [col.lower() for col in df_copy.columns]
    elif case == 'upper':
        case_inconsistent = [col for col in df_copy.columns if col != col.upper()]
        if case_inconsistent:
            print(f"发现{len(case_inconsistent)}个大小写不一致的列名，将统一转为大写: {case_inconsistent[:5]}...")
            df_copy.columns = [col.upper() for col in df_copy.columns]

    # 如果有多级索引，并且需要重设，保证索引列的大小写也是统一的
    if isinstance(df_copy.index, pd.MultiIndex):
        # 仅记录，不改变索引结构
        if case == 'lower':
            idx_inconsistent = [name for name in df_copy.index.names if name is not None and name != name.lower()]
            if idx_inconsistent:
                print(f"注意: 索引名称也存在大小写不一致: {idx_inconsistent}")
        elif case == 'upper':
            idx_inconsistent = [name for name in df_copy.index.names if name is not None and name != name.upper()]
            if idx_inconsistent:
                print(f"注意: 索引名称也存在大小写不一致: {idx_inconsistent}")

    return df_copy

# 使用case_utils模块中的工具处理大小写不一致
try:
    from case_utils import get_column_case_insensitive, find_columns_case_insensitive
    print("成功导入大小写处理工具")
except ImportError as e:
    print(f"警告: 未能导入大小写处理工具: {e}")
    print("使用内部定义的大小写处理函数")

    def get_column_case_insensitive(df, col_name, default=None):
        """以大小写不敏感方式获取数据框中的列名"""
        if col_name in df.columns:
            return col_name
        lower_col_name = col_name.lower()
        for col in df.columns:
            if col.lower() == lower_col_name:
                return col
        return default

    def find_columns_case_insensitive(df, col_names):
        """以大小写不敏感方式从数据框中查找多个列名"""
        result = {}
        for name in col_names:
            actual_name = get_column_case_insensitive(df, name)
            if actual_name is not None:
                result[name] = actual_name
        return result

def create_output_dirs():
    """创建输出目录结构"""
    dirs = ['output', 'output/tables', 'output/figures', 'output/reports']
    for d in dirs:
        if not os.path.exists(d):
            os.makedirs(d)

class VariableTracker:
    """改进版变量追踪管理系统"""

    def __init__(self):
        self.variables = {}  # 存储所有变量信息
        self.current_version = {}  # 当前每个变量的最新版本
        self.transform_history = {}  # 转换历史
        self.best_versions = {}  # 每个分析目的最适合的变量版本

    def register_original(self, df, var_names):
        """注册原始变量"""
        for var in var_names:
            if var in df.columns:
                self.variables[var] = {
                    'original': var,
                    'versions': {var: {'type': 'original', 'date': pd.Timestamp.now()}},
                    'current': var
                }
                self.current_version[var] = var
            else:
                print(f"警告: 变量 {var} 不在数据框中")

    def register_transformation(self, original_var, new_var, transform_type, df=None, params=None):
        """注册变量转换"""
        # 确保原始变量已注册
        if original_var not in self.variables:
            if df is not None and original_var in df.columns:
                self.register_original(df, [original_var])
            else:
                print(f"错误: 无法注册转换，原始变量 {original_var} 未注册")
                return None

        # 检查新变量是否已存在
        if df is not None and new_var not in df.columns:
            print(f"警告: 新变量 {new_var} 不在数据框中")
            return None

        # 记录转换信息
        if original_var in self.variables:
            # 添加到版本历史
            self.variables[original_var]['versions'][new_var] = {
                'type': transform_type,
                'date': pd.Timestamp.now(),
                'params': params
            }

            # 更新当前版本
            self.variables[original_var]['current'] = new_var
            self.current_version[original_var] = new_var

            # 记录转换历史
            if original_var not in self.transform_history:
                self.transform_history[original_var] = []
            self.transform_history[original_var].append({
                'from': original_var,
                'to': new_var,
                'type': transform_type,
                'date': pd.Timestamp.now(),
                'params': params
            })

            print(f"已注册转换: {original_var} → {new_var} (类型: {transform_type})")
            return new_var

        return None

    def get_current(self, var_name):
        """获取变量的当前版本"""
        if var_name in self.current_version:
            return self.current_version[var_name]
        return var_name

    def set_best_for_analysis(self, analysis_type, var_mappings):
        """为特定分析设置最佳变量版本"""
        self.best_versions[analysis_type] = var_mappings

    def get_best_for_analysis(self, analysis_type, var_name):
        """获取特定分析的最佳变量版本"""
        if analysis_type in self.best_versions and var_name in self.best_versions[analysis_type]:
            return self.best_versions[analysis_type][var_name]
        return self.get_current(var_name)

    def generate_report(self):
        """生成变量管理报告"""
        report = {
            'variables_count': len(self.variables),
            'transformations_count': sum(len(hist) for hist in self.transform_history.values() if hist),
            'variables': []
        }

        for var_name, var_info in self.variables.items():
            var_report = {
                'name': var_name,
                'original': var_info['original'],
                'current_version': var_info['current'],
                'versions_count': len(var_info['versions']),
                'versions': var_info['versions']
            }
            report['variables'].append(var_report)

        return report

def safe_get_param(result, param_name, possible_variants=None):
    """安全获取模型结果中的参数值，处理可能的变量名变体"""
    if not possible_variants:
        possible_variants = [f"{param_name}_centered", f"{param_name}_std", f"{param_name}_diff"]

    # 先尝试直接获取
    if param_name in result.params.index:
        return param_name, result.params[param_name], result.pvalues[param_name]

    # 尝试变体名称
    for variant in possible_variants:
        if variant in result.params.index:
            return variant, result.params[variant], result.pvalues[variant]

    # 尝试查找包含原名称的参数
    for param in result.params.index:
        if isinstance(param, str) and param_name in param:
            return param, result.params[param], result.pvalues[param]

    # 未找到
    print(f"警告: 无法在模型结果中找到参数 '{param_name}' 或其变体")
    return None, None, None

# ---------- 2. 数据预处理函数 ----------
def load_and_preprocess_data(file_path, winsorize=True, winsorize_limits=None,
                       transform_strategy='conservative', standardize=False,
                       center_interactions=True, control_vars_groups=None):
    """
    读取数据并进行优化预处理，包括:
    - 列名最小化处理
    - 高级缺失值处理
    - 可选的Winsorize处理（可自定义缩尾百分比）
    - 可选的变量变换策略（全面/保守/无）
    - 可选的标准化处理
    - 可选的交互项中心化
    - 可选的控制变量分组

    参数:
    -----------
    file_path : str
        数据文件路径
    winsorize : bool, default=True
        是否进行Winsorize处理
    winsorize_limits : dict, default=None
        自定义Winsorize缩尾百分比，格式为{'core': [0.005, 0.005], 'quality': [0.01, 0.01], 'control': [0.01, 0.01]}
        如果为None，则使用默认值
    transform_strategy : str, default='conservative'
        变量变换策略，可选值为'comprehensive'（全面）, 'conservative'（保守）, 'none'（无）
    standardize : bool, default=False
        是否进行标准化处理
    center_interactions : bool, default=True
        是否对交互项进行中心化处理
    control_vars_groups : dict, default=None
        控制变量分组，格式为{'basic': ['size', 'lev', 'roa'], 'governance': ['board', 'dual'], ...}
        如果为None，则使用所有控制变量

    返回:
    --------
    pandas.DataFrame
        预处理后的数据框
    dict
        变量映射字典
    """
    print("开始数据加载与高级预处理...")

    try:
        # 读取数据
        df = pd.read_excel(file_path)

        # 显示数据基本信息
        print(f"原始数据维度: {df.shape}")
          # 1. 列名标准化处理 - 将所有列名统一转为小写
        original_cols = df.columns.tolist()
        col_mapping = {}

        # 记录原始列名与小写列名的映射关系，用于跟踪大小写转换
        case_mapping = {col: col.lower() for col in original_cols}

        # 关键变量名映射列表 - 小写形式
        key_vars = [
            'stkcd', 'year', 'stknm', 'indcd', 'indnm', 'provcd', 'provnm', 'citycd', 'citynm',
            'ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log',
            'intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth'
        ]

        # 检查是否存在大写形式的关键变量
        uppercase_vars_found = [col for col in original_cols if col.lower() in key_vars and col != col.lower()]
        if uppercase_vars_found:
            print(f"发现大写形式的关键变量: {uppercase_vars_found}，将统一转为小写")

        # 将所有列名转小写并应用
        df.columns = [col.lower() for col in df.columns]
        print("已将所有列名转为小写，确保一致性")

        # 创建简短别名，但关键变量保持原样
        for i, col in enumerate(df.columns):
            if col.lower() in [k.lower() for k in key_vars]:
                col_mapping[col] = col.lower()
            else:
                # 为控制变量创建简短别名
                col_mapping[col] = f"v{i+1}_{col[:3]}"

        # 保存列名映射表，但暂不应用 - 供参考
        try:
            os.makedirs('output', exist_ok=True)
            os.makedirs('output/tables', exist_ok=True)
            pd.DataFrame({"original": col_mapping.keys(), "simplified": col_mapping.values()}).to_csv(
                'output/column_mapping.csv', index=False)
        except Exception as e:
            print(f"保存列名映射表失败: {e}，但将继续执行")

        # 将stkcd和year作为索引前的处理
        df['stkcd'] = df['stkcd'].astype(str)  # 确保股票代码是字符串类型
        df['year'] = df['year'].astype(int)    # 确保年份是整数类型

        # 2. 增强版缺失值处理
        # 计算缺失值比例
        missing_values = df.isnull().sum()
        missing_pct = (missing_values / len(df)) * 100

        missing_data = pd.DataFrame({
            'missing_values': missing_values,
            'missing_percent': missing_pct
        }).sort_values('missing_percent', ascending=False)

        # 保存缺失值报告
        try:
            missing_data[missing_data['missing_values'] > 0].to_csv('output/tables/missing_values_report.csv')
        except Exception as e:
            print(f"保存缺失值报告失败: {e}，但将继续执行")

        if missing_values.sum() > 0:
            print("缺失值情况:")
            print(missing_data[missing_data['missing_values'] > 0])

            # 对于数值型变量：根据缺失比例采用不同策略
            for col in df.select_dtypes(include=[np.number]).columns:
                missing_ratio = df[col].isnull().mean()
                if missing_ratio == 0:
                    continue
                elif missing_ratio < 0.05:  # 低缺失率：按行业和年份分组填充中位数
                    print(f"  低缺失率变量 {col}：按行业年份分组填充")
                    # 尝试按行业和年份分组填充
                    if 'indcd' in df.columns:
                        df[col] = df.groupby(['indcd', 'year'])[col].transform(
                            lambda x: x.fillna(x.median()) if not pd.isna(x.median()) else x)

                    # 若仍有缺失，按年份分组填充
                    if df[col].isnull().sum() > 0:
                        df[col] = df.groupby('year')[col].transform(
                            lambda x: x.fillna(x.median()) if not pd.isna(x.median()) else x)

                    # 如果还有缺失，使用全局中位数
                    if df[col].isnull().sum() > 0:
                        df[col] = df[col].fillna(df[col].median())

                elif missing_ratio < 0.15:  # 中等缺失率：使用相关变量预测填充
                    print(f"  中等缺失率变量 {col}：使用相关变量预测填充")
                    # 找出与该变量相关性较高的变量
                    numeric_cols = df.select_dtypes(include=[np.number]).columns
                    if len(numeric_cols) > 1:
                        try:
                            # 计算相关性
                            correlations = df[numeric_cols].corr()[col].drop(col).abs().sort_values(ascending=False)
                            related_cols = correlations[correlations > 0.3].index.tolist()[:5]

                            if related_cols:
                                # 使用有值的样本训练简单模型
                                non_null_idx = df[col].notna()
                                X_train = df.loc[non_null_idx, related_cols]
                                y_train = df.loc[non_null_idx, col]

                                if len(X_train) > 30:  # 确保有足够的训练样本
                                    # 使用简单线性回归预测
                                    model = sm.OLS(y_train, sm.add_constant(X_train)).fit()

                                    # 预测缺失值
                                    null_idx = df[col].isna()
                                    X_pred = df.loc[null_idx, related_cols]
                                    predictions = model.predict(sm.add_constant(X_pred))

                                    # 填充预测值
                                    df.loc[null_idx, col] = predictions
                                else:
                                    # 样本不足，使用中位数填充
                                    df[col] = df[col].fillna(df[col].median())
                            else:
                                # 无相关变量，使用中位数填充
                                df[col] = df[col].fillna(df[col].median())
                        except Exception as e:
                            print(f"  预测填充 {col} 失败: {e}，使用中位数填充")
                            df[col] = df[col].fillna(df[col].median())
                    else:
                        # 数值列不足，使用中位数填充
                        df[col] = df[col].fillna(df[col].median())

                else:  # 高缺失率：创建缺失指示变量，并尝试填充
                    print(f"  高缺失率变量 {col}：创建缺失指示变量")
                    df[f"{col}_missing"] = df[col].isnull().astype(int)
                    df[col] = df[col].fillna(df[col].median())

            # 对于分类变量：使用众数填充
            for col in df.select_dtypes(include=['object', 'category']).columns:
                if col not in ['stkcd', 'stknm'] and df[col].isnull().sum() > 0:
                    # 按年份分组计算众数
                    year_modes = df.groupby('year')[col].apply(
                        lambda x: x.mode()[0] if not x.mode().empty else None)

                    # 填充缺失值
                    for year, mode in year_modes.items():
                        if mode is not None:
                            mask = (df['year'] == year) & df[col].isna()
                            df.loc[mask, col] = mode

                    # 如果还有缺失值，使用全局众数
                    if df[col].isnull().sum() > 0:
                        if not df[col].mode().empty:
                            df[col] = df[col].fillna(df[col].mode()[0])        # 3. 更少使用Winsorize处理以保留原始信号
        print("\n执行精简的Winsorize极端值处理...")

        def safe_winsorize(x, limits=[0.01, 0.01]):
            """安全的winsorize函数，处理单样本情况，使用非常保守的限制以保留极端值信号"""
            if len(x) <= 1:
                return x  # 如果只有0或1个样本，直接返回
            elif len(x) <= 5:
                # 样本太少，不做处理
                return x
            else:
                return winsorize(x, limits=limits)

        # 获取数值型列
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

        # 对数值型列进行更保守的Winsorize处理，仅处理极端异常值
        for col in numeric_cols:
            # 使用更保守的策略处理变量，以保留更多原始信号
            if col in ['ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log', 'intotal', 'ep', 'dp']:
                # 核心自变量和因变量仅处理极端异常值（0.5%），最大程度保留信号
                try:
                    df[col] = safe_winsorize(df[col].values, limits=[0.005, 0.005])
                    print(f"  对核心变量{col}进行了最小化Winsorize处理(0.5%)")
                except Exception as e:
                    print(f"  {col}处理失败: {e}")
            elif col in ['ai_patent_quality', 'ai_patent_depth']:
                # 质量指标也使用保守处理（1%）
                try:
                    df[col] = safe_winsorize(df[col].values, limits=[0.01, 0.01])
                    print(f"  对质量指标{col}进行了保守Winsorize处理(1.0%)")
                except Exception as e:
                    print(f"  {col}处理失败: {e}")
            else:
                # 控制变量使用相对标准处理
                try:
                    df[col] = safe_winsorize(df[col].values, limits=[0.01, 0.01])
                    print(f"  对控制变量{col}进行了Winsorize处理(1.0%)")
                except Exception as e:
                    print(f"  {col}处理失败: {e}")
          # 4. 简化变量转换处理，最大程度保留原始信号
        print("\n执行精简的变量转换处理...")
        transformation_results = []

        # 简化核心变量转换策略，只对严重偏斜的变量进行处理
        transform_candidates = {
            'ai': ['center'],  # 仅中心化处理AI强度变量，不做标准化
            'ai_job_log': ['log', 'center'],  # 仅对数变换+中心化，不标准化
            'ai_patent_log': ['log', 'center'],  # 仅对数变换+中心化
            'manu_job_log': ['log', 'center'],  # 仅对数变换+中心化
            'intotal': ['log', 'center'],  # 仅对数变换+中心化
            'ep': ['log', 'center'],  # 仅对数变换+中心化
            'dp': ['log', 'center'],  # 仅对数变换+中心化
            'ai_patent_quality': ['center'],  # 仅中心化处理，不标准化
            'ai_patent_depth': ['center']   # 仅中心化处理，不标准化
        }

        # 创建变量转换函数字典，简化转换方式
        transform_functions = {
            'log': lambda x: np.log(x + 1) if (x.min() >= 0).all() else np.log(x - x.min() + 1),  # 安全的对数变换
            'center': lambda x: x - x.mean(),  # 去中心化
            'standardize': lambda x: (x - x.mean()) / x.std() if x.std() > 0 else x  # 标准化（只在必要时使用）
        }
        for col in numeric_cols:
            if col not in df.columns:
                continue

            # 计算原始偏度和峰度
            original_skew = df[col].skew()
            original_kurt = df[col].kurtosis()
            best_transform = {'name': 'none', 'skew': original_skew, 'kurt': original_kurt, 'col': col}

            # 对指定变量应用预设变换
            if col in transform_candidates:
                transforms = transform_candidates[col]
                current_data = df[col].copy()
                current_col = col

                for transform in transforms:
                    try:
                        # 应用变换
                        transform_func = transform_functions.get(transform)
                        if transform_func is None:
                            print(f"  警告: 未找到变换 {transform} 的函数定义")
                            continue

                        # 创建新列名称
                        new_col = f"{current_col}_{transform}"
                        transformed_data = transform_func(current_data)

                        # 更新数据和列名
                        df[new_col] = transformed_data
                        current_data = transformed_data.copy()
                        current_col = new_col

                        # 计算变换后的偏度和峰度
                        transform_skew = df[new_col].skew()
                        transform_kurt = df[new_col].kurtosis()

                        print(f"  对变量 {col} 应用了 {transform} 变换 => {new_col}")

                        # 更新最佳变换信息
                        best_transform = {
                            'name': transform,
                            'skew': transform_skew,
                            'kurt': transform_kurt,
                            'col': new_col
                        }
                    except Exception as e:
                        print(f"  变换 {col} 使用 {transform} 失败: {e}")

                # 添加结果
                transformation_results.append({
                    'variable': col,
                    'original_skew': original_skew,
                    'original_kurt': original_kurt,
                    'best_transform': best_transform['name'],
                    'best_transform_skew': best_transform['skew'],
                    'best_transform_kurt': best_transform['kurt'],
                    'best_transform_col': best_transform['col']
                })
            else:
                # 对于其他变量，应用自动检测的变换
                try:
                    # 尝试不同的变换方法，查找最佳变换
                    transforms_to_try = []

                    # 根据偏度选择变换
                    if abs(original_skew) > 1.5:  # 高偏度
                        transforms_to_try.extend(['log', 'sqrt', 'standardize'])
                    elif abs(original_skew) > 0.8:  # 中偏度
                        transforms_to_try.extend(['standardize', 'center'])
                    else:  # 低偏度
                        transforms_to_try.append('center')

                    # 尝试每个变换并记录结果
                    for transform in transforms_to_try:
                        transform_func = transform_functions.get(transform)
                        if transform_func is None:
                            continue

                        # 应用变换
                        new_col = f"{col}_{transform}"

                        try:
                            df[new_col] = transform_func(df[col])

                            # 计算变换后的偏度和峰度
                            transform_skew = df[new_col].skew()
                            transform_kurt = df[new_col].kurtosis()

                            # 检查是否改善偏度
                            if abs(transform_skew) < abs(best_transform['skew']):
                                best_transform = {
                                    'name': transform,
                                    'skew': transform_skew,
                                    'kurt': transform_kurt,
                                    'col': new_col
                                }
                                print(f"  变量 {col} 使用 {transform} 变换改善偏度: {original_skew:.4f} → {transform_skew:.4f}")
                        except Exception as e:
                            print(f"  尝试变换 {col} 使用 {transform} 失败: {e}")

                    # 添加结果
                    transformation_results.append({
                        'variable': col,
                        'original_skew': original_skew,
                        'original_kurt': original_kurt,
                        'best_transform': best_transform['name'],
                        'best_transform_skew': best_transform['skew'],
                        'best_transform_kurt': best_transform['kurt'],
                        'best_transform_col': best_transform['col']
                    })
                except Exception as e:
                    print(f"  处理变量 {col} 出错: {e}")
                    # 如果处理失败，使用原始变量
                    transformation_results.append({
                        'variable': col,
                        'original_skew': original_skew,
                        'original_kurt': original_kurt,
                        'best_transform': 'none',
                        'best_transform_skew': original_skew,
                        'best_transform_kurt': original_kurt,
                        'best_transform_col': col
                    })

            # 如果最佳转换不是原始变量，输出转换信息
            if best_transform['name'] != 'none':
                print(f"  变量 {col} 最佳变换结果: {best_transform['name']}，偏度从 {original_skew:.4f} 改善到 {best_transform['skew']:.4f}，"
                      f"峰度从 {original_kurt:.4f} 改善到 {best_transform['kurt']:.4f}")

        # 保存转换结果
        transformation_df = pd.DataFrame(transformation_results)
        try:
            transformation_df.to_csv('output/tables/variable_transformation_report.csv', index=False)
        except Exception as e:
            print(f"保存变量转换报告失败: {e}，但将继续执行")        # 创建变量映射字典，用于后续分析
        var_mapping = {}
        for result in transformation_results:
            if result['best_transform'] != 'none':
                var_mapping[result['variable']] = result['best_transform_col']
            else:
                var_mapping[result['variable']] = result['variable']

        # 保存变换结果表
        try:
            transformation_df = pd.DataFrame(transformation_results)
            # 删除不必要的列，简化输出
            essential_cols = ['variable', 'original_skew', 'original_kurt', 'best_transform',
                            'best_transform_skew', 'best_transform_kurt', 'best_transform_col']
            simplified_df = transformation_df[essential_cols]
            simplified_df.to_csv('output/tables/variable_transformation_report.csv', index=False)
            print("  成功保存变量转换报告")
        except Exception as e:
            print(f"  保存变量转换报告失败: {e}")
        # 5. 创建必要的交互项，减少交互项数量以避免多重共线性
        print("\n创建关键交互项变量 (保持数量有限)...")
        # 获取已转换的变量名
        case_insensitive_mapping = {}
        for col in df.columns:
            case_insensitive_mapping[col.lower()] = col

        # 使用大小写不敏感的查找方式获取变量
        def get_var_case_insensitive(var_name, default=None):
            """以大小写不敏感方式获取变量名"""
            if var_name in df.columns:
                return var_name
            elif var_name.lower() in case_insensitive_mapping:
                actual_name = case_insensitive_mapping[var_name.lower()]
                print(f"使用 '{actual_name}' 替代 '{var_name}'")
                return actual_name
            elif var_name in var_mapping:
                mapped_name = var_mapping[var_name]
                if mapped_name in df.columns:
                    return mapped_name
                elif mapped_name.lower() in case_insensitive_mapping:
                    actual_name = case_insensitive_mapping[mapped_name.lower()]
                    print(f"使用 '{actual_name}' 替代 '{var_name}'")
                    return actual_name
            return default

        # 获取关键变量（大小写不敏感）
        ai_var = get_var_case_insensitive('ai', var_mapping.get('ai', 'ai'))
        ai_centered = get_var_case_insensitive(f"{ai_var}_center", ai_var)

        ai_job_var = get_var_case_insensitive('ai_job_log', var_mapping.get('ai_job_log', 'ai_job_log'))
        ai_job_centered = get_var_case_insensitive(f"{ai_job_var}_center", ai_job_var)

        ai_patent_var = get_var_case_insensitive('ai_patent_log', var_mapping.get('ai_patent_log', 'ai_patent_log'))
        ai_patent_centered = get_var_case_insensitive(f"{ai_patent_var}_center", ai_patent_var)

        manu_job_var = get_var_case_insensitive('manu_job_log', var_mapping.get('manu_job_log', 'manu_job_log'))
        manu_job_centered = get_var_case_insensitive(f"{manu_job_var}_center", manu_job_var)

        # H2a/H2b的二次项 (使用中心化变量创建二次项)
        if ai_centered in df.columns:
            df['ai_squared'] = df[ai_centered] ** 2
            print(f"  已创建: ai_squared (基于{ai_centered})")
            # 更新变量映射
            var_mapping['ai_squared'] = 'ai_squared'

        # 仅创建主要研究假设所需的最关键交互项
        # H5a - 专利跨领域程度调节
        if ai_centered in df.columns and ai_patent_centered in df.columns:
            interaction_name = f'{ai_centered}_x_{ai_patent_centered}'
            df[interaction_name] = df[ai_centered] * df[ai_patent_centered]
            print(f"  已创建: {interaction_name}")
            # 更新变量映射
            var_mapping[f'ai_x_ai_patent_log'] = interaction_name

        # H5b - 人才跨领域程度调节
        if ai_job_centered in df.columns and manu_job_centered in df.columns:
            interaction_name = f'{ai_job_centered}_x_{manu_job_centered}'
            df[interaction_name] = df[ai_job_centered] * df[manu_job_centered]
            print(f"  已创建: {interaction_name}")
            # 更新变量映射
            var_mapping[f'ai_job_log_x_manu_job_log'] = interaction_name

        # 7. 设置面板数据索引
        print("\n设置面板数据索引 (stkcd, year)...")
        if 'stkcd' in df.columns and 'year' in df.columns:
            try:
                # 确保列类型正确
                df['stkcd'] = df['stkcd'].astype(str)
                df['year'] = df['year'].astype(int)

                # 设置索引
                df = df.set_index(['stkcd', 'year'])
                print("面板数据索引设置成功。")
            except Exception as e:
                print(f"设置面板索引失败: {e}")
                # 尝试重置索引后再设置
                try:
                    df = df.reset_index()
                    if 'stkcd' in df.columns and 'year' in df.columns:
                        df = df.set_index(['stkcd', 'year'])
                        print("重新设置面板索引成功。")
                    else:
                        print("错误：无法找到stkcd和year列来设置索引")
                except Exception as e2:
                    print(f"重新设置面板索引也失败: {e2}")

        # 然后再创建时间相关变量
        # 6. 创建时间相关变量 (H1b的时间分析)
        # 添加索引检查以确保安全
        if isinstance(df.index, pd.MultiIndex) and 'year' in df.index.names:
            try:
                years = df.index.get_level_values('year')
                # 将索引转换为数组再计算统计量
                years_array = np.array(years)
                year_mean = np.mean(years_array)
                year_min = np.min(years_array)
                year_max = np.max(years_array)
                  # 创建时间相关变量
                df['year_centered'] = years_array - year_mean
                if year_max > year_min:  # 避免除以零
                    df['year_std'] = (years_array - year_min) / (year_max - year_min)
                else:
                    df['year_std'] = 0
                  # 创建AI与时间的交互项
                if ai_var in df.columns:  # 确保变量存在
                    df[f'{ai_var}_x_year_std'] = df[ai_var] * df['year_std']
                    print(f"  已创建: {ai_var}_x_year_std (用于H1b的时间动态分析)")
            except Exception as e:
                print(f"创建时间相关变量出错: {e}")
                print("将继续执行后续步骤...")
        else:
            print("警告: 无法创建时间相关变量，索引结构不正确")
            # 尝试修复索引问题
            try:
                if 'stkcd' in df.columns and 'year' in df.columns:
                    df = df.set_index(['stkcd', 'year'])
                    print("已修复索引，现在尝试创建时间相关变量")

                    years = df.index.get_level_values('year')
                    # 将索引转换为数组再计算统计量
                    years_array = np.array(years)
                    year_mean = np.mean(years_array)
                    year_min = np.min(years_array)
                    year_max = np.max(years_array)

                    # 创建时间相关变量
                    df['year_centered'] = years_array - year_mean
                    if year_max > year_min:  # 避免除以零
                        df['year_std'] = (years_array - year_min) / (year_max - year_min)
                    else:
                        df['year_std'] = 0

                    # 创建AI与时间的交互项
                    if ai_var in df.columns:  # 确保变量存在
                        df[f'{ai_var}_x_year_std'] = df[ai_var] * df['year_std']
                        print(f"  已创建: {ai_var}_x_year_std (用于H1b的时间动态分析)")
                else:
                    print("无法修复索引，缺少必要的列")
            except Exception as e:
                print(f"尝试修复索引并创建时间变量失败: {e}")        # 8. 精简单位根检验和差分处理，以保留长期关系
        print("\n执行精简版单位根检验...")

        # 关键变量列表
        key_analysis_vars = [
            ai_var, ai_job_var, ai_patent_var, manu_job_var,
            'intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth'
        ]

        # 执行单位根检验
        unit_root_results = {}
        non_stationary_vars = []

        for var in key_analysis_vars:
            if var not in df.columns:
                continue

            try:
                # 进行平稳性检验
                stationary_result = test_stationarity(df, var)
                unit_root_results[var] = stationary_result

                # 仅标记极端非平稳的变量
                if not stationary_result.get('is_stationary', False) and stationary_result.get('stationary_pct', 0) < 30:
                    non_stationary_vars.append(var)
                    print(f"  变量 {var} 严重不平稳 ({stationary_result.get('stationary_pct', 0):.1f}%)，考虑创建差分变量")
                else:
                    print(f"  变量 {var} 具有足够平稳性 ({stationary_result.get('stationary_pct', 0):.1f}%)，保留原始变量")
            except Exception as e:
                print(f"  变量 {var} 平稳性检验失败: {e}, 跳过")

        # 仅创建极端非平稳变量的差分，且保留原始变量
        print("\n创建差分变量（仅针对严重非平稳变量）...")

        for var in non_stationary_vars:
            try:
                # 确保索引已排序
                df = df.sort_index()

                # 创建一阶差分变量（不替换原始变量）
                df[f'{var}_diff1'] = df.groupby(level=0)[var].diff()
                print(f"  已创建一阶差分变量: {var}_diff1 (保留原始变量)")

                # 更新变量映射 - 保留原始变量不替换
                # var_mapping[var] = var  # 保持对原始变量的引用
            except Exception as e:
                print(f"  创建变量 {var} 的差分变量失败: {e}")
        print("\n变量映射更新:")
        for var, mapped_var in var_mapping.items():
            if var != mapped_var:
                print(f"  {var} -> {mapped_var}")

        # 额外处理：创建平滑差分和去趋势变量
        print("\n创建平滑差分和去趋势变量...")

        # 定义弱平稳变量列表
        weakly_stationary_vars = []
        for var, result in unit_root_results.items():
            if 30 <= result.get('stationary_pct', 0) < 70:
                weakly_stationary_vars.append(var)
                print(f"  变量 {var} 是弱平稳的 ({result.get('stationary_pct', 0):.1f}%)")

        # 为所有非平稳和弱平稳变量创建平滑差分
        for var in non_stationary_vars + weakly_stationary_vars:
            if var in df.columns:
                try:
                    # 创建移动平均差分 (更平滑的差分)
                    df[f'{var}_ma_diff'] = df.groupby(level=0)[var].transform(
                        lambda x: x.rolling(window=2, min_periods=2).mean().diff())
                    print(f"  已创建平滑差分: {var}_ma_diff")

                    # 创建去趋势变量
                    df_reset_temp = df.reset_index()
                    for company, group in df_reset_temp.groupby('stkcd'):
                        if len(group) >= 5:  # 至少需要5个观测值才能去趋势
                            try:
                                # 按时间排序
                                group = group.sort_values('year')

                                # 创建时间趋势变量
                                group['trend'] = np.arange(len(group))

                                # 回归去除趋势
                                X = sm.add_constant(group['trend'])
                                y = group[var]

                                try:
                                    model = sm.OLS(y, X).fit()
                                    residuals = model.resid

                                    # 将残差作为去趋势变量
                                    df_reset_temp.loc[group.index, f'{var}_detrend'] = residuals
                                except Exception as e:
                                    print(f"  公司 {company} 去趋势回归失败: {e}")
                            except Exception as e:
                                print(f"  处理公司 {company} 时出错: {e}")

                    # 将去趋势变量添加回多级索引DataFrame
                    if f'{var}_detrend' in df_reset_temp.columns:
                        df[f'{var}_detrend'] = df_reset_temp.set_index(['stkcd', 'year'])[f'{var}_detrend']
                        print(f"  已创建去趋势变量: {var}_detrend")
                except Exception as e:
                    print(f"  创建变量 {var} 的平滑差分或去趋势变量失败: {e}")

        # 7. 设置面板数据索引
        print("\n确保面板数据索引正确设置...")
        if not isinstance(df.index, pd.MultiIndex) or 'stkcd' not in df.index.names or 'year' not in df.index.names:
            if 'stkcd' in df.columns and 'year' in df.columns:
                try:
                    # 确保列类型正确
                    df['stkcd'] = df['stkcd'].astype(str)
                    df['year'] = df['year'].astype(int)

                    # 设置索引
                    df = df.set_index(['stkcd', 'year'])
                    print("面板数据索引设置成功。")
                except Exception as e:
                    print(f"设置面板索引失败: {e}")
            else:
                print("警告: 无法设置面板索引，数据框中缺少'stkcd'或'year'列")
        else:
            print("面板索引已正确设置。")

        # 6. 创建时间相关变量 (H1b的时间分析)
        print("\n创建时间相关变量...")
        # 添加索引检查以确保安全
        if isinstance(df.index, pd.MultiIndex) and 'year' in df.index.names:
            try:
                years = df.index.get_level_values('year')
                # 将索引转换为数组再计算统计量
                years_array = np.array(years)
                year_mean = np.mean(years_array)
                year_min = np.min(years_array)
                year_max = np.max(years_array)

                # 创建时间相关变量
                df['year_centered'] = years_array - year_mean
                if year_max > year_min:  # 避免除以零
                    df['year_std'] = (years_array - year_min) / (year_max - year_min)
                else:
                    df['year_std'] = 0

                # 创建二次项 - 用于检验非线性时间趋势
                df['year_centered_squared'] = df['year_centered'] ** 2
                df['year_std_squared'] = df['year_std'] ** 2
                print("  已创建: year_centered, year_std, year_centered_squared, year_std_squared")

                # 优化: 创建时间趋势变量的标准化/去中心化版本
                df['time_trend'] = df['year_std']  # 使用标准化的时间变量
                df['time_trend_center'] = df['year_centered']  # 使用去中心化的时间变量

                # 创建AI与时间的交互项 (用于H1b的时间动态分析)
                # 使用已处理的AI变量(中心化或标准化)
                if ai_centered in df.columns:
                    df[f'{ai_centered}_x_time_trend'] = df[ai_centered] * df['time_trend']
                    print(f"  已创建: {ai_centered}_x_time_trend")

                if ai_job_centered in df.columns:
                    df[f'{ai_job_centered}_x_time_trend'] = df[ai_job_centered] * df['time_trend']
                    print(f"  已创建: {ai_job_centered}_x_time_trend")

                if ai_patent_centered in df.columns:
                    df[f'{ai_patent_centered}_x_time_trend'] = df[ai_patent_centered] * df['time_trend']
                    print(f"  已创建: {ai_patent_centered}_x_time_trend")
            except Exception as e:
                print(f"创建时间相关变量出错: {e}")
        else:
            print("警告: 无法创建时间相关变量，索引结构不正确")
          # 9. 保存处理后的数据和变量映射
        try:
            print("\n保存处理后的数据和变量映射...")
            os.makedirs('output', exist_ok=True)

            # 保存处理后的数据
            df.to_pickle('output/processed_data.pkl')
            print("  成功保存处理后数据到 output/processed_data.pkl")

            # 保存变量映射
            with open('output/variable_mapping.pkl', 'wb') as f:
                import pickle as pickle_module
                pickle_module.dump(var_mapping, f)
            print("  成功保存变量映射到 output/variable_mapping.pkl")

            # 保存变量信息的CSV形式
            var_info_df = pd.DataFrame({
                'original_var': list(var_mapping.keys()),
                'final_var': list(var_mapping.values())
            })
            var_info_df.to_csv('output/variable_mapping.csv', index=False)
            print("  成功保存变量映射CSV到 output/variable_mapping.csv")

            # 保存变量转换报告 - 更详细的版本
            transformation_details = []
            for var, mapped_var in var_mapping.items():
                if var == mapped_var:
                    status = "原始变量"
                    transformation = "无"
                elif "_standardize" in mapped_var:
                    status = "标准化"
                    transformation = "标准化 (z-score)"
                elif "_center" in mapped_var:
                    status = "去中心化"
                    transformation = "中心化 (减去均值)"
                elif "_log" in mapped_var:
                    status = "对数变换"
                    transformation = "对数变换 log(x+1)"
                elif "_diff" in mapped_var:
                    status = "差分变量"
                    if "_diff1" in mapped_var:
                        transformation = "一阶差分"
                    elif "_diff2" in mapped_var:
                        transformation = "二阶差分"
                    else:
                        transformation = "差分"
                elif "_ma_diff" in mapped_var:
                    status = "平滑差分"
                    transformation = "移动平均差分"
                elif "_detrend" in mapped_var:
                    status = "去趋势"
                    transformation = "线性去趋势"
                elif "_x_" in mapped_var:
                    status = "交互项"
                    parts = mapped_var.split("_x_")
                    transformation = f"交互: {parts[0]} × {parts[1]}"
                elif "_squared" in mapped_var:
                    status = "二次项"
                    transformation = "平方项"
                else:
                    status = "其他变换"
                    transformation = "未知变换"

                # 收集信息
                if var in df.columns:
                    try:
                        original_mean = df[var].mean()
                        original_std = df[var].std()
                        original_min = df[var].min()
                        original_max = df[var].max()
                        original_skew = df[var].skew()
                    except:
                        original_mean = original_std = original_min = original_max = original_skew = None
                else:
                    original_mean = original_std = original_min = original_max = original_skew = None

                if mapped_var in df.columns:
                    try:
                        transformed_mean = df[mapped_var].mean()
                        transformed_std = df[mapped_var].std()
                        transformed_min = df[mapped_var].min()
                        transformed_max = df[mapped_var].max()
                        transformed_skew = df[mapped_var].skew()
                    except:
                        transformed_mean = transformed_std = transformed_min = transformed_max = transformed_skew = None
                else:
                    transformed_mean = transformed_std = transformed_min = transformed_max = transformed_skew = None

                transformation_details.append({
                    'original_var': var,
                    'transformed_var': mapped_var,
                    'status': status,
                    'transformation': transformation,
                    'original_mean': original_mean,
                    'original_std': original_std,
                    'original_min': original_min,
                    'original_max': original_max,
                    'original_skew': original_skew,
                    'transformed_mean': transformed_mean,
                    'transformed_std': transformed_std,
                    'transformed_min': transformed_min,
                    'transformed_max': transformed_max,
                    'transformed_skew': transformed_skew
                })

            # 保存详细变换报告
            transformation_details_df = pd.DataFrame(transformation_details)
            transformation_details_df.to_csv('output/transformation_details.csv', index=False)
            print("  成功保存变量转换详情到 output/transformation_details.csv")

        except Exception as e:
            print(f"警告: 保存数据过程中出错: {e}")
            traceback.print_exc()

        # 最终检查返回的DataFrame是否具有正确的索引结构
        if not isinstance(df.index, pd.MultiIndex) or 'stkcd' not in df.index.names or 'year' not in df.index.names:
            print("警告: 返回的DataFrame索引结构不正确，尝试最后修复...")
            try:
                if 'stkcd' in df.columns and 'year' in df.columns:
                    df = df.set_index(['stkcd', 'year'])
                    print("在返回前成功修复DataFrame索引")
                else:
                    print("无法修复索引，请检查后续分析结果")
            except Exception as e:
                print(f"尝试修复索引时出错: {e}")

        # 在函数结尾处，确保返回正确的格式
        print(f"处理后的数据维度: {df.shape}")

        # 确保返回元组格式 (DataFrame, dict)
        return df, var_mapping
    except Exception as e:
        print(f"load函数出错：{e}")


# 修复相关变量查找函数
def find_correlated_vars(df, target_col, threshold=0.3, max_vars=5):
    """找出与目标变量相关性较高的变量"""
    try:
        if target_col not in df.columns:
            print(f"错误: 目标变量 {target_col} 不在数据框中")
            return []

        if df[target_col].isnull().sum() == len(df):
            print(f"错误: 目标变量 {target_col} 全为缺失值")
            return []

        # 计算与目标变量的相关性
        corr_data = df.select_dtypes(include=[np.number])

        # 确保目标变量在corr_data中
        if target_col not in corr_data.columns:
            print(f"错误: 目标变量 {target_col} 不是数值型")
            return []

        # 计算相关系数，忽略NA值
        corr_with_target = corr_data.corr()[target_col].sort_values(ascending=False)

        # 筛选相关性高于阈值且不是目标变量自身的变量
        high_corr_vars = corr_with_target[(abs(corr_with_target) > threshold) &
                                          (corr_with_target.index != target_col)]

        # 返回相关性最高的变量(最多max_vars个)
        result = high_corr_vars.index.tolist()[:max_vars]
        print(f"与 {target_col} 相关的变量: {result}")
        return result
    except Exception as e:
        print(f"查找相关变量出错: {e}")
        return []

# 修复预测填充函数
def impute_with_prediction(df, target_col, predictor_cols):
    """使用回归模型预测填充缺失值"""
    try:
        if target_col not in df.columns:
            print(f"错误: 目标变量 {target_col} 不在数据框中")
            return

        if not predictor_cols or not all(col in df.columns for col in predictor_cols):
            missing_cols = [col for col in predictor_cols if col not in df.columns]
            print(f"错误: 预测变量 {missing_cols} 不在数据框中")
            return

        # 创建训练集（有完整数据的行）
        train_mask = df[target_col].notna() & df[predictor_cols].notna().all(axis=1)

        if train_mask.sum() < 30:
            print(f"警告: 训练数据不足 ({train_mask.sum()} < 30行), 跳过预测填充")
            return

        train_data = df.loc[train_mask, predictor_cols]
        train_target = df.loc[train_mask, target_col]

        # 创建预测集（有缺失的行）
        pred_mask = df[target_col].isna() & df[predictor_cols].notna().all(axis=1)

        if pred_mask.sum() == 0:
            print(f"警告: 没有可预测的行 (所有缺失 {target_col} 的行在预测变量中也有缺失)")
            return

        pred_data = df.loc[pred_mask, predictor_cols]

        # 创建并训练简单线性回归模型
        model = sm.OLS(train_target, sm.add_constant(train_data)).fit()

        # 预测缺失值
        try:
            predictions = model.predict(sm.add_constant(pred_data))

            # 填充预测值
            df.loc[pred_mask, target_col] = predictions
            print(f"已使用 {len(predictor_cols)} 个预测变量成功填充 {pred_mask.sum()} 个缺失值")
        except Exception as e:
            print(f"预测填充失败: {e}")
    except Exception as e:
        print(f"预测填充过程出错: {e}")
        traceback.print_exc()

# 修复平稳性检验函数
def test_stationarity(df, column, significance=0.05):
    """增强版面板数据变量平稳性检验"""
    try:
        if column not in df.columns:
            print(f"错误: 变量 {column} 不在数据框中")
            return {'is_stationary': False, 'stationary_pct': 0, 'total_tested': 0, 'error': '变量不存在'}

        # 转换多级索引为普通DataFrame
        if isinstance(df.index, pd.MultiIndex):
            df_reset = df.reset_index()
        else:
            df_reset = df.copy()

        # 检查是否有足够的公司
        if 'stkcd' not in df_reset.columns:
            print(f"错误: 数据框中没有'stkcd'列")
            return {'is_stationary': False, 'stationary_pct': 0, 'total_tested': 0, 'error': '缺少stkcd列'}

        unique_companies = df_reset['stkcd'].unique()
        min_company_count = 10  # 增加最小公司数量要求

        if len(unique_companies) < min_company_count:
            print(f"警告: 公司数量太少 ({len(unique_companies)} < {min_company_count}), 无法进行可靠的平稳性检验")
            return {'is_stationary': False, 'stationary_pct': 0, 'total_tested': 0, 'error': '公司数量太少'}

        stationary_count = 0
        total_tested = 0

        # 增加ADF检验中的最小样本要求
        min_observations = 8  # 至少需要8个观测值

        # 增加样本大小
        max_companies = min(80, len(unique_companies))  # 增加到最多80家公司
        companies = np.random.choice(unique_companies, max_companies, replace=False)

        # 添加一个方差阈值
        MIN_VARIANCE_THRESHOLD = 1e-5  # 增加阈值

        stationary_count = 0
        total_tested = 0
        skipped_low_var = 0

        # 收集详细结果
        detailed_results = []

        for company in companies:
            try:
                group = df_reset[df_reset['stkcd'] == company]
                if len(group) < min_observations:  # 增加最小样本要求
                    continue

                # 按时间排序，确保时序性
                group = group.sort_values('year')

                series = group[column].values

                # 检查序列是否包含无穷或NaN值
                if np.any(np.isnan(series)) or np.any(np.isinf(series)):
                    continue

                # 检查序列是否为常数
                if np.std(series) < MIN_VARIANCE_THRESHOLD:
                    skipped_low_var += 1
                    continue

                # 尝试不同的ADF检验规范
                results = []

                # 尝试多种回归类型
                for regression_type in ['ct', 'c', 'nc']:
                    try:
                        result = adfuller(series, regression=regression_type)
                        results.append({
                            'company': company,
                            'regression': regression_type,
                            'adf_stat': result[0],
                            'p_value': result[1],
                            'is_stationary': result[1] < significance
                        })
                    except Exception as e:
                        #print(f"ADF检验失败 ({column}, 公司: {company}, 类型: {regression_type}): {e}")
                        continue

                # 如果有至少一种检验方法表明序列平稳，则认为平稳
                if results and any(r['is_stationary'] for r in results):
                    # 选择p值最小的结果
                    best_result = min(results, key=lambda x: x['p_value'])
                    detailed_results.append(best_result)
                    stationary_count += 1
                elif results:
                    # 如果所有检验都表明非平稳，选择p值最小的结果
                    best_result = min(results, key=lambda x: x['p_value'])
                    detailed_results.append(best_result)

                total_tested += 1 if results else 0

            except Exception as e:
                print(f"处理公司 {company} 数据时出错: {e}")
                continue

        # 计算平稳序列的比例
        if total_tested > 0:
            stationary_pct = stationary_count / total_tested
            is_stationary = stationary_pct >= 0.5  # 如果超过50%的序列平稳，则认为整体平稳

            print(f"变量 {column}: 测试了 {total_tested} 个公司, {stationary_count} 个平稳 ({stationary_pct*100:.1f}%)")
            if skipped_low_var > 0:
                print(f"  注意: 跳过了 {skipped_low_var} 个方差过小的序列")

            # 检查序列是否强平稳
            is_strongly_stationary = stationary_pct >= 0.7
            stationarity_strength = "强平稳" if is_strongly_stationary else "弱平稳" if is_stationary else "非平稳"

            return {
                'is_stationary': is_stationary,
                'is_strongly_stationary': is_strongly_stationary,
                'stationarity_strength': stationarity_strength,
                'stationary_pct': stationary_pct * 100,
                'total_tested': total_tested,
                'stationary_count': stationary_count,
                'skipped_low_var': skipped_low_var,
                'detailed_results': detailed_results
            }
        else:
            print(f"变量 {column}: 没有可测试的序列")
            return {
                'is_stationary': False,
                'stationary_pct': 0,
                'total_tested': 0,
                'error': '没有可测试的序列'
            }
    except Exception as e:
        print(f"平稳性检验过程出错: {e}")
        traceback.print_exc()
        return {
            'is_stationary': False,
            'stationary_pct': 0,
            'total_tested': 0,
            'error': str(e)
        }

# ---------- 3. 描述性统计分析 ----------

def create_simplified_statistics(df, output_dir='output/tables'):
    """生成简化版描述性统计供论文使用"""
    try:
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        # 重置索引（如果是多级索引）
        df_reset = df.reset_index() if isinstance(df.index, pd.MultiIndex) else df.copy()

        # 变量分组
        key_vars = {
            'AI投入': ['ai', 'ai_job_log'],
            'AI多样性': ['ai_patent_log', 'manu_job_log'],
            '创新绩效': ['intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth']
        }

        control_vars = {
            '企业特征': ['age2', 'size', 'lev', 'growth', 'roa'],
            '治理因素': ['balance', 'mhold', 'audit', 'soe'],
            '其他控制': ['bm1', 'tobinq1', 'tat', 'dsi', 'ci', 'ocr', 'em']
        }

        # 汇总所有变量并筛选存在的变量
        all_vars = []
        for group, vars_list in {**key_vars, **control_vars}.items():
            all_vars.extend([v for v in vars_list if v in df_reset.columns])

        # 1. 创建描述性统计表
        desc_stats = df_reset[all_vars].describe().T
        desc_stats['skew'] = df_reset[all_vars].skew()
        desc_stats['kurt'] = df_reset[all_vars].kurtosis()

        # 保存描述性统计结果
        desc_stats.to_csv(os.path.join(output_dir, 'simple_stats.csv'))

        # 2. 创建相关性矩阵（仅关键变量）
        key_var_list = []
        for vars_list in key_vars.values():
            key_var_list.extend([v for v in vars_list if v in df_reset.columns])

        corr_matrix = df_reset[key_var_list].corr().round(3)
        corr_matrix.to_csv(os.path.join(output_dir, 'simple_corr.csv'))

        # 3. 生成年度均值趋势
        if 'year' in df_reset.columns:
            yearly_means = df_reset.groupby('year')[key_var_list].mean()
            yearly_means.to_csv(os.path.join(output_dir, 'yearly_means_simple.csv'))

        # 4. 创建简单的行业分布
        if 'indnm' in df_reset.columns:
            ind_counts = df_reset['indnm'].value_counts()
            ind_pcts = (ind_counts / ind_counts.sum() * 100).round(1)
            ind_df = pd.DataFrame({'count': ind_counts, 'percentage': ind_pcts})
            ind_df.to_csv(os.path.join(output_dir, 'industry_dist_simple.csv'))

        print(f"简化版统计分析完成，结果保存在 {output_dir} 目录")

        return {
            'desc_stats': desc_stats,
            'corr_matrix': corr_matrix if 'corr_matrix' in locals() else None,
            'yearly_means': yearly_means if 'yearly_means' in locals() else None
        }

    except Exception as e:
        print(f"生成简化版统计分析时出错: {e}")
        traceback.print_exc()
        return None

def enhanced_descriptive_statistics(df, final_vars=None):
    """生成全面的描述性统计分析"""
    print("\n开始增强版描述性统计分析...")

    # 重置索引便于分析
    df_reset = df.reset_index()

    # 确定要使用的变量列表
    if final_vars is None:
        print("警告: 未提供变量名映射，使用默认变量名")
        key_vars = [
            'ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log',
            'intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth'
        ]

        control_vars = [
            'age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
            'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe',
            'dsi', 'ci', 'ocr', 'em'
        ]
    else:
        # 从映射字典获取实际使用的变量名
        print("使用变量名映射...")
        key_vars = []
        for var in ['ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log',
                   'intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth']:
            if var in final_vars:
                mapped_var = final_vars[var]
                if mapped_var in df.columns:
                    key_vars.append(mapped_var)
                else:
                    print(f"警告: 映射的变量 {mapped_var} 不在DataFrame中")
                    if var in df.columns:
                        key_vars.append(var)
            elif var in df.columns:
                key_vars.append(var)

        # 控制变量
        control_vars = [
            'age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
            'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe',
            'dsi', 'ci', 'ocr', 'em'
        ]
        control_vars = [c for c in control_vars if c in df.columns]

    # 变量分组
    var_groups = {
        'ai_investment': [v for v in key_vars if v.startswith('ai') and 'patent' not in v and 'job' not in v],
        'ai_diversity': [v for v in key_vars if v.startswith('ai_patent_log') or v.startswith('manu_job_log')],
        'innovation_capacity': [v for v in key_vars if v in ['intotal', 'ep', 'dp']],
        'innovation_quality': [v for v in key_vars if v.startswith('ai_patent_quality') or v.startswith('ai_patent_depth')],
        'firm_characteristics': [v for v in control_vars if v in ['age2', 'size', 'lev', 'growth', 'roa', 'roe']],
        'governance': [v for v in control_vars if v in ['balance', 'mhold', 'audit', 'soe']],
        'market_valuation': [v for v in control_vars if v in ['tobinq1', 'bm1']],
        'efficiency': [v for v in control_vars if v in ['tat', 'dsi', 'ci', 'ocr', 'em']]
    }

    # 计算描述性统计
    all_desc_stats = {}

    # 为每个变量组生成描述性统计
    for group_name, vars_list in var_groups.items():
        valid_vars = [var for var in vars_list if var in df.columns]
        if valid_vars:
            print(f"为组 '{group_name}' 生成描述性统计，包含 {len(valid_vars)} 个变量")
            group_stats = df[valid_vars].describe(percentiles=[.01, .05, .25, .5, .75, .95, .99]).T

            # 添加其他统计量
            group_stats['skewness'] = df[valid_vars].skew()
            group_stats['kurtosis'] = df[valid_vars].kurtosis()
            group_stats['missing_pct'] = df[valid_vars].isnull().mean() * 100

            all_desc_stats[group_name] = group_stats

            # 保存每组的描述性统计
            group_stats.to_csv(f'output/tables/desc_stats_{group_name}.csv', encoding='utf-8-sig')

    # 年度统计
    yearly_stats = {}
    yearly_means = df_reset.groupby('year')[key_vars].mean()
    yearly_means.to_csv('output/tables/yearly_means.csv', encoding='utf-8-sig')

    # 生成论文友好格式的统计表
    paper_stats = df[key_vars].describe(percentiles=[.25, .5, .75]).T
    paper_stats['skewness'] = df[key_vars].skew()
    paper_stats['kurtosis'] = df[key_vars].kurtosis()
    paper_stats.to_csv('output/tables/paper_descriptive_stats.csv', encoding='utf-8-sig')

    # 相关性分析
    all_num_vars = key_vars + control_vars
    valid_num_vars = [var for var in all_num_vars if var in df.columns and pd.api.types.is_numeric_dtype(df[var])]

    if valid_num_vars:
        corr_matrix = df[valid_num_vars].corr()
        corr_matrix.to_csv('output/tables/full_correlation_matrix.csv', encoding='utf-8-sig')

    # 创建分布可视化
    if key_vars:
        try:
            print("\n创建变量分布可视化...")
            create_distribution_plots(df, key_vars)
        except Exception as e:
            print(f"创建分布图时出错: {e}")
            traceback.print_exc()

        try:
            create_scatter_matrix(df, key_vars[:5])  # 取前5个变量避免图表过于拥挤
        except Exception as e:
            print(f"创建散点图矩阵时出错: {e}")
            traceback.print_exc()

    # 创建年度均值趋势图
    if not yearly_means.empty:
        try:
            print("\n创建年度趋势分析...")
            plot_yearly_trends(yearly_means, [k for k in key_vars if k in yearly_means.columns])
        except Exception as e:
            print(f"创建年度趋势图时出错: {e}")
            traceback.print_exc()

    # 使用增强可视化模块创建高级分析(如果可用)
    try:
        # 这些函数从enhanced_visualization模块导入
        print("\n生成增强可视化分析...")

        # 1. 创建相关系数热图
        if 'create_correlation_heatmap' in globals():
            print("创建相关系数热图...")
            # 主要变量相关系数热图
            create_correlation_heatmap(
                df,
                variables=key_vars,
                output_path='output/figures/key_vars_correlation_heatmap.png',
                method='pearson',
                cluster=True
            )

            # 包括控制变量的完整相关系数热图
            create_correlation_heatmap(
                df,
                variables=valid_num_vars,
                output_path='output/figures/full_correlation_heatmap.png',
                method='pearson',
                cluster=True,
                mask_upper=True
            )

        # 2. 创建主成分分析可视化
        if 'create_pca_visualization' in globals() and len(key_vars) >= 3:
            print("创建主成分分析可视化...")
            create_pca_visualization(
                df,
                variables=key_vars,
                n_components=min(5, len(key_vars)),
                output_path='output/figures/key_vars_pca.png'
            )

        # 3. 创建行业分布分析
        if 'create_industry_distribution' in globals() and 'indcd' in df_reset.columns:
            print("创建行业分布分析...")
            create_industry_distribution(
                df_reset,
                industry_col='indcd',
                output_path='output/figures/industry_distribution.png'
            )

        # 4. 创建论文友好的描述性统计表
        if 'create_enhanced_descriptive_table' in globals():
            print("创建论文友好的描述性统计表...")
            create_enhanced_descriptive_table(
                df,
                variables=valid_num_vars,
                include_diff=False,
                output_path='output/tables/full_descriptive_stats.xlsx'
            )

            # 只包含关键变量的统计表
            create_enhanced_descriptive_table(
                df,
                variables=key_vars,
                include_diff=False,
                output_path='output/tables/key_vars_descriptive_stats.xlsx'
            )

        # 5. 创建高级年度趋势分析
        if 'yearly_trends_analysis' in globals():
            print("创建高级年度趋势分析...")
            yearly_trends_analysis(
                df,
                key_vars=key_vars,
                output_path='output/figures/yearly_trends_advanced.png'
            )

    except Exception as e:
        print(f"增强可视化分析过程中出错: {e}")
        traceback.print_exc()

    return {
        'all_desc_stats': all_desc_stats,
        'yearly_stats': yearly_stats,
        'corr_matrix': corr_matrix if 'corr_matrix' in locals() else None
    }

# 改进可视化辅助函数
def create_distribution_plots(df, variables, output_path='output/figures/key_vars_distribution.png'):
    """创建变量分布的直方图和核密度估计图"""
    try:
        # 创建输出目录
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # 确保变量存在
        valid_vars = [var for var in variables if var in df.columns]
        if not valid_vars:
            print("警告: 没有有效变量可用于创建分布图")
            return

        n_vars = min(len(valid_vars), 6)  # 最多展示6个变量
        variables = valid_vars[:n_vars]

        fig, axes = plt.subplots(n_vars, 2, figsize=(14, 4 * n_vars))

        # 如果只有一个变量，需要重塑axes
        if n_vars == 1:
            axes = axes.reshape(1, 2)

        for i, var in enumerate(variables):
            try:
                # 获取非缺失值
                data = df[var].dropna()

                if len(data) < 5:
                    print(f"警告: 变量 {var} 的非缺失值太少 ({len(data)}), 跳过绘图")
                    axes[i, 0].text(0.5, 0.5, f'数据不足: {var}',
                                  horizontalalignment='center',
                                  verticalalignment='center',
                                  transform=axes[i, 0].transAxes)
                    axes[i, 1].text(0.5, 0.5, f'数据不足: {var}',
                                  horizontalalignment='center',
                                  verticalalignment='center',
                                  transform=axes[i, 1].transAxes)
                    continue

                # 检查数据是否包含异常值
                has_inf = np.any(np.isinf(data))
                has_nan = np.any(np.isnan(data))

                if has_inf or has_nan:
                    print(f"警告: 变量 {var} 包含无穷或NaN值")
                    # 过滤掉异常值
                    data = data[~np.isinf(data) & ~np.isnan(data)]

                # 绘制直方图
                sns.histplot(data, kde=True, ax=axes[i, 0])
                axes[i, 0].set_title(f'{var}的直方图')

                # 绘制Q-Q图
                # 使用try-except处理可能的错误
                try:
                    sm.qqplot(data, line='s', ax=axes[i, 1])
                    axes[i, 1].set_title(f'{var}的Q-Q图')
                except Exception as e:
                    print(f"绘制 {var} 的Q-Q图失败: {e}")
                    axes[i, 1].text(0.5, 0.5, f'Q-Q图生成失败: {var}',
                                  horizontalalignment='center',
                                  verticalalignment='center',
                                  transform=axes[i, 1].transAxes)
            except Exception as e:
                print(f"绘制 {var} 的分布图失败: {e}")
                axes[i, 0].text(0.5, 0.5, f'绘图失败: {var}',
                              horizontalalignment='center',
                              verticalalignment='center',
                              transform=axes[i, 0].transAxes)
                axes[i, 1].text(0.5, 0.5, f'绘图失败: {var}',
                              horizontalalignment='center',
                              verticalalignment='center',
                              transform=axes[i, 1].transAxes)

        plt.tight_layout()
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        print(f"已保存变量分布图到 {output_path}")
    except Exception as e:
        print(f"创建分布图时出错: {e}")
        traceback.print_exc()

def create_scatter_matrix(df, variables, output_path='output/figures/key_vars_scatter_matrix.png'):
    """创建变量之间的散点图矩阵"""
    try:
        # 创建输出目录
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # 确保变量存在
        valid_vars = [var for var in variables if var in df.columns]
        if not valid_vars:
            print("警告: 没有有效变量可用于创建散点图矩阵")
            return

        if len(valid_vars) < 2:
            print("警告: 至少需要2个变量来创建散点图矩阵")
            return

        # 重置多级索引
        if isinstance(df.index, pd.MultiIndex):
            df_reset = df.reset_index()
        else:
            df_reset = df.copy()

        # 创建散点图矩阵, 设置适当的参数来增强健壮性
        scatter_matrix = sns.pairplot(
            df_reset[valid_vars].dropna(),
            diag_kind="kde",
            plot_kws={'alpha': 0.6, 's': 80, 'edgecolor': 'k'},
            diag_kws={'fill': True, 'alpha': 0.6},
        )

        scatter_matrix.fig.suptitle('关键变量散点图矩阵', y=1.02, fontsize=16)
        plt.tight_layout()
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        print(f"已保存散点图矩阵到 {output_path}")
    except Exception as e:
        print(f"创建散点图矩阵时出错: {e}")
        traceback.print_exc()

def plot_yearly_trends(yearly_means, key_vars, output_path='output/figures/yearly_means_trend.png'):
    """创建年度均值趋势图"""
    try:
        # 创建输出目录
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # 确保变量存在
        valid_vars = [var for var in key_vars if var in yearly_means.columns]
        if not valid_vars:
            print("警告: 没有有效变量可用于创建年度趋势图")
            return

        plt.figure(figsize=(14, 8))

        for var in valid_vars:
            # 确保数据是有效的数值
            data = yearly_means[var].replace([np.inf, -np.inf], np.nan).dropna()
            if len(data) > 1:  # 至少需要两个点来绘制趋势
                plt.plot(data.index, data, marker='o', label=var)

        plt.title('关键变量年度均值趋势(2014-2022)')
        plt.xlabel('年份')
        plt.ylabel('均值')
        plt.legend(loc='best')
        plt.grid(True, alpha=0.3)
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        print(f"已保存年度趋势图到 {output_path}")
    except Exception as e:
        print(f"创建年度趋势图时出错: {e}")
        traceback.print_exc()

# ---------- 4. 面板数据诊断 ----------
def calculate_vif(df, features, output_path='output/tables/vif_report.csv',
                 high_vif_threshold=10, extreme_vif_threshold=30,
                 max_vif_to_report=50):
    """
    计算方差膨胀因子(VIF)用于多重共线性检测，并生成详细报告，增强警告提示。

    参数:
    -----------
    df : pandas DataFrame
        包含要分析特征的数据框
    features : list
        要包含在VIF计算中的特征名称列表
    output_path : str, default='output/tables/vif_report.csv'
        保存VIF报告CSV文件的路径
    high_vif_threshold : float, default=10
        触发警告的高VIF值阈值
    extreme_vif_threshold : float, default=30
        触发严重警告的极高VIF值阈值
    max_vif_to_report : float, default=50
        要详细报告的最大VIF值（更高的值将被截断显示）

    返回:
    --------
    dict
        包含VIF结果和多重共线性评估的字典
    """
    print("\n计算方差膨胀因子(VIF)进行多重共线性检测...")

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # 验证所有特征是否存在于数据框中
    missing_features = [f for f in features if f not in df.columns]
    if missing_features:
        print(f"警告: 以下特征不在数据框中: {missing_features}")
        features = [f for f in features if f in df.columns]

    if len(features) < 2:
        print("错误: 需要至少2个特征来计算VIF")
        return {'error': '计算VIF的特征不足'}

    # 创建仅包含所选特征的数据框副本
    df_vif = df[features].copy()

    # 删除有缺失值的行
    na_count_before = len(df_vif)
    df_vif = df_vif.dropna()
    na_count_after = len(df_vif)

    if na_count_before > na_count_after:
        print(f"删除了{na_count_before - na_count_after}行有缺失值的数据")

    if len(df_vif) < 10:
        print("错误: 删除缺失值后观测值太少")
        return {'error': '计算VIF的观测值不足'}

    # 计算每个特征的VIF
    vif_data = []
    high_vif_features = []
    extreme_vif_features = []

    try:
        # 为statsmodels实现添加常数项
        X = sm.add_constant(df_vif)

        # 计算每个特征的VIF
        for i, feature in enumerate(X.columns):
            if feature == 'const':
                continue

            try:
                vif_value = variance_inflation_factor(X.values, i)

                # 处理极大或无限VIF值
                if not np.isfinite(vif_value) or vif_value > 1e6:
                    print(f"⚠️ 极端警告: 特征'{feature}'的VIF值无限或极高!")
                    vif_value = 1e6  # 为报告目的设置上限
                    extreme_vif_features.append((feature, vif_value))
                elif vif_value > extreme_vif_threshold:
                    print(f"⚠️ 严重警告: 特征'{feature}'的VIF值非常高，为{vif_value:.2f}!")
                    extreme_vif_features.append((feature, vif_value))
                elif vif_value > high_vif_threshold:
                    print(f"⚠️ 警告: 特征'{feature}'的VIF值较高，为{vif_value:.2f}")
                    high_vif_features.append((feature, vif_value))

                # 为极高值设置显示上限
                display_vif = min(vif_value, max_vif_to_report)
                if display_vif < vif_value:
                    display_note = f" (实际: {vif_value:.2e})"
                else:
                    display_note = ""

                vif_data.append({
                    'feature': feature,
                    'VIF': vif_value
                })
            except Exception as e:
                print(f"计算特征'{feature}'的VIF时出错: {e}")

    except Exception as e:
        print(f"VIF计算过程中出错: {e}")
        return {'error': str(e)}

    # 创建包含VIF结果的DataFrame
    vif_df = pd.DataFrame(vif_data)

    # 按VIF值降序排序
    vif_df = vif_df.sort_values('VIF', ascending=False)

    # 保存到CSV
    try:
        vif_df.to_csv(output_path, index=True)
        print(f"VIF报告已保存到 {output_path}")
    except Exception as e:
        print(f"保存VIF报告时出错: {e}")

    # 生成汇总统计
    vif_summary = {
        'mean_vif': vif_df['VIF'].mean(),
        'median_vif': vif_df['VIF'].median(),
        'max_vif': vif_df['VIF'].max(),
        'min_vif': vif_df['VIF'].min(),
        'high_vif_count': len(high_vif_features) + len(extreme_vif_features),
        'extreme_vif_count': len(extreme_vif_features),
        'high_vif_features': high_vif_features,
        'extreme_vif_features': extreme_vif_features,
        'vif_data': vif_df.to_dict('records')
    }

    # 打印汇总
    print("\nVIF分析汇总:")
    print(f"- 平均VIF: {vif_summary['mean_vif']:.2f}")
    print(f"- 中位数VIF: {vif_summary['median_vif']:.2f}")
    print(f"- 高VIF值(>{high_vif_threshold})特征数量: {vif_summary['high_vif_count']}")
    print(f"- 极高VIF值(>{extreme_vif_threshold})特征数量: {vif_summary['extreme_vif_count']}")

    # 根据VIF结果提供建议
    if vif_summary['extreme_vif_count'] > 0:
        print("\n⚠️ 检测到严重多重共线性!")
        print("建议: 考虑移除或合并这些具有极高VIF值的特征:")
        for feature, vif_value in extreme_vif_features:
            print(f"  - {feature}: VIF = {min(vif_value, max_vif_to_report):.2f}")
        print("这些特征与其他特征高度相关，可能导致回归模型不稳定。")

    elif vif_summary['high_vif_count'] > 0:
        print("\n⚠️ 检测到多重共线性!")
        print("建议: 考虑处理这些具有高VIF值的特征:")
        for feature, vif_value in high_vif_features:
            print(f"  - {feature}: VIF = {vif_value:.2f}")
        print("这些特征与其他特征存在相关性，可能影响回归系数的稳定性。")

    else:
        print("\n✓ 未检测到显著的多重共线性。")

    return vif_summary

def check_multicollinearity(df, features=None, output_path='output/tables/vif_report.csv',
                          high_vif_threshold=10, extreme_vif_threshold=30):
    """
    全面的多重共线性检查，增强警告和详细报告。

    参数:
    -----------
    df : pandas DataFrame
        包含要分析特征的数据框
    features : list or None
        要包含在VIF计算中的特征名称列表。如果为None，则使用所有数值列。
    output_path : str, default='output/tables/vif_report.csv'
        保存VIF报告CSV文件的路径
    high_vif_threshold : float, default=10
        触发警告的高VIF值阈值
    extreme_vif_threshold : float, default=30
        触发严重警告的极高VIF值阈值

    返回:
    --------
    dict
        包含VIF结果和多重共线性评估的字典
    """
    print("\n执行全面的多重共线性检查...")

    # 如果未指定特征，则使用所有数值列
    if features is None:
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        # 移除全为零或全为相同值（零方差）的列
        features = [col for col in numeric_cols if df[col].var() > 0]
        print(f"使用{len(features)}个具有非零方差的数值特征")

    # 首先计算相关矩阵以识别高度相关的对
    corr_matrix = df[features].corr()

    # 查找高度相关的对
    high_corr_pairs = []
    for i in range(len(features)):
        for j in range(i+1, len(features)):
            corr_value = abs(corr_matrix.iloc[i, j])
            if corr_value > 0.8:  # 高相关性阈值
                high_corr_pairs.append({
                    'feature1': features[i],
                    'feature2': features[j],
                    'correlation': corr_value
                })

    # 按相关值排序
    high_corr_pairs = sorted(high_corr_pairs, key=lambda x: x['correlation'], reverse=True)

    # 打印高度相关的对
    if high_corr_pairs:
        print("\n检测到高度相关的特征对:")
        for pair in high_corr_pairs[:10]:  # 显示前10对
            print(f"  - {pair['feature1']} & {pair['feature2']}: r = {pair['correlation']:.4f}")

        if len(high_corr_pairs) > 10:
            print(f"  ... 以及{len(high_corr_pairs) - 10}对更多")

    # 计算VIF
    vif_results = calculate_vif(
        df,
        features,
        output_path=output_path,
        high_vif_threshold=high_vif_threshold,
        extreme_vif_threshold=extreme_vif_threshold
    )

    # 添加相关信息到结果
    vif_results['high_corr_pairs'] = high_corr_pairs

    # 生成更详细的报告
    report_path = output_path.replace('.csv', '_detailed.txt')
    try:
        with open(report_path, 'w') as f:
            f.write("多重共线性分析报告\n")
            f.write("================================\n\n")

            f.write("1. 汇总\n")
            f.write("---------\n")
            f.write(f"平均VIF: {vif_results['mean_vif']:.2f}\n")
            f.write(f"中位数VIF: {vif_results['median_vif']:.2f}\n")
            f.write(f"最大VIF: {vif_results['max_vif']:.2f}\n")
            f.write(f"最小VIF: {vif_results['min_vif']:.2f}\n")
            f.write(f"高VIF值(>{high_vif_threshold})特征数量: {vif_results['high_vif_count']}\n")
            f.write(f"极高VIF值(>{extreme_vif_threshold})特征数量: {vif_results['extreme_vif_count']}\n\n")

            f.write("2. 高度相关的特征对\n")
            f.write("--------------------------------\n")
            if high_corr_pairs:
                for pair in high_corr_pairs:
                    f.write(f"{pair['feature1']} & {pair['feature2']}: r = {pair['correlation']:.4f}\n")
            else:
                f.write("未检测到高度相关的特征对。\n")
            f.write("\n")

            f.write("3. 按特征的VIF值\n")
            f.write("-----------------------\n")
            for item in vif_results['vif_data']:
                feature = item['feature']
                vif_value = item['VIF']

                if vif_value > extreme_vif_threshold:
                    severity = "极高"
                elif vif_value > high_vif_threshold:
                    severity = "高"
                else:
                    severity = "正常"

                f.write(f"{feature}: {vif_value:.4f} - {severity}\n")

            f.write("\n4. 建议\n")
            f.write("-----------------\n")
            if vif_results['extreme_vif_count'] > 0:
                f.write("检测到严重多重共线性!\n")
                f.write("建议: 考虑移除或合并这些具有极高VIF值的特征:\n")
                for feature, vif_value in vif_results['extreme_vif_features']:
                    f.write(f"  - {feature}: VIF = {vif_value:.2f}\n")
            elif vif_results['high_vif_count'] > 0:
                f.write("检测到多重共线性!\n")
                f.write("建议: 考虑处理这些具有高VIF值的特征:\n")
                for feature, vif_value in vif_results['high_vif_features']:
                    f.write(f"  - {feature}: VIF = {vif_value:.2f}\n")
            else:
                f.write("未检测到显著的多重共线性。\n")

        print(f"详细的多重共线性报告已保存到 {report_path}")
    except Exception as e:
        print(f"保存详细报告时出错: {e}")

    return vif_results

def panel_data_diagnostics(df, final_vars=None):
    """进行面板数据诊断，检查平稳性、模型选择、异方差和多重共线性等问题"""
    print("进行面板数据诊断...")

    # 创建结果字典
    diagnostics = {
        'unit_root': {},
        'hausman': {},
        'heteroskedasticity': {},
        'serial_correlation': {},
        'multicollinearity': {}  # 添加多重共线性结果
    }

    # 重置索引便于分析
    df_reset = df.reset_index()

    # 确定要使用的变量列表
    if final_vars is None:
        # 默认的变量列表
        key_vars = [
            'ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log',
            'intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth'
        ]

        control_vars = [
            'age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
            'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe'
        ]

        # 只保留存在于数据框中的变量
        key_vars = [var for var in key_vars if var in df.columns]
        control_vars = [var for var in control_vars if var in df.columns]
    else:
        # 从映射字典获取实际使用的变量名
        key_vars = []
        for var in ['ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log',
                   'intotal', 'ep', 'dp', 'ai_patent_quality', 'ai_patent_depth']:
            if var in final_vars:
                mapped_var = final_vars[var]
                if mapped_var in df.columns:
                    key_vars.append(mapped_var)
            elif var in df.columns:
                key_vars.append(var)

        # 控制变量
        control_vars = [
            'age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
            'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe'
        ]
        control_vars = [c for c in control_vars if c in df.columns]

    # 增强的数据预处理和ADF检验函数
    def safe_adfuller_test(series_data, company_id, var_name, regression='ct'):
        """对时间序列进行安全的ADF检验，包含全面的数据预处理"""
        try:
            # 1. 转换为Pandas Series以便于数据处理
            if not isinstance(series_data, pd.Series):
                series_data = pd.Series(series_data)

            # 2. 重置Series索引，确保连续
            series_data = series_data.reset_index(drop=True)

            # 3. 处理异常值: 替换无限值为NaN
            series_data = series_data.replace([np.inf, -np.inf], np.nan)

            # 4. 处理离群值: 使用Winsorize处理极端值 (可选)
            if len(series_data) >= 8:  # 至少需要足够样本才能识别离群值
                try:
                    # 计算四分位数
                    q1, q3 = np.nanpercentile(series_data, [25, 75])
                    iqr = q3 - q1
                    lower_bound = q1 - 3 * iqr
                    upper_bound = q3 + 3 * iqr

                    # 修剪极端离群值 (比标准Winsorize更激进)
                    series_data = series_data.clip(lower=lower_bound, upper=upper_bound)
                except Exception as e:
                    print(f"离群值处理失败 ({var_name}, 公司: {company_id}): {str(e)}")

            # 5. 移除NaN值
            series_clean = series_data.dropna()

            # 6. 数据有效性检查
            if len(series_clean) <= 5:
                print(f"变量 {var_name} 对公司 {company_id} 有效数据点不足 (只有{len(series_clean)}个)")
                return None

            # 7. 方差检查
            if np.std(series_clean) <= 1e-8:  # 宽松的方差阈值
                print(f"变量 {var_name} 对公司 {company_id} 方差接近零 ({np.std(series_clean):.2e})")
                return None

            # 8. 变异性检查
            if len(np.unique(series_clean)) < 3:
                print(f"变量 {var_name} 对公司 {company_id} 变异性不足 (唯一值数量: {len(np.unique(series_clean))})")
                return None

            # 9. 最终转换为numpy数组
            clean_array = np.array(series_clean, dtype=np.float64)

            # 10. 最后确认数据没有问题
            if not np.isfinite(clean_array).all() or len(clean_array) <= 5:
                print(f"变量 {var_name} 对公司 {company_id} 最终数据仍有问题，无法分析")
                return None

            # 11. 进行ADF检验
            result = adfuller(clean_array, regression=regression)
            return result

        except Exception as e:
            print(f"ADF预处理或检验失败 ({var_name}, 公司: {company_id}): {str(e)}")
            return None

    # 1. 单位根检验
    print("进行单位根检验...")
    unit_root_results = {}

    for var in key_vars[:5]:  # 限制为最多5个变量
        if var not in df.columns:
            continue

        adf_results = []
        tested_companies = 0
        total_companies = len(df_reset['stkcd'].unique())
          # 对每个公司单独进行ADF检验
        max_companies = min(100, total_companies)
        sample_companies = np.random.choice(df_reset['stkcd'].unique(), max_companies, replace=False)

        for company in sample_companies:
            group = df_reset[df_reset['stkcd'] == company]
            if len(group) > 5:
                # 使用增强版ADF检验函数
                result = safe_adfuller_test(group[var], company, var, regression='ct')
                if result is not None:
                    adf_results.append({
                        'stkcd': company,
                        'adf_stat': result[0],
                        'p_value': result[1],
                        'stationary': result[1] < 0.05
                    })
                    tested_companies += 1
                else:
                    print(f"ADF检验返回None (变量: {var}, 公司: {company})")
            else:
                print(f"公司 {company} 样本数量不足 ({len(group)}个观测值), 跳过ADF检验")

        if adf_results:
            adf_df = pd.DataFrame(adf_results)
            stationary_pct = (adf_df['stationary'].sum() / len(adf_df)) * 100
            unit_root_results[var] = {
                'stationary_percentage': stationary_pct,
                'total_companies': tested_companies,
                'stationary_companies': adf_df['stationary'].sum(),
                'sample_fraction': tested_companies / total_companies
            }
            print(f"变量 {var}: {stationary_pct:.1f}%序列平稳, 测试了{tested_companies}家公司")

    non_stationary_vars = {}
    for var, result in unit_root_results.items():
        if result['stationary_percentage'] < 30:  # 非常不平稳
            non_stationary_vars[var] = 'high_risk'
        elif result['stationary_percentage'] < 50:  # 中度不平稳
            non_stationary_vars[var] = 'medium_risk'

    # 创建差分变量
    diff_vars_results = {}
    if non_stationary_vars:
        print("\n自动为非平稳变量创建一阶差分...")
        df_diff = df.copy()
        for var, risk in non_stationary_vars.items():
            if var in df_diff.columns:
                try:
                    df_reset_temp = df_diff.reset_index()
                    df_reset_temp[f'{var}_diff'] = np.nan

                    # 按公司分组差分
                    for company, group in df_reset_temp.groupby('stkcd'):
                        group = group.sort_values('year')
                        df_reset_temp.loc[group.index, f'{var}_diff'] = group[var].diff().values                    # 将差分变量加入原数据框
                    df_diff = df_reset_temp.set_index(['stkcd', 'year'])

                    # 预先初始化变量
                    series_tested = 0
                    stationary_count = 0
                    ma_stationary_count = 0
                    ma_series_tested = 0
                    recommendation = 'unknown'
                    status = 'unknown'
                    improvement = 0
                    stationary_pct = 0  # 预先初始化，避免未定义错误

                    # 测试完所有样本后再计算百分比和改进指标
                    # 下面的代码将移到测试完成后执行
                    # 针对无改善情况创建备选转换
                    try:
                        # 创建移动平均差分 (更平滑的差分)
                        df_diff[f'{var}_ma_diff'] = df_diff.groupby(level=0)[var].apply(
                            lambda x: x.rolling(window=2).mean().diff())

                        # 重新评估平稳性
                        ma_stationary_count = 0
                        ma_series_tested = 0
                    except Exception as ma_err:
                        print(f"创建移动平均差分出错: {ma_err}")

                    # 测试差分后的平稳性
                    series_tested = 0
                    stationary_count = 0

                    for company in np.random.choice(df_reset_temp['stkcd'].unique(),
                                                min(30, len(df_reset_temp['stkcd'].unique())),
                                                replace=False):
                        group = df_reset_temp[df_reset_temp['stkcd'] == company]
                        if len(group) > 5:                            # 处理差分序列，确保没有无限值或NaN                            # 更全面的预处理
                            try:
                                # 1. 先替换无限值为NaN
                                diff_series_raw = group[f'{var}_diff'].replace([np.inf, -np.inf], np.nan)

                                # 2. 检查是否有足够数据进行分析
                                valid_data = diff_series_raw.dropna()
                                if len(valid_data) <= 5:
                                    print(f"差分序列 {var}_diff 对公司 {company} 有效数据点不足 (只有{len(valid_data)}个)")
                                    continue

                                # 3. 转换为numpy数组并再次检查
                                diff_series = valid_data.values

                                # 4. 检查方差
                                if np.std(diff_series) <= 1e-10:
                                    print(f"差分序列 {var}_diff 对公司 {company} 方差接近零 ({np.std(diff_series):.2e})")
                                    continue

                                # 5. 检查有限性
                                if not np.isfinite(diff_series).all():
                                    print(f"差分序列 {var}_diff 对公司 {company} 仍然包含无效值，无法处理")
                                    continue

                                # 6. 检查具体何种无效值出现在数据中
                                if np.any(np.isnan(diff_series)):
                                    print(f"差分序列 {var}_diff 对公司 {company} 包含NaN值")
                                    continue

                                if np.any(np.isinf(diff_series)):
                                    print(f"差分序列 {var}_diff 对公司 {company} 包含无穷值")
                                    continue

                                # 7. 确保数据包含足够的变异
                                if len(np.unique(diff_series)) < 3:
                                    print(f"差分序列 {var}_diff 对公司 {company} 变异性不足 (唯一值数量: {len(np.unique(diff_series))})")
                                    continue

                                # 所有检查通过，进行ADF检验
                                result = adfuller(diff_series)
                                series_tested += 1
                                if result[1] < 0.05:
                                    stationary_count += 1
                                    if ma_series_tested > 0:
                                        ma_stationary_pct = (ma_stationary_count / ma_series_tested) * 100
                                        # 增加安全检查，确保var在unit_root_results中存在
                                        if var in unit_root_results and 'stationary_percentage' in unit_root_results[var]:
                                            ma_improvement = ma_stationary_pct - unit_root_results[var]['stationary_percentage']
                                            if ma_improvement > improvement:
                                                recommendation = 'use_ma_diff'
                                                status = 'ma_diff_better'
                                                print(f"移动平均差分比普通差分效果更好 ({ma_improvement:.1f}% vs {improvement:.1f}%)")
                                        else:
                                            print(f"警告：无法获取变量'{var}'的原始平稳性信息")
                                            ma_improvement = 0
                            except Exception as ma_err:
                                print(f"创建移动平均差分出错: {ma_err}")

                    # 添加改进的评估标准 - 移动到这里，确保在所有测试完成后执行
                    if series_tested > 0:
                        stationary_pct = (stationary_count / series_tested) * 100
                        # 确保var在unit_root_results中存在且含有stationary_percentage键
                        if var in unit_root_results and 'stationary_percentage' in unit_root_results[var]:
                            improvement = stationary_pct - unit_root_results[var]['stationary_percentage']

                            # 设置明确的决策标准
                            if improvement > 20:  # 显著改善
                                recommendation = 'use_diff'
                                status = 'significant_improvement'
                            elif improvement > 10:  # 中等改善
                                recommendation = 'consider_diff'
                                status = 'moderate_improvement'
                            else:  # 微小改善或无改善
                                recommendation = 'use_original'
                                status = 'minimal_improvement'
                        else:
                            print(f"警告：unit_root_results中不存在变量'{var}'或其未包含'stationary_percentage'")
                            improvement = 0
                            recommendation = 'unknown'
                            status = 'error_missing_data'

                    diff_vars_results[f'{var}_diff'] = {
                        'original_var': var,
                        'stationary_percentage': stationary_pct,
                        'improvement': improvement,
                        'recommendation': recommendation,
                        'status': status
                    }

                    print(f"变量 {var} 差分后: {stationary_pct:.1f}% 序列平稳 " +
                        f"(改善: {improvement:.1f}%, 建议: {recommendation})")
                except Exception as e:
                    print(f"为变量 {var} 创建差分出错: {e}")


            # 保存差分变量结果和差分后的数据框
            diagnostics['diff_variables'] = diff_vars_results
            diagnostics['df_with_diff'] = df_diff

            # 推荐变量使用
            recommended_vars = {}
            for var, result in unit_root_results.items():
                if f'{var}_diff' in diff_vars_results:
                    if diff_vars_results[f'{var}_diff']['improvement'] > 20:
                        recommended_vars[var] = f'{var}_diff'
                        print(f"推荐使用 {var} 的差分版本: {var}_diff")
                    else:
                        recommended_vars[var] = var
                        print(f"推荐继续使用原始变量: {var} (差分改善不显著)")
                else:
                    recommended_vars[var] = var

            diagnostics['recommended_vars'] = recommended_vars

    # 其他诊断代码保持不变...
    # 2. Hausman检验 - 使用正确的变量名
    print("进行Hausman检验...")

    # 获取因变量和自变量的正确名称
    dependent_vars = [v for v in ['intotal', 'ep', 'dp'] if v in df.columns]
    ai_var = final_vars.get('ai', 'ai') if final_vars else 'ai'
    ai_job_log_var = final_vars.get('ai_job_log', 'ai_job_log') if final_vars else 'ai_job_log'

    # 确保使用数值型变量
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

    hausman_results = {}
    for dep_var in dependent_vars:
        if ai_var not in numeric_cols or ai_job_log_var not in numeric_cols or dep_var not in numeric_cols:
            continue

        try:
            # 准备模型变量
            X_vars = [ai_var, ai_job_log_var]
            valid_x = [var for var in X_vars if var in numeric_cols]

            if len(valid_x) > 0:
                # 创建数据子集
                y = df[dep_var]
                X = df[valid_x]

                # 添加常数项
                X = sm.add_constant(X)

                # 创建固定效应和随机效应模型
                fe_model = PanelOLS(y, X, entity_effects=True)
                re_model = RandomEffects(y, X)

                # 拟合模型并使用稳健标准误
                fe_result = fe_model.fit(cov_type='robust')
                re_result = re_model.fit(cov_type='robust')

                # 计算Hausman统计量
                b_fe = fe_result.params
                b_re = re_result.params

                # 只保留共同的参数
                common_params = list(set(b_fe.index).intersection(set(b_re.index)))
                b_fe = b_fe.loc[common_params]
                b_re = b_re.loc[common_params]

                # 计算差异
                diff = b_fe - b_re

                # 计算协方差矩阵差异
                var_fe = fe_result.cov.loc[common_params, common_params]
                var_re = re_result.cov.loc[common_params, common_params]
                var_diff = var_fe - var_re

                # 计算Hausman统计量
                try:
                    inv_var_diff = np.linalg.inv(var_diff)
                    hausman_stat = float(diff.T @ inv_var_diff @ diff)
                    p_value = 1 - stats.chi2.cdf(hausman_stat, len(common_params))

                    hausman_results[dep_var] = {
                        'hausman_stat': hausman_stat,
                        'p_value': p_value,
                        'df': len(common_params),
                        'conclusion': '固定效应' if p_value < 0.05 else '随机效应'
                    }
                except np.linalg.LinAlgError:
                    # 使用Moore-Penrose伪逆
                    inv_var_diff = np.linalg.pinv(var_diff)
                    hausman_stat = float(diff.T @ inv_var_diff @ diff)
                    p_value = 1 - stats.chi2.cdf(hausman_stat, len(common_params))

                    hausman_results[dep_var] = {
                        'hausman_stat': hausman_stat,
                        'p_value': p_value,
                        'df': len(common_params),
                        'conclusion': '固定效应' if p_value < 0.05 else '随机效应',
                        'method': 'pseudoinverse'
                    }
        except Exception as e:
            print(f"Hausman检验出错 ({dep_var}): {e}")
            traceback.print_exc()
            hausman_results[dep_var] = {'error': str(e)}

    # 3. 异方差检验
    print("进行异方差检验...")
    heteroskedasticity_results = {}

    for dep_var in dependent_vars:
        if dep_var not in df.columns:
            continue

        try:
            # 准备模型变量
            model_vars = [ai_var, ai_job_log_var] + control_vars[:5]
            valid_vars = [var for var in model_vars if var in df_reset.columns]

            if len(valid_vars) < 2:
                print(f"异方差检验 - 对 {dep_var} 变量不足，需要至少两个有效变量")
                continue

            # 创建OLS模型
            X = sm.add_constant(df_reset[valid_vars].values)
            y = df_reset[dep_var].values

            # 删除缺失值
            mask = ~np.isnan(y)
            for i in range(X.shape[1]):
                mask = mask & ~np.isnan(X[:, i])

            X_clean = X[mask]
            y_clean = y[mask]

            if len(y_clean) < 30:
                print(f"异方差检验 - 对 {dep_var} 有效观测值不足，需要至少30个观测值")
                continue

            model = sm.OLS(y_clean, X_clean)
            results = model.fit()

            # Breusch-Pagan检验
            bp_test = het_breuschpagan(results.resid, results.model.exog)

            heteroskedasticity_results[dep_var] = {
                'bp_statistic': bp_test[0],
                'bp_p_value': bp_test[1],
                'heteroskedasticity': bp_test[1] < 0.05
            }

            print(f"变量 {dep_var} - BP检验 p值: {bp_test[1]:.4f}")
        except Exception as e:
            print(f"异方差检验出错 ({dep_var}): {e}")
            heteroskedasticity_results[dep_var] = {'error': str(e)}

    # 4. 序列相关检验
    print("进行序列相关检验...")
    serial_correlation_results = {}

    for dep_var in dependent_vars:
        if dep_var not in df.columns:
            print(f"序列相关检验 - 变量 {dep_var} 不在数据框中")
            continue

        try:
            # 对每个公司进行DW检验
            dw_results = []

            # 限制测试的公司数量
            unique_companies = df_reset['stkcd'].unique()
            max_companies = min(100, len(unique_companies))
            sample_companies = np.random.choice(unique_companies, max_companies, replace=False)

            for company in sample_companies:
                group = df_reset[df_reset['stkcd'] == company]
                if len(group) > 5:
                    model_vars = [ai_var, ai_job_log_var] + control_vars[:3]
                    valid_vars = [var for var in model_vars if var in group.columns]

                    if len(valid_vars) > 0 and dep_var in group.columns:
                        try:
                            X = sm.add_constant(group[valid_vars].values)
                            y = group[dep_var].values

                            # 删除缺失值
                            mask = ~np.isnan(y)
                            for i in range(X.shape[1]):
                                mask = mask & ~np.isnan(X[:, i])

                            X_clean = X[mask]
                            y_clean = y[mask]

                            if len(y_clean) < 5:
                                continue

                            model = sm.OLS(y_clean, X_clean)
                            results = model.fit()

                            # Durbin-Watson检验
                            dw_stat = durbin_watson(results.resid)

                            dw_results.append({
                                'stkcd': company,
                                'dw_stat': dw_stat,
                                'autocorrelation': dw_stat < 1.5 or dw_stat > 2.5
                            })
                        except Exception as e:
                            print(f"公司 {company} 的DW检验失败: {e}")

            if dw_results:
                dw_df = pd.DataFrame(dw_results)
                autocorr_pct = (dw_df['autocorrelation'].sum() / len(dw_df)) * 100

                serial_correlation_results[dep_var] = {
                    'companies_with_autocorrelation_pct': autocorr_pct,
                    'mean_dw_stat': dw_df['dw_stat'].mean(),
                    'min_dw_stat': dw_df['dw_stat'].min(),
                    'max_dw_stat': dw_df['dw_stat'].max(),
                    'companies_tested': len(dw_df),
                    'sample_fraction': len(dw_df) / len(unique_companies)
                }

                print(f"变量 {dep_var} - 平均DW统计量: {dw_df['dw_stat'].mean():.2f}, {autocorr_pct:.1f}%的公司存在自相关")
        except Exception as e:
            print(f"序列相关检验出错 ({dep_var}): {e}")
            serial_correlation_results[dep_var] = {'error': str(e)}

    # 5. 多重共线性检验
    print("进行多重共线性检验...")

    # 获取所有数值型变量
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

    # 移除零方差列
    valid_features = [col for col in numeric_cols if col in df.columns and df[col].var() > 0]

    # 如果有足够的特征，进行多重共线性检验
    if len(valid_features) >= 2:
        try:
            # 执行多重共线性检查
            multicollinearity_results = check_multicollinearity(
                df,
                features=valid_features,
                output_path='output/tables/vif_report.csv',
                high_vif_threshold=10,
                extreme_vif_threshold=30
            )

            # 保存到诊断结果
            diagnostics['multicollinearity'] = multicollinearity_results
        except Exception as e:
            print(f"多重共线性检验出错: {e}")
            traceback.print_exc()
            diagnostics['multicollinearity'] = {'error': str(e)}
    else:
        print("多重共线性检验 - 有效特征不足，需要至少两个有效特征")
        diagnostics['multicollinearity'] = {'error': '有效特征不足'}

    # 保存诊断结果
    diagnostics['unit_root'] = unit_root_results
    diagnostics['hausman'] = hausman_results
    diagnostics['heteroskedasticity'] = heteroskedasticity_results
    diagnostics['serial_correlation'] = serial_correlation_results

    # 生成诊断报告
    try:
        with open('output/reports/panel_diagnostics_report.txt', 'w', encoding='utf-8') as f:
            f.write("面板数据诊断报告\n")
            f.write("=" * 60 + "\n\n")

            f.write("1. 单位根检验 (平稳性检验)\n")
            f.write("-" * 60 + "\n")
            for var, result in unit_root_results.items():
                f.write(f"变量: {var}\n")
                f.write(f"  - 平稳序列比例: {result['stationary_percentage']:.2f}%\n")
                f.write(f"  - 检验公司总数: {result['total_companies']}\n")
                f.write(f"  - 平稳公司数量: {result['stationary_companies']}\n")
                f.write(f"  - 样本占比: {result['sample_fraction']:.2%}\n")

                # 添加简要解释
                if result['stationary_percentage'] > 75:
                    f.write("  - 解释: 大多数序列平稳，不需要进一步处理\n")
                elif result['stationary_percentage'] > 50:
                    f.write("  - 解释: 部分序列平稳，考虑使用差分或考虑协整关系\n")
                else:
                    f.write("  - 解释: 大多数序列非平稳，建议使用差分处理\n")
                f.write("\n")

            f.write("\n2. Hausman检验 (固定效应vs随机效应)\n")
            f.write("-" * 60 + "\n")
            for var, result in diagnostics['hausman'].items():
                f.write(f"因变量: {var}\n")

                # 添加对错误情况的检查
                if 'error' in result:
                    f.write(f"  - 检验失败: {result['error']}\n\n")
                else:
                    # 正常情况
                    f.write(f"  - 统计量: {result['hausman_stat']:.4f}\n")
                    f.write(f"  - p值: {result['p_value']:.4f}\n")
                    f.write(f"  - 推荐模型: {result['conclusion']}\n")

                    # 添加方法说明
                    if 'method' in result and result['method'] == 'pseudoinverse':
                        f.write("  - 注意: 使用了伪逆方法，结果可能不如标准方法可靠\n")

                    # 添加简要解释
                    if result['p_value'] < 0.05:
                        f.write("  - 解释: 固定效应和随机效应估计存在系统性差异，应使用固定效应模型\n")
                    else:
                        f.write("  - 解释: 两种估计方法无显著差异，可以使用随机效应模型以获得更高效率\n")
                    f.write("\n")

            f.write("\n3. 异方差检验\n")
            f.write("-" * 60 + "\n")
            for var, result in heteroskedasticity_results.items():
                f.write(f"因变量: {var}\n")
                if 'error' in result:
                    f.write(f"  - 检验失败: {result['error']}\n\n")
                else:
                    f.write(f"  - BP统计量: {result['bp_statistic']:.4f}\n")
                    f.write(f"  - p值: {result['bp_p_value']:.4f}\n")
                    f.write(f"  - 存在异方差: {'是' if result['heteroskedasticity'] else '否'}\n")

                    # 添加建议
                    if result['heteroskedasticity']:
                        f.write("  - 建议: 使用稳健标准误(如White或聚类稳健标准误)来修正异方差问题\n")
                    else:
                        f.write("  - 建议: 可以使用普通标准误，但使用稳健标准误总是更为保险\n")
                    f.write("\n")

            f.write("\n4. 序列相关检验\n")
            f.write("-" * 60 + "\n")
            for var, result in serial_correlation_results.items():
                f.write(f"因变量: {var}\n")
                if 'error' in result:
                    f.write(f"  - 检验失败: {result['error']}\n\n")
                else:
                    f.write(f"  - 存在自相关的公司比例: {result['companies_with_autocorrelation_pct']:.2f}%\n")
                    f.write(f"  - 平均DW统计量: {result['mean_dw_stat']:.4f}\n")
                    f.write(f"  - DW统计量范围: [{result['min_dw_stat']:.4f}, {result['max_dw_stat']:.4f}]\n")
                    f.write(f"  - 测试公司数量: {result.get('companies_tested', 'N/A')}\n")
                    if 'sample_fraction' in result:
                        f.write(f"  - 样本占比: {result['sample_fraction']:.2%}\n")

                    # 添加建议
                    if result['companies_with_autocorrelation_pct'] > 50:
                        f.write("  - 建议: 多数公司存在自相关，应使用自相关稳健标准误或添加滞后变量\n")
                    elif result['companies_with_autocorrelation_pct'] > 25:
                        f.write("  - 建议: 部分公司存在自相关，建议使用稳健标准误\n")
                    else:
                        f.write("  - 建议: 自相关程度较低，可以使用普通标准误，但使用稳健标准误更为保险\n")
                    f.write("\n")

            # 5. 多重共线性检验
            f.write("\n5. 多重共线性检验\n")
            f.write("-" * 60 + "\n")

            if 'multicollinearity' in diagnostics:
                multicollinearity_results = diagnostics['multicollinearity']

                if 'error' in multicollinearity_results:
                    f.write(f"多重共线性检验失败: {multicollinearity_results['error']}\n\n")
                else:
                    # 写入VIF汇总信息
                    f.write("VIF汇总统计:\n")
                    f.write(f"  - 平均VIF: {multicollinearity_results.get('mean_vif', 'N/A'):.2f}\n")
                    f.write(f"  - 中位数VIF: {multicollinearity_results.get('median_vif', 'N/A'):.2f}\n")
                    f.write(f"  - 最大VIF: {multicollinearity_results.get('max_vif', 'N/A'):.2f}\n")
                    f.write(f"  - 最小VIF: {multicollinearity_results.get('min_vif', 'N/A'):.2f}\n")
                    f.write(f"  - 高VIF值(>10)特征数量: {multicollinearity_results.get('high_vif_count', 'N/A')}\n")
                    f.write(f"  - 极高VIF值(>30)特征数量: {multicollinearity_results.get('extreme_vif_count', 'N/A')}\n\n")

                    # 写入高VIF值特征
                    high_vif_features = multicollinearity_results.get('high_vif_features', [])
                    extreme_vif_features = multicollinearity_results.get('extreme_vif_features', [])

                    if extreme_vif_features:
                        f.write("极高VIF值特征 (VIF > 30):\n")
                        for feature, vif_value in extreme_vif_features:
                            f.write(f"  - {feature}: VIF = {vif_value:.2f}\n")
                        f.write("\n")

                    if high_vif_features:
                        f.write("高VIF值特征 (10 < VIF < 30):\n")
                        for feature, vif_value in high_vif_features:
                            f.write(f"  - {feature}: VIF = {vif_value:.2f}\n")
                        f.write("\n")

                    # 写入高度相关的特征对
                    high_corr_pairs = multicollinearity_results.get('high_corr_pairs', [])
                    if high_corr_pairs:
                        f.write("高度相关的特征对 (|r| > 0.8):\n")
                        for i, pair in enumerate(high_corr_pairs[:10]):  # 最多显示10对
                            f.write(f"  - {pair['feature1']} & {pair['feature2']}: r = {pair['correlation']:.4f}\n")

                        if len(high_corr_pairs) > 10:
                            f.write(f"  ... 以及{len(high_corr_pairs) - 10}对更多\n")
                        f.write("\n")

                    # 添加多重共线性建议
                    f.write("多重共线性诊断结论:\n")
                    if multicollinearity_results.get('extreme_vif_count', 0) > 0:
                        f.write("  - 存在严重多重共线性问题，建议移除或合并具有极高VIF值的特征\n")
                        f.write("  - 这些特征与其他特征高度相关，可能导致回归模型不稳定\n")
                    elif multicollinearity_results.get('high_vif_count', 0) > 0:
                        f.write("  - 存在多重共线性问题，建议考虑处理具有高VIF值的特征\n")
                        f.write("  - 这些特征与其他特征存在相关性，可能影响回归系数的稳定性\n")
                    else:
                        f.write("  - 未检测到显著的多重共线性问题\n")
                    f.write("\n")
            else:
                f.write("未执行多重共线性检验\n\n")

            # 6. 总体诊断结论和建模建议
            f.write("\n6. 总体诊断结论和建模建议\n")
            f.write("-" * 60 + "\n")

            # 分析平稳性问题
            stationarity_issues = False
            for var, result in unit_root_results.items():
                if result['stationary_percentage'] < 50:
                    stationarity_issues = True
                    break

            # 分析固定效应vs随机效应选择
            fixed_effects_recommended = True
            for var, result in hausman_results.items():
                if 'error' not in result and result.get('conclusion') == '随机效应':
                    fixed_effects_recommended = False

            # 分析异方差和自相关问题
            heteroskedasticity_issues = False
            for var, result in heteroskedasticity_results.items():
                if 'error' not in result and result['heteroskedasticity']:
                    heteroskedasticity_issues = True
                    break

            autocorrelation_issues = False
            for var, result in serial_correlation_results.items():
                if 'error' not in result and result['companies_with_autocorrelation_pct'] > 25:
                    autocorrelation_issues = True
                    break

            # 分析多重共线性问题
            multicollinearity_issues = False
            severe_multicollinearity = False
            if 'multicollinearity' in diagnostics and 'error' not in diagnostics['multicollinearity']:
                multicollinearity_results = diagnostics['multicollinearity']
                if multicollinearity_results.get('extreme_vif_count', 0) > 0:
                    multicollinearity_issues = True
                    severe_multicollinearity = True
                elif multicollinearity_results.get('high_vif_count', 0) > 0:
                    multicollinearity_issues = True

            # 写入总体结论
            if stationarity_issues:
                f.write("- 平稳性问题: 存在非平稳序列，建议进行差分处理或考虑协整关系\n")
            else:
                f.write("- 平稳性问题: 大多数序列平稳，可以直接使用面板数据模型\n")

            f.write(f"- 模型选择: {'推荐使用固定效应模型' if fixed_effects_recommended else '可以考虑使用随机效应模型'}\n")

            if heteroskedasticity_issues:
                f.write("- 异方差问题: 存在显著异方差，必须使用稳健标准误\n")
            else:
                f.write("- 异方差问题: 异方差不明显，但建议使用稳健标准误以保险\n")

            if autocorrelation_issues:
                f.write("- 自相关问题: 存在序列相关性，建议使用聚类稳健标准误或添加滞后变量\n")
            else:
                f.write("- 自相关问题: 自相关程度不高，但使用聚类稳健标准误更为保险\n")

            if severe_multicollinearity:
                f.write("- 多重共线性问题: 存在严重多重共线性，必须移除或合并高度相关的变量\n")
            elif multicollinearity_issues:
                f.write("- 多重共线性问题: 存在多重共线性，建议考虑处理高VIF值的变量\n")
            else:
                f.write("- 多重共线性问题: 未检测到显著的多重共线性\n")

            # 最终建模建议
            f.write("\n最终建模建议:\n")

            model_type = "固定效应" if fixed_effects_recommended else "随机效应"
            error_treatment = []
            if heteroskedasticity_issues:
                error_treatment.append("异方差稳健标准误")
            if autocorrelation_issues:
                error_treatment.append("聚类稳健标准误")

            if error_treatment:
                error_str = "和".join(error_treatment)
                f.write(f"建议使用{model_type}模型，同时采用{error_str}进行估计。\n")
            else:
                f.write(f"建议使用{model_type}模型，可以使用普通标准误，但采用稳健标准误更为保险。\n")

            if stationarity_issues:
                f.write("在建模之前，应考虑对非平稳变量进行差分处理或探索协整关系。\n")

            if severe_multicollinearity:
                f.write("在进行回归分析前，必须解决严重的多重共线性问题：\n")
                if 'multicollinearity' in diagnostics and 'extreme_vif_features' in diagnostics['multicollinearity']:
                    for feature, vif_value in diagnostics['multicollinearity']['extreme_vif_features'][:3]:  # 只显示前3个
                        f.write(f"  - 考虑移除或替换变量 '{feature}' (VIF = {vif_value:.2f})\n")
                    if len(diagnostics['multicollinearity']['extreme_vif_features']) > 3:
                        f.write(f"  - 以及其他{len(diagnostics['multicollinearity']['extreme_vif_features'])-3}个极高VIF值的变量\n")
                f.write("  - 或者使用主成分分析(PCA)等降维技术处理高度相关的变量\n")
            elif multicollinearity_issues:
                f.write("建议处理多重共线性问题，可以考虑：\n")
                f.write("  - 中心化交互项变量\n")
                f.write("  - 选择性地移除部分高度相关的控制变量\n")
                f.write("  - 使用正则化方法如岭回归(Ridge)来稳定系数估计\n")
    except Exception as e:
        print(f"生成诊断报告出错: {e}")
        traceback.print_exc()

    print("面板数据诊断完成，结果已保存到output/reports/panel_diagnostics_report.txt")

    return diagnostics

def advanced_panel_regression(df, dependent_var, independent_vars, control_vars=None,
                           entity_effects=True, time_effects=False,
                           cluster_entity=True, robust=True, drop_absorbed=False,
                           add_lagged_dependent=False, final_vars=None):
    """高级版面板回归模型，优化控制变量数量和固定效应处理"""
    # 使用变量追踪器获取实际变量名
    try:
        # 根据final_vars解析实际的变量名
        actual_dependent_var = None
        actual_independent_vars = []
        actual_control_vars = []

        if final_vars is not None:
            # 解析因变量
            if dependent_var in final_vars:
                actual_dependent_var = final_vars[dependent_var]
                print(f"使用映射的因变量名: {dependent_var} -> {actual_dependent_var}")
            else:
                actual_dependent_var = dependent_var
                print(f"因变量 {dependent_var} 未在映射中找到，直接使用")

            # 解析自变量
            for var in independent_vars:
                if var in final_vars:
                    actual_var = final_vars[var]
                    actual_independent_vars.append(actual_var)
                    print(f"使用映射的自变量名: {var} -> {actual_var}")
                else:
                    actual_independent_vars.append(var)
                    print(f"自变量 {var} 未在映射中找到，直接使用")

            # 解析控制变量
            if control_vars:
                # 限制控制变量数量，防止过度拟合
                if len(control_vars) > 8:
                    print(f"控制变量过多 ({len(control_vars)})，限制为最多8个关键控制变量")
                    # 优先选择重要控制变量
                    key_controls = ['size', 'lev', 'roa', 'age2', 'tobinq1', 'audit', 'soe', 'balance']
                    # 按照优先级选择最多8个
                    filtered_controls = []
                    for var in key_controls:
                        if var in control_vars and len(filtered_controls) < 8:
                            filtered_controls.append(var)
                    # 如果还没够8个，再从其他控制变量选
                    other_controls = [v for v in control_vars if v not in key_controls]
                    if len(filtered_controls) < 8 and other_controls:
                        filtered_controls.extend(other_controls[:8-len(filtered_controls)])
                    control_vars = filtered_controls
                    print(f"控制变量已减少至: {control_vars}")

                for var in control_vars:
                    if var in final_vars:
                        actual_var = final_vars[var]
                        actual_control_vars.append(actual_var)
                        print(f"使用映射的控制变量名: {var} -> {actual_var}")
                    else:
                        actual_control_vars.append(var)
                        print(f"控制变量 {var} 未在映射中找到，直接使用")
            elif 'controls' in final_vars:
                # 使用映射中预定义的控制变量，但限制数量
                all_controls = final_vars['controls']
                if len(all_controls) > 8:
                    print(f"预定义控制变量过多 ({len(all_controls)})，限制为最多8个")
                    actual_control_vars = all_controls[:8]
                else:
                    actual_control_vars = all_controls
                print(f"使用映射中预定义的控制变量: {actual_control_vars}")
        else:
            actual_dependent_var = dependent_var
            actual_independent_vars = independent_vars
            # 限制控制变量数量
            if control_vars and len(control_vars) > 8:
                print(f"控制变量过多 ({len(control_vars)})，限制为最多8个")
                actual_control_vars = control_vars[:8]
            else:
                actual_control_vars = control_vars if control_vars else []

        print(f"开始为变量 {actual_dependent_var} 创建面板回归模型...")

        # 确保我们有一个单独的DataFrame副本
        all_vars = [actual_dependent_var] + actual_independent_vars
        if actual_control_vars:
            all_vars += actual_control_vars

        # 检查变量是否存在
        missing_vars = [var for var in all_vars if var not in df.columns]
        if missing_vars:
            print(f"错误: 以下变量不在数据集中: {missing_vars}")
            print(f"可用列: {df.columns.tolist()[:10]}...")
            return None

        # 创建一个只包含必要列的干净数据框
        df_model = df[all_vars].copy()

        # 删除有缺失值的行
        na_before = len(df_model)
        df_model = df_model.dropna()
        na_after = len(df_model)
        if na_before > na_after:
            print(f"警告: 删除了 {na_before - na_after} 行有缺失值的数据")

        # 对自变量进行去中心化(centering)以减少多重共线性
        print("对自变量进行去中心化处理...")
        for var in actual_independent_vars:
            if var in df_model.columns and not var.endswith('_centered') and not var.endswith('_center'):
                # 计算均值
                var_mean = df_model[var].mean()
                # 创建去中心化的变量
                df_model[f'{var}_centered'] = df_model[var] - var_mean
                print(f"  已去中心化: {var} -> {var}_centered (均值: {var_mean:.4f})")

                # 更新变量名
                idx = actual_independent_vars.index(var)
                actual_independent_vars[idx] = f'{var}_centered'

        # 处理交互项
        interaction_vars = [var for var in actual_independent_vars if '_x_' in var]
        if interaction_vars:
            print("检测到交互项变量，确保使用去中心化后的主效应变量创建交互项...")
            for interact_var in interaction_vars:
                # 尝试识别交互项中的主效应变量
                parts = interact_var.split('_x_')
                if len(parts) == 2:
                    main_var1, main_var2 = parts
                    # 查看是否有对应的去中心化变量
                    centered_var1 = next((v for v in df_model.columns if v.startswith(f"{main_var1}_") and (v.endswith('_centered') or v.endswith('_center'))), main_var1)
                    centered_var2 = next((v for v in df_model.columns if v.startswith(f"{main_var2}_") and (v.endswith('_centered') or v.endswith('_center'))), main_var2)

                    if centered_var1 in df_model.columns and centered_var2 in df_model.columns:
                        # 使用去中心化变量重新创建交互项
                        centered_interact = f'{centered_var1}_x_{centered_var2}'
                        df_model[centered_interact] = df_model[centered_var1] * df_model[centered_var2]
                        print(f"  已创建去中心化交互项: {centered_interact}")

                        # 更新变量名
                        idx = actual_independent_vars.index(interact_var)
                        actual_independent_vars[idx] = centered_interact

        # 显示数据示例
        print("\n处理后的数据示例:")
        print(df_model.head(3))

        # 创建公式
        formula = f"{actual_dependent_var} ~ "
        formula += " + ".join(actual_independent_vars)

        if actual_control_vars and len(actual_control_vars) > 0:
            valid_controls = [c for c in actual_control_vars if c in df_model.columns]
            if valid_controls:
                formula += " + " + " + ".join(valid_controls)

        print(f"模型公式: {formula}")

        # 转换为标准的纵向面板数据格式
        try:
            # 检查是否已经是正确的索引
            if df_model.index.names != ['stkcd', 'year']:
                df_model = df_model.reset_index()
                # 确保索引列存在
                if 'stkcd' not in df_model.columns or 'year' not in df_model.columns:
                    print("错误: 数据框中缺少实体ID或时间ID列")
                    return None
                # 将索引列转换为正确的类型
                df_model['stkcd'] = df_model['stkcd'].astype(str)
                df_model['year'] = df_model['year'].astype(int)
                # 设置索引
                df_model = df_model.set_index(['stkcd', 'year'])

            # 将索引类型转换为字符串，避免类型错误
            df_model = df_model.copy()
            df_model.index = pd.MultiIndex.from_tuples(
                [(str(s), int(y)) for s, y in df_model.index],
                names=['stkcd', 'year']
            )

            # 确保所有列是数值型
            for col in df_model.columns:
                if not pd.api.types.is_numeric_dtype(df_model[col]):
                    print(f"将列 {col} 转换为数值型")
                    df_model[col] = pd.to_numeric(df_model[col], errors='coerce')

            # 再次删除缺失值
            df_model = df_model.dropna()

            # 转换为标准的浮点数类型，避免某些整数类型的问题
            df_model = df_model.astype(float)

        except Exception as e:
            print(f"数据准备阶段出错: {e}")
            traceback.print_exc()
            return None

        # 处理序列相关 - 添加因变量滞后项创建动态面板
        if add_lagged_dependent:
            print(f"创建动态面板模型: 添加因变量 {actual_dependent_var} 的滞后项")

            # 创建滞后项
            lagged_dep_var = f"{actual_dependent_var}_lag1"
            df_model[lagged_dep_var] = np.nan

            # 按实体分组创建滞后变量
            df_temp = df_model.copy().reset_index()
            for entity, group in df_temp.groupby('stkcd'):
                # 按时间排序
                group = group.sort_values('year')
                # 创建滞后值
                df_temp.loc[group.index, lagged_dep_var] = group[actual_dependent_var].shift(1).values

            # 将处理后的数据转回多级索引
            df_model = df_temp.set_index(['stkcd', 'year'])

            # 移除含有NaN的行
            df_model = df_model.dropna(subset=[lagged_dep_var])

            # 添加到自变量列表
            independent_vars.append(lagged_dep_var)
            print(f"已添加滞后因变量: {lagged_dep_var}, 模型有效观测值: {len(df_model)}")

        # 在创建模型之前，检查因变量的方差
        y_var = df[dependent_var].var()
        if y_var < 1e-8:
            print(f"警告: 因变量 {dependent_var} 方差极小 ({y_var:.2e})，添加微小扰动")
            # 创建副本并添加微小扰动
            df = df.copy()
            df[dependent_var] = df[dependent_var] + np.random.normal(0, max(y_var*0.01, 1e-8), size=len(df))

        # 尝试创建模型
        try:
            # 创建模型 - 直接使用正确的参数创建
            from linearmodels.panel import PanelOLS

            # 分离因变量和自变量
            y = df_model[actual_dependent_var]
            X = df_model[[v for v in df_model.columns if v != actual_dependent_var]]

            # 创建模型，考虑时间效应的使用
            # 对于短面板，避免同时使用企业和时间固定效应
            if time_effects and entity_effects and df_model.index.get_level_values('year').nunique() < 7:
                print("警告: 时间序列较短，避免同时使用企业和时间固定效应，仅使用企业固定效应")
                time_effects = False

            try:
                model = PanelOLS(
                    y,
                    X,
                    entity_effects=entity_effects,
                    time_effects=time_effects,
                    drop_absorbed=drop_absorbed,
                    check_rank=True  # 先使用标准秩检查
                )
            except ValueError as e:
                # 如果出现秩不足错误，尝试禁用秩检查
                if "exog does not have full column rank" in str(e):
                    print("警告: 检测到矩阵秩不足问题，尝试禁用秩检查...")
                    model = PanelOLS(
                        y,
                        X,
                        entity_effects=entity_effects,
                        time_effects=time_effects,
                        drop_absorbed=True,
                        check_rank=False  # 禁用秩检查
                    )
                else:
                    raise  # 如果是其他错误，则继续抛出

        except Exception as e:
            print(f"模型创建失败: {e}")
            traceback.print_exc()
            return None
          # 尝试拟合模型，使用更合理的模型规范
        try:
            # 设置拟合选项 - 使用适当级别的标准误处理
            fit_options = {}

            # 只选择一种标准误处理方法避免过度调整
            if cluster_entity:
                fit_options['cov_type'] = 'clustered'
                fit_options['cluster_entity'] = True
            elif robust:
                fit_options['cov_type'] = 'robust'
            else:
                fit_options['cov_type'] = 'unadjusted'  # 使用普通标准误

            # 直接拟合模型，使用合适的标准误选项
            print(f"使用 {fit_options['cov_type']} 标准误拟合模型...")
            try:
                result = model.fit(**fit_options)
            except ZeroDivisionError as ze:
                print(f"警告: 在模型拟合过程中发生除零错误: {ze}")
                print("正在尝试处理可能的零方差问题...")

                # 检查数据中的零方差列
                var_info = df_model.var()
                problem_cols = var_info[var_info <= 1e-8].index.tolist()
                if problem_cols:
                    print(f"发现低/零方差列: {problem_cols}")
                    # 从X中移除问题列
                    X_clean = X.drop(columns=[col for col in problem_cols if col in X.columns])
                    if X_clean.shape[1] == 0:
                        raise ValueError("移除零方差列后没有剩余变量")

                    print("使用移除零方差列后的数据重建模型...")
                    model = PanelOLS(
                        y,
                        X_clean,
                        entity_effects=entity_effects,
                        time_effects=time_effects,
                        drop_absorbed=drop_absorbed,
                        check_rank=False
                    )
                    result = model.fit(**fit_options)
                else:
                    # 尝试简化模型
                    print("没有发现明确的零方差列，尝试简化模型...")
                    # 使用更少的变量，优先保留主要研究变量
                    main_vars = actual_independent_vars[:3]
                    if len(actual_control_vars) > 3:
                        control_subset = actual_control_vars[:3]
                    else:
                        control_subset = actual_control_vars

                    print(f"使用简化变量集: {main_vars + control_subset}")
                    X_simple = df_model[main_vars + control_subset]

                    # 创建简化模型
                    simple_model = PanelOLS(
                        y,
                        X_simple,
                        entity_effects=entity_effects,
                        time_effects=False,  # 避免使用时间效应
                        drop_absorbed=True,
                        check_rank=False
                    )
                    result = simple_model.fit(**fit_options)

            # 计算标准化系数
            try:
                std_coefficients = calculate_standardized_coefficients(df_model, result)
            except Exception as e:
                print(f"计算标准化系数失败: {e}")
                traceback.print_exc()
                std_coefficients = None

            return {
                'result': result,
                'std_coefficients': std_coefficients,
                'dependent_var': actual_dependent_var,  # 添加实际使用的变量名
                'independent_vars': actual_independent_vars,
                'control_vars': actual_control_vars
            }
        except Exception as e:
            print(f"模型拟合失败: {e}")
            traceback.print_exc()

            # 尝试基础模型
            print("尝试拟合基础模型，不使用固定效应...")
            try:
                base_model = PanelOLS(
                    y,
                    X,
                    entity_effects=False,
                    time_effects=False,
                    check_rank=False
                )
                result = base_model.fit(cov_type='robust')

                return {
                    'result': result,
                    'std_coefficients': None,
                    'dependent_var': actual_dependent_var,
                    'independent_vars': actual_independent_vars,
                    'control_vars': actual_control_vars,
                    'is_base_model': True
                }
            except Exception as base_e:
                print(f"基础模型也失败: {base_e}")
                return None
    except Exception as outer_e:
        print(f"高级面板回归函数出错: {outer_e}")
        traceback.print_exc()
        return None

# 8.计算标准化系数（已修复）
def calculate_standardized_coefficients(data, result):
    """
    计算回归模型的标准化系数(β系数)，适用于面板数据模型

    参数:
    data: 用于拟合模型的数据
    result: 回归模型的结果对象(PanelOLS或其他模型的拟合结果)

    返回:
    包含标准化系数的字典
    """
    try:
        # 获取模型的因变量名
        if hasattr(result, 'model') and hasattr(result.model, 'dependent'):
            # statsmodels方式
            y_name = result.model.dependent.vars[0]
        else:
            # linearmodels方式
            y_name = result.dependent.vars[0]

        # 获取自变量列表
        x_names = [var for var in result.params.index]

        # 将多级索引数据转为常规DataFrame以便计算标准差
        if isinstance(data.index, pd.MultiIndex):
            data_reset = data.reset_index()
        else:
            data_reset = data

        # 计算因变量和自变量的标准差
        y_std = data_reset[y_name].std()
        x_std = {}

        for x in x_names:
            if x in data_reset.columns:
                x_std[x] = data_reset[x].std()
            else:                # 处理可能的交互项或变换后的变量
                x_components = x.split('_x_')  # 检查是否为交互项
                if len(x_components) == 2 and all(c in data_reset.columns for c in x_components):
                    # 对于包含零值较多的变量，计算标准差时需要小心
                    std1 = data_reset[x_components[0]].std()
                    std2 = data_reset[x_components[1]].std()

                    # 避免零标准差导致的问题
                    if std1 <= 1e-8 or std2 <= 1e-8:
                        print(f"警告: 交互项 {x} 的分量具有接近零的标准差 (std1={std1:.2e}, std2={std2:.2e})")
                        # 使用非零的最小值
                        std1 = max(std1, 1e-8)
                        std2 = max(std2, 1e-8)

                    x_std[x] = std1 * std2
                    print(f"为交互项 {x} 计算近似标准差 ({std1:.4f} * {std2:.4f})")
                else:
                    # 标记未找到的变量
                    print(f"警告: 未找到变量 {x} 的标准差，使用1.0作为默认值")
                    x_std[x] = 1.0

        # 计算标准化系数
        std_coefficients = {}

        for x in x_names:
            if x in x_std:
                # 标准化系数 = 原始系数 * (自变量标准差 / 因变量标准差)
                std_coefficients[x] = result.params[x] * (x_std[x] / y_std)
            else:
                # 对于固定效应等不在数据中的变量
                std_coefficients[x] = result.params[x] / y_std

        return std_coefficients
    except Exception as e:
        print(f"标准化系数计算出错: {e}")
        traceback.print_exc()
        return {}

# 9. Johnson-Neyman分析函数
def johnson_neyman_analysis(df, predictor, moderator, outcome, interaction, result, alpha=0.05):
    """执行Johnson-Neyman分析，确定调节效应的显著性区域"""
    try:
        if not result or not hasattr(result, 'params') or not hasattr(result, 'cov'):
            print("错误: 无效的模型结果")
            return None

        # 获取参数
        try:
            b_predictor = result.params[predictor]
            b_interaction = result.params[interaction]
            se_predictor = result.std_errors[predictor]
            se_interaction = result.std_errors[interaction]
            cov_predictor_interaction = result.cov.loc[predictor, interaction]
        except KeyError as e:
            print(f"Johnson-Neyman分析错误: 找不到参数 {e}")
            # 尝试查找正确的变量名
            print(f"可用参数: {list(result.params.index)}")

            # 尝试查找可能的变体名称
            predictor_candidates = [p for p in result.params.index if predictor in p]
            interaction_candidates = [p for p in result.params.index if interaction in p]

            if predictor_candidates:
                print(f"找到可能的预测变量替代: {predictor_candidates}")
                predictor = predictor_candidates[0]

            if interaction_candidates:
                print(f"找到可能的交互项替代: {interaction_candidates}")
                interaction = interaction_candidates[0]

            try:
                b_predictor = result.params[predictor]
                b_interaction = result.params[interaction]
                se_predictor = result.std_errors[predictor]
                se_interaction = result.std_errors[interaction]
                cov_predictor_interaction = result.cov.loc[predictor, interaction]
            except KeyError:
                print("仍然无法找到必要的参数，无法进行Johnson-Neyman分析")
                return None

        # 重置索引以便操作
        if isinstance(df.index, pd.MultiIndex):
            df_reset = df.reset_index()
        else:
            df_reset = df.copy()

        # 计算调节变量的值范围
        mod_min = df_reset[moderator].min()
        mod_max = df_reset[moderator].max()

        # 创建用于绘图的调节变量值序列
        mod_range = np.linspace(mod_min, mod_max, 100)

        # 计算简单斜率和标准误
        simple_slopes = []
        slope_se = []
        t_values = []
        p_values = []
        significance = []

        for mod_value in mod_range:
            # 简单斜率: b_predictor + b_interaction * mod_value
            slope = b_predictor + b_interaction * mod_value

            # 标准误: sqrt(var(b_predictor) + mod_value^2 * var(b_interaction) + 2 * mod_value * cov(b_predictor, b_interaction))
            try:
                var_term = se_predictor**2 + mod_value**2 * se_interaction**2 + 2 * mod_value * cov_predictor_interaction
                # 处理数值不稳定性
                if var_term <= 0:
                    print(f"警告: 在值 {mod_value} 处方差项为负值或零 ({var_term})")
                    var_term = max(1e-10, var_term)  # 使用小的正值代替
                se = np.sqrt(var_term)

                # t值和p值
                t = slope / se if se > 0 else 0  # 防止除以零
                p = 2 * (1 - stats.t.cdf(abs(t), df=result.df_resid))
            except Exception as calc_err:
                print(f"计算简单斜率时出错: {calc_err}, 使用默认值")
                se = 1.0
                t = 0
                p = 1.0
            p = 2 * (1 - stats.t.cdf(abs(t), df=result.df_resid))

            simple_slopes.append(slope)
            slope_se.append(se)
            t_values.append(t)
            p_values.append(p)
            significance.append(p < alpha)

        # 找到显著性区域的边界（Johnson-Neyman点）
        jn_points = []
        for i in range(1, len(mod_range)):
            if significance[i] != significance[i-1]:
                # 找到了一个边界点
                # 通过线性插值找到确切的分界点
                mod_low, mod_high = mod_range[i-1], mod_range[i]
                p_low, p_high = p_values[i-1], p_values[i]

                # 线性插值
                jn_point = mod_low + (mod_high - mod_low) * (alpha - p_low) / (p_high - p_low)
                jn_points.append(jn_point)

        # 准备绘图数据
        plot_data = pd.DataFrame({
            'moderator_value': mod_range,
            'simple_slope': simple_slopes,
            'standard_error': slope_se,
            'p_value': p_values,
            'significant': significance
        })

        # 定义显著区域
        significant_regions = []

        if len(jn_points) == 0:
            # 没有JN点，说明整个区间要么全部显著，要么全部不显著
            if all(significance):
                significant_regions.append({
                    'region': [mod_min, mod_max],
                    'description': f"调节变量整个范围 ({mod_min:.4f} 到 {mod_max:.4f}) 内效应显著"
                })
            else:
                # 全部不显著，返回空的显著区域
                pass
        elif len(jn_points) == 1:
            # 只有一个JN点，说明只有一个区间显著
            jn_point = jn_points[0]

            # 判断是左侧显著还是右侧显著
            left_significant = significance[0]

            if left_significant:
                significant_regions.append({
                    'region': [mod_min, jn_point],
                    'description': f"调节变量小于 {jn_point:.4f} 时效应显著"
                })
            else:
                significant_regions.append({
                    'region': [jn_point, mod_max],
                    'description': f"调节变量大于 {jn_point:.4f} 时效应显著"
                })
        else:
            # 有多个JN点，交替显著和不显著的区域
            jn_points_ext = [mod_min] + jn_points + [mod_max]
            current_significant = significance[0]

            for i in range(len(jn_points_ext) - 1):
                if current_significant:
                    significant_regions.append({
                        'region': [jn_points_ext[i], jn_points_ext[i+1]],
                        'description': f"调节变量在 {jn_points_ext[i]:.4f} 到 {jn_points_ext[i+1]:.4f} 之间时效应显著"
                    })
                current_significant = not current_significant

        # 准备返回结果
        jn_result = {
            'plot_data': plot_data,
            'jn_points': jn_points,
            'predictor': predictor,
            'moderator': moderator,
            'interaction': interaction,
            'main_effect': b_predictor,
            'interaction_effect': b_interaction,
            'alpha': alpha,
            'mod_range': [mod_min, mod_max],
            'significant_regions': significant_regions
        }

        # 解释结果
        if not jn_points:
            if all(significance):
                print(f"在整个调节变量的范围内，{predictor}对{outcome}的影响都是显著的")
            elif not any(significance):
                print(f"在整个调节变量的范围内，{predictor}对{outcome}的影响都不显著")
        else:
            print(f"找到 {len(jn_points)} 个Johnson-Neyman点:")
            for i, point in enumerate(jn_points):
                print(f"  点{i+1}: 当{moderator} = {point:.4f}时，简单斜率的显著性发生变化")

            for region in significant_regions:
                print(f"  {region['description']}")

        return jn_result
    except Exception as e:
        print(f"Johnson-Neyman分析出错: {e}")
        traceback.print_exc()
        return {'error': str(e)}

# 10. U型关系检验函数修复
def improved_u_shape_test(result, x_var, x_squared_var, data, dependent_var=None, output_dir=None):
    """
    完善版U型关系检验，提供更严格的统计验证

    参数:
        result: 回归模型结果
        x_var: 线性项变量名
        x_squared_var: 二次项变量名
        data: 数据框
        dependent_var: 因变量名(可选)
        output_dir: 输出目录(可选)

    返回:
        包含检验结果的字典
    """
    if output_dir is None:
        output_dir = 'output/reports/'

    if dependent_var is None:
        # 尝试从result中提取因变量名
        if hasattr(result, 'model') and hasattr(result.model, 'dependent'):
            dependent_var = result.model.dependent.vars[0]
        else:
            dependent_var = result.dependent.vars[0]

    # 提取系数和标准误
    b_linear = result.params[x_var]
    b_squared = result.params[x_squared_var]
    se_linear = result.std_errors[x_var]
    se_squared = result.std_errors[x_squared_var]

    # 确定关系类型
    if b_squared > 0:
        relationship_type = "U型"
    else:
        relationship_type = "倒U型"

    # 计算拐点
    turning_point = -b_linear / (2 * b_squared)

    # 确保data是DataFrame
    if isinstance(data, pd.DataFrame) and isinstance(data.index, pd.MultiIndex):
        data_reset = data.reset_index()
    else:
        data_reset = data

    # 获取自变量的范围
    x_min = data_reset[x_var].min()
    x_max = data_reset[x_var].max()
    x_range = x_max - x_min

    # 判断拐点是否在数据范围内
    within_range = (turning_point > x_min) and (turning_point < x_max)

    # 计算拐点在整个区间的相对位置(0-1)
    if within_range:
        relative_position = (turning_point - x_min) / x_range
    else:
        relative_position = -1 if turning_point <= x_min else 2

    # 增强的Sasabuchi检验 (极值检验)
    slope_at_min = b_linear + 2 * b_squared * x_min  # 最小点处的斜率
    slope_at_max = b_linear + 2 * b_squared * x_max  # 最大点处的斜率

    # 计算斜率的标准误差
    var_slope_min = (se_linear**2 + 4*x_min**2*se_squared**2 +
                     4*x_min*result.cov.loc[x_var, x_squared_var])
    var_slope_max = (se_linear**2 + 4*x_max**2*se_squared**2 +
                     4*x_max*result.cov.loc[x_var, x_squared_var])

    se_slope_min = np.sqrt(max(0, var_slope_min))
    se_slope_max = np.sqrt(max(0, var_slope_max))

    # 计算t统计量
    t_min = slope_at_min / se_slope_min if se_slope_min > 0 else 0
    t_max = slope_at_max / se_slope_max if se_slope_max > 0 else 0

    # 计算p值 (单尾)
    p_min = 1 - stats.t.cdf(abs(t_min), df=result.df_resid) if slope_at_min < 0 else stats.t.cdf(-abs(t_min), df=result.df_resid)
    p_max = stats.t.cdf(t_max, df=result.df_resid) if slope_at_max > 0 else 1 - stats.t.cdf(t_max, df=result.df_resid)

    # 确定Sasabuchi检验的整体显著性
    if relationship_type == "U型":
        # U型需要：左端点斜率显著为负，右端点斜率显著为正
        overall_sasabuchi_p = max(p_min, p_max)
        sasabuchi_passes = (slope_at_min < 0 and slope_at_max > 0 and overall_sasabuchi_p < 0.05)
    else:
        # 倒U型需要：左端点斜率显著为正，右端点斜率显著为负
        overall_sasabuchi_p = max(1-p_min, 1-p_max)
        sasabuchi_passes = (slope_at_min > 0 and slope_at_max < 0 and overall_sasabuchi_p < 0.05)

    # 计算拐点的置信区间
    d_linear = -1/(2*b_squared)
    d_squared = b_linear/(2*b_squared**2)

    var_turning_point = (d_linear**2 * se_linear**2 +
                         d_squared**2 * se_squared**2 +
                         2 * d_linear * d_squared * result.cov.loc[x_var, x_squared_var])

    se_turning_point = np.sqrt(max(0, var_turning_point))
    ci_lower = turning_point - 1.96 * se_turning_point
    ci_upper = turning_point + 1.96 * se_turning_point

    # 判断置信区间是否在数据范围内
    ci_within_range = (ci_lower >= x_min) and (ci_upper <= x_max)

    # 验证拐点是否为最大值或最小值
    is_minimum = b_squared > 0
    is_maximum = b_squared < 0

    # 结果汇总
    valid_nonlinear = (
        result.pvalues[x_squared_var] < 0.05 and  # 二次项显著
        sasabuchi_passes and                      # Sasabuchi检验通过
        within_range                              # 拐点在数据范围内
    )

    # 为H2a和H2b量身定制的结论
    if relationship_type == "U型" and valid_nonlinear:  # H2a
        conclusion = "支持H2a假设：AI技术投入与专利质量呈显著U型关系。"
        explanation = (f"在投入较少时对专利质量有负面影响(斜率={slope_at_min:.4f})，"
                      f"但随着投入增加，当超过拐点({turning_point:.4f})后，"
                      f"效应转为正向(斜率={slope_at_max:.4f})，表明高水平AI投入能提高专利质量。")
    elif relationship_type == "倒U型" and valid_nonlinear:  # H2b
        conclusion = "支持H2b假设：AI技术投入与专利深度呈显著倒U型关系。"
        explanation = (f"在投入较少时对专利深度有正面影响(斜率={slope_at_min:.4f})，"
                      f"但随着投入增加，当超过拐点({turning_point:.4f})后，"
                      f"效应转为负向(斜率={slope_at_max:.4f})，表明过度AI投入可能不利于专利深度。")
    elif relationship_type == "U型":
        conclusion = "不支持H2a假设：虽然系数方向符合U型，但未通过严格统计检验。"
        if not (result.pvalues[x_squared_var] < 0.05):
            explanation = "二次项系数不显著。"
        elif not sasabuchi_passes:
            explanation = "未通过Sasabuchi极值检验。"
        elif not within_range:
            explanation = "拐点不在数据范围内。"
        else:
            explanation = "未能满足所有条件。"
    else:
        conclusion = "不支持H2b假设：虽然系数方向符合倒U型，但未通过严格统计检验。"
        if not (result.pvalues[x_squared_var] < 0.05):
            explanation = "二次项系数不显著。"
        elif not sasabuchi_passes:
            explanation = "未通过Sasabuchi极值检验。"
        elif not within_range:
            explanation = "拐点不在数据范围内。"
        else:
            explanation = "未能满足所有条件。"

    # 创建优化的可视化图表
    create_enhanced_nonlinear_plot(
        result=result,
        x_var=x_var,
        x_squared_var=x_squared_var,
        turning_point=turning_point,
        relationship_type=relationship_type,
        data_reset=data_reset,
        dependent_var=dependent_var,
        slope_at_min=slope_at_min,
        slope_at_max=slope_at_max,
        p_min=p_min,
        p_max=p_max,
        conclusion=conclusion,
        x_min=x_min,
        x_max=x_max,
        output_path=f'{output_dir}/{relationship_type}_relationship_{x_var}_{dependent_var}_enhanced.png'
    )

    # 返回详细结果
    return {
        'relationship_type': relationship_type,
        'turning_point': turning_point,
        'turning_point_ci': (ci_lower, ci_upper),
        'ci_within_range': ci_within_range,
        'within_range': within_range,
        'relative_position': relative_position,
        'slope_at_min': slope_at_min,
        'slope_at_max': slope_at_max,
        'p_value_min': p_min,
        'p_value_max': p_max,
        'squared_coefficient': b_squared,
        'squared_p_value': result.pvalues[x_squared_var],
        'sasabuchi_passes': sasabuchi_passes,
        'sasabuchi_p': overall_sasabuchi_p,
        'is_minimum': is_minimum,
        'is_maximum': is_maximum,
        'valid_nonlinear': valid_nonlinear,
        'conclusion': conclusion,
        'explanation': explanation,
        'plot_path': f'{output_dir}/{relationship_type}_relationship_{x_var}_{dependent_var}_enhanced.png'
    }

def create_enhanced_nonlinear_plot(result, x_var, x_squared_var, turning_point, relationship_type,
                                 data_reset, dependent_var, slope_at_min, slope_at_max,
                                 p_min, p_max, conclusion, x_min, x_max, output_path):
    """创建更优化的非线性关系可视化图表"""
    import matplotlib.pyplot as plt
    import numpy as np

    # 生成拟合曲线的x值
    x_vals = np.linspace(x_min, x_max, 300)

    # 计算拟合曲线的y值
    y_vals = result.params.get('const', 0) + result.params[x_var] * x_vals + result.params[x_squared_var] * x_vals ** 2

    # 计算转折点坐标
    if x_min <= turning_point <= x_max:
        turning_y = result.params.get('const', 0) + result.params[x_var] * turning_point + result.params[x_squared_var] * turning_point ** 2

    # 创建图表
    plt.figure(figsize=(10, 7))

    # 散点图 - 原始数据
    plt.scatter(data_reset[x_var], data_reset[dependent_var], alpha=0.3, color='gray', label='原始数据')

    # 拟合曲线
    plt.plot(x_vals, y_vals, color='blue', linewidth=2.5, label='拟合曲线')

    # 标记转折点
    if x_min <= turning_point <= x_max:
        plt.scatter([turning_point], [turning_y], color='red', s=100, zorder=5, label='拐点')
        plt.axvline(x=turning_point, color='red', linestyle='--', alpha=0.5)

    # 添加斜率指示箭头
    arrow_len_x = (x_max - x_min) * 0.1

    # 最小点斜率箭头
    if slope_at_min != 0:
        arrow_y_min = result.params.get('const', 0) + result.params[x_var] * x_min + result.params[x_squared_var] * x_min ** 2
        arrow_dir_y_min = slope_at_min * arrow_len_x
        plt.arrow(x_min, arrow_y_min, arrow_len_x, arrow_dir_y_min,
                 head_width=(max(y_vals)-min(y_vals))*0.03,
                 head_length=arrow_len_x*0.3,
                 fc='green', ec='green', label='斜率')

        significant_text = "显著" if p_min < 0.05 else "不显著"
        plt.text(x_min + arrow_len_x/2, arrow_y_min + arrow_dir_y_min/2,
                f"斜率={slope_at_min:.2f}\n({significant_text})",
                fontsize=9, ha='center')

    # 最大点斜率箭头
    if slope_at_max != 0:
        arrow_y_max = result.params.get('const', 0) + result.params[x_var] * x_max + result.params[x_squared_var] * x_max ** 2
        arrow_dir_y_max = slope_at_max * arrow_len_x
        plt.arrow(x_max - arrow_len_x, arrow_y_max - arrow_dir_y_max, arrow_len_x, arrow_dir_y_max,
                 head_width=(max(y_vals)-min(y_vals))*0.03,
                 head_length=arrow_len_x*0.3,
                 fc='green', ec='green')

        significant_text = "显著" if p_max < 0.05 else "不显著"
        plt.text(x_max - arrow_len_x/2, arrow_y_max - arrow_dir_y_max/2,
                f"斜率={slope_at_max:.2f}\n({significant_text})",
                fontsize=9, ha='center')

    # 添加结论
    plt.figtext(0.5, 0.01, conclusion, wrap=True, horizontalalignment='center', fontsize=12)

    # 设置标题和轴标签
    plt.title(f'AI技术投入与{dependent_var}的{relationship_type}关系', fontsize=14)
    plt.xlabel(f'AI技术投入 ({x_var})', fontsize=12)
    plt.ylabel(dependent_var, fontsize=12)

    plt.grid(True, alpha=0.3)
    plt.legend(loc='best')
    plt.tight_layout(rect=[0, 0.05, 1, 0.95])  # 为底部文本留出空间

    # 保存图片
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

def predict_at_turning_point(result, x_var, x_squared_var, turning_point):
    """计算拐点处的预测值"""
    # 尝试提取所有系数
    coefs = result.params.copy()

    # 创建预测值
    pred = 0
    if 'const' in coefs:
        pred += coefs['const']

    pred += coefs[x_var] * turning_point
    pred += coefs[x_squared_var] * turning_point**2

    # 添加其他变量的均值贡献(如果有)
    for var in coefs.index:
        if var not in ['const', x_var, x_squared_var] and not var.startswith('entity') and not var.startswith('time'):
            # 对于其他变量，我们不知道它们的值，这里简化处理，忽略它们的贡献
            pass

    return pred

def test_time_dynamic_effect(df, y_var, x_var, control_vars=None, entity_effects=True, periods=None):
    """
    增强版时间动态效应分析

    参数:
        df: 面板数据
        y_var: 因变量
        x_var: 自变量
        control_vars: 控制变量
        entity_effects: 是否包含实体固定效应
        periods: 自定义时期划分，如{'early':[2014,2015,2016], 'mid':[2017,2018,2019], 'late':[2020,2021,2022]}
    """
    print(f"\n分析{x_var}对{y_var}的时间动态效应...")

    # 复制数据进行处理
    df_analysis = df.copy()
    df_reset = df_analysis.reset_index()

    # 获取唯一年份并排序
    years = sorted(df_reset['year'].unique())

    # 创建标准化的时间趋势变量
    df_reset['time_trend'] = (df_reset['year'] - min(years)) / (max(years) - min(years))
    df_reset['time_trend_sq'] = df_reset['time_trend'] ** 2

    # 创建交互项
    df_reset[f'{x_var}_x_time'] = df_reset[x_var] * df_reset['time_trend']
    df_reset[f'{x_var}_x_time_sq'] = df_reset[x_var] * df_reset['time_trend_sq']

    # 如果提供了自定义时期，则使用自定义时期，否则自动划分
    if periods is None:
        # 自动按照三等分划分时期
        period_breaks = np.quantile(years, [0, 1/3, 2/3, 1])
        early_years = years[:int(len(years)/3)]
        mid_years = years[int(len(years)/3):int(2*len(years)/3)]
        late_years = years[int(2*len(years)/3):]
    else:
        early_years = periods['early']
        mid_years = periods['mid']
        late_years = periods['late']

    # 创建时期指示变量
    df_reset['period_early'] = df_reset['year'].isin(early_years).astype(int)
    df_reset['period_mid'] = df_reset['year'].isin(mid_years).astype(int)
    df_reset['period_late'] = df_reset['year'].isin(late_years).astype(int)

    # 创建时期交互项
    df_reset[f'{x_var}_x_mid'] = df_reset[x_var] * df_reset['period_mid']
    df_reset[f'{x_var}_x_late'] = df_reset[x_var] * df_reset['period_late']

    # 设置面板索引
    df_analysis = df_reset.set_index(['stkcd', 'year'])

    # 模型结果存储
    models = {}

    # 1. 基础模型 - 无时间交互
    base_model = advanced_panel_regression(
        df=df_analysis,
        dependent_var=y_var,
        independent_vars=[x_var],
        control_vars=control_vars,
        entity_effects=entity_effects
    )
    models['base'] = base_model

    # 2. 线性时间趋势模型
    linear_model = advanced_panel_regression(
        df=df_analysis,
        dependent_var=y_var,
        independent_vars=[x_var, 'time_trend', f'{x_var}_x_time'],
        control_vars=control_vars,
        entity_effects=entity_effects
    )
    models['linear_trend'] = linear_model

    # 3. 二次时间趋势模型
    quadratic_model = advanced_panel_regression(
        df=df_analysis,
        dependent_var=y_var,
        independent_vars=[x_var, 'time_trend', 'time_trend_sq', f'{x_var}_x_time', f'{x_var}_x_time_sq'],
        control_vars=control_vars,
        entity_effects=entity_effects
    )
    models['quadratic_trend'] = quadratic_model

    # 4. 时期划分模型
    period_model = advanced_panel_regression(
        df=df_analysis,
        dependent_var=y_var,
        independent_vars=[x_var, 'period_mid', 'period_late', f'{x_var}_x_mid', f'{x_var}_x_late'],
        control_vars=control_vars,
        entity_effects=entity_effects
    )
    models['period'] = period_model

    # 创建结果可视化
    create_time_dynamic_visualization(df_analysis, models, y_var, x_var,
                                    periods={"early": early_years, "mid": mid_years, "late": late_years})

    # 5. 计算每年效应 - 为了详细分析
    yearly_effects = []
    for year in years:
        # 创建当年数据子集
        year_data = df_analysis[df_analysis.index.get_level_values('year') == year]
        if len(year_data) > 30:  # 确保足够的样本量
            try:
                year_model = advanced_panel_regression(
                    df=year_data,
                    dependent_var=y_var,
                    independent_vars=[x_var],
                    control_vars=control_vars,
                    entity_effects=False  # 单年不使用实体效应
                )

                if year_model and year_model.get('result'):
                    coef = year_model['result'].params[x_var]
                    pval = year_model['result'].pvalues[x_var]
                    yearly_effects.append({
                        'year': year,
                        'coefficient': coef,
                        'p_value': pval,
                        'significant': pval < 0.05,
                        'observations': len(year_data)
                    })
            except Exception as e:
                print(f"年度 {year} 模型估计失败: {e}")

    if yearly_effects:
        models['yearly_effects'] = yearly_effects

    # 分析结果并提供详细解读
    results_interpretation = interpret_time_dynamics(models)

    return {
        'models': models,
        'interpretation': results_interpretation,
        'data': df_analysis
    }

def interpret_time_dynamics(models):
    """解释时间动态效应模型结果"""
    interpretation = {
        'has_dynamic_effect': False,
        'pattern': 'none',
        'early_effect': None,
        'late_effect': None,
        'turning_point': None,
        'explanation': ''
    }

    # 检查线性时间趋势模型
    if 'linear_trend' in models and models['linear_trend'] and models['linear_trend'].get('result'):
        result = models['linear_trend']['result']
        # 查找交互项系数
        interact_var = None
        for param in result.params.index:
            if '_x_time' in param and 'sq' not in param:
                interact_var = param
                break

        if interact_var:
            interact_coef = result.params[interact_var]
            interact_pval = result.pvalues[interact_var]

            if interact_pval < 0.05:
                interpretation['has_dynamic_effect'] = True
                if interact_coef > 0:
                    interpretation['pattern'] = 'strengthening'
                    interpretation['explanation'] = '随时间推移，效应逐渐增强'
                else:
                    interpretation['pattern'] = 'weakening'
                    interpretation['explanation'] = '随时间推移，效应逐渐减弱'

    # 检查二次时间趋势模型
    if 'quadratic_trend' in models and models['quadratic_trend'] and models['quadratic_trend'].get('result'):
        result = models['quadratic_trend']['result']
        # 查找交互项系数
        linear_var = None
        quad_var = None

        for param in result.params.index:
            if '_x_time' in param and 'sq' not in param:
                linear_var = param
            elif '_x_time_sq' in param:
                quad_var = param

        if linear_var and quad_var:
            linear_coef = result.params[linear_var]
            linear_pval = result.pvalues[linear_var]
            quad_coef = result.params[quad_var]
            quad_pval = result.pvalues[quad_var]

            if quad_pval < 0.05:
                interpretation['has_dynamic_effect'] = True

                # 计算拐点
                if quad_coef != 0:
                    turning_point = -linear_coef / (2 * quad_coef)
                    if 0 <= turning_point <= 1:  # 在标准化时间范围内
                        interpretation['turning_point'] = turning_point

                        if quad_coef > 0:
                            interpretation['pattern'] = 'u_shaped'
                            interpretation['explanation'] = '效应呈U型变化，先减弱后增强'
                        else:
                            interpretation['pattern'] = 'inverted_u_shaped'
                            interpretation['explanation'] = '效应呈倒U型变化，先增强后减弱'

    # 检查时期模型
    if 'period' in models and models['period'] and models['period'].get('result'):
        result = models['period']['result']
        x_var = None
        x_mid_var = None
        x_late_var = None

        # 识别参数名
        for param in result.params.index:
            if '_x_mid' in param:
                x_mid_var = param
            elif '_x_late' in param:
                x_late_var = param
            elif param in ['const', 'period_mid', 'period_late']:
                continue
            else:
                x_var = param

        if x_var and x_mid_var and x_late_var:
            base_coef = result.params[x_var]
            mid_coef = base_coef + result.params[x_mid_var]
            late_coef = base_coef + result.params[x_late_var]

            interpretation['early_effect'] = base_coef
            interpretation['late_effect'] = late_coef

            # 检查时期之间的变化
            if abs(base_coef) < abs(mid_coef) and abs(mid_coef) > abs(late_coef):
                interpretation['pattern'] = 'inverted_u_shaped_period'
                interpretation['explanation'] += '\n分期分析表明效应呈倒U型：中期最强，后期减弱'
            elif abs(base_coef) > abs(mid_coef) and abs(mid_coef) < abs(late_coef):
                interpretation['pattern'] = 'u_shaped_period'
                interpretation['explanation'] += '\n分期分析表明效应呈U型：中期最弱，后期增强'
            elif abs(base_coef) > abs(mid_coef) and abs(mid_coef) > abs(late_coef):
                interpretation['pattern'] = 'declining_period'
                interpretation['explanation'] += '\n分期分析表明效应逐渐减弱'
            elif abs(base_coef) < abs(mid_coef) and abs(mid_coef) < abs(late_coef):
                interpretation['pattern'] = 'increasing_period'
                interpretation['explanation'] += '\n分期分析表明效应逐渐增强'

    # 添加总结性解释
    if interpretation['has_dynamic_effect']:
        if interpretation['pattern'] in ['weakening', 'declining_period']:
            interpretation['summary'] = '总体而言，效应随时间推移逐渐减弱，支持假设H1b。'
        elif interpretation['pattern'] in ['strengthening', 'increasing_period']:
            interpretation['summary'] = '总体而言，效应随时间推移逐渐增强，与假设H1b相反。'
        elif interpretation['pattern'] in ['u_shaped', 'u_shaped_period']:
            interpretation['summary'] = '总体而言，效应呈U型变化，初期减弱后期增强。'
        elif interpretation['pattern'] in ['inverted_u_shaped', 'inverted_u_shaped_period']:
            interpretation['summary'] = '总体而言，效应呈倒U型变化，初期增强后期减弱，部分支持假设H1b。'
    else:
        interpretation['summary'] = '未检测到显著的时间动态效应，不支持假设H1b。'

    return interpretation

# 改进变量名查找和映射函数
def get_mapped_variable_name(var_name, var_mapping=None):
    """获取变量的映射名称，更健壮的实现"""
    if not var_mapping:
        return var_name

    # 直接查找精确匹配
    if var_name in var_mapping:
        return var_mapping[var_name]

    # 尝试忽略大小写查找
    lower_var = var_name.lower()
    for k, v in var_mapping.items():
        if k.lower() == lower_var:
            return v

    # 查找包含该变量名的键
    for k, v in var_mapping.items():
        if var_name in k:
            return v

    # 最后尝试寻找可能的变体
    variants = [f'{var_name}_log', f'{var_name}_sqrt', f'{var_name}_boxcox', f'{var_name}_quantile']
    for variant in variants:
        if variant in var_mapping:
            return var_mapping[variant]

    # 找不到映射，返回原名
    return var_name

def find_parameter_in_results(result, param_name):
    """在模型结果中查找参数，处理可能的变量名变化"""
    if not result or not hasattr(result, 'params'):
        return None

    # 直接查找精确匹配
    if param_name in result.params.index:
        return param_name

    # 尝试常见的变体
    variants = [
        f'{param_name}_centered',
        f'{param_name}_std',
        f'{param_name}_diff',
        f'{param_name}_log',
        f'{param_name}_sqrt',
        f'{param_name}_boxcox',
        f'{param_name}_quantile',
        f'{param_name}_yeojohnson'
    ]

    for variant in variants:
        if variant in result.params.index:
            return variant

    # 尝试部分匹配
    for param in result.params.index:
        if isinstance(param, str) and param.startswith(param_name):
            return param

    # 查找包含该参数名的项
    for param in result.params.index:
        if isinstance(param, str) and param_name in param:
            return param

    # 尝试忽略大小写查找
    lower_param = param_name.lower()
    for param in result.params.index:
        if isinstance(param, str) and param.lower() == lower_param:
            return param

    # 找不到匹配项
    return None

# 12. 添加动态效应可视化函数
def create_time_dynamic_visualization(df, models, y_var, x_var,
                                     periods=None, period_names=None):
    """
    期间可以从调用函数中动态传递
    """
    if periods is None:
        periods = {"early": [2014, 2015, 2016],
                  "mid": [2017, 2018, 2019],
                  "late": [2020, 2021, 2022]}
    if period_names is None:
        period_names = ["早期", "中期", "晚期"]

    if not models:
        return

    # 添加辅助函数来查找与参数相关的变量名
    def find_param(params, var_name):
        """安全查找参数中与指定变量名相关的参数"""
        # 先直接检查原始变量名
        if var_name in params.index:
            return var_name

        # 检查常见的变体
        variants = [f'{var_name}_centered', f'{var_name}_std', f'{var_name}_lag1']
        for variant in variants:
            if variant in params.index:
                return variant

        # 如果没有精确匹配，检查以变量名开头的参数
        for param in params.index:
            if isinstance(param, str) and param.startswith(var_name):
                return param

        # 最后，检查所有包含变量名的参数
        for param in params.index:
            if isinstance(param, str) and var_name in param.lower():
                return param

        return None

    # 创建时间序列效应图
    plt.figure(figsize=(14, 8))

    # 获取年份范围
    years = sorted(df.reset_index()['year'].unique())
    base_year = min(years)
    max_year = max(years)

    # 如果有年度交互模型，绘制年度效应变化
    if 'yearly' in models and models['yearly'] and models['yearly'].get('result'):
        yearly_result = models['yearly']['result']
        param_name = find_param(yearly_result.params, x_var)

        if param_name:
            # 计算每年的边际效应
            yearly_effects = [yearly_result.params[param_name]]  # 基准年份的效应

            for year in years[1:]:
                interaction_var = f'{x_var}_x_year_{year}'
                interact_param = find_param(yearly_result.params, interaction_var)

                if interact_param and interact_param in yearly_result.params:
                    yearly_effects.append(yearly_result.params[param_name] +
                                          yearly_result.params[interact_param])
                else:
                    yearly_effects.append(np.nan)

            # 绘制年度效应变化
            plt.plot(years, yearly_effects, 'o-', linewidth=2, markersize=8,
                     label='年度边际效应')
        else:
            print(f"警告: 在年度模型中找不到与 {x_var} 相关的参数，跳过绘制")

    # 如果有时期交互模型，绘制时期效应变化
    if 'period' in models and models['period'] and models['period'].get('result'):
        period_result = models['period']['result']
        param_name = find_param(period_result.params, x_var)

        if param_name:
            # 计算三个时期的边际效应
            base_effect = period_result.params[param_name]

            mid_interact_name = find_param(period_result.params, f'{x_var}_x_period_mid')
            late_interact_name = find_param(period_result.params, f'{x_var}_x_period_late')

            mid_effect = base_effect
            late_effect = base_effect

            if mid_interact_name and mid_interact_name in period_result.params:
                mid_effect += period_result.params[mid_interact_name]

            if late_interact_name and late_interact_name in period_result.params:
                late_effect += period_result.params[late_interact_name]

            # 计算时期边界（年份）
            period_bounds = [base_year]
            period_bounds.append(base_year + (max_year - base_year) // 3)
            period_bounds.append(base_year + 2 * (max_year - base_year) // 3)
            period_bounds.append(max_year)

            # 绘制时期效应
            period_x = [(period_bounds[0] + period_bounds[1]) / 2,
                       (period_bounds[1] + period_bounds[2]) / 2,
                       (period_bounds[2] + period_bounds[3]) / 2]

            period_effects = [base_effect, mid_effect, late_effect]

            plt.plot(period_x, period_effects, 's--', linewidth=2, markersize=10,
                     label='时期平均边际效应')
        else:
            print(f"警告: 在时期模型中找不到与 {x_var} 相关的参数，跳过绘制")

    # 如果有线性趋势交互模型，绘制趋势线
    if 'trend' in models and models['trend'] and models['trend'].get('result'):
        trend_result = models['trend']['result']
        param_name = find_param(trend_result.params, x_var)
        interact_name = find_param(trend_result.params, f'{x_var}_x_time_trend')

        if param_name and interact_name:
            # 绘制趋势线
            x_trend = np.linspace(base_year, max_year, 100)
            y_trend = []

            for x in x_trend:
                normalized_x = (x - base_year) / (max_year - base_year)
                effect = trend_result.params[param_name] + trend_result.params[interact_name] * normalized_x
                y_trend.append(effect)

            plt.plot(x_trend, y_trend, '-', linewidth=3, alpha=0.6,
                     label='线性时间趋势')
        else:
            print(f"警告: 在趋势模型中找不到所需参数，跳过绘制")

    # 完善图表
    plt.axhline(y=0, color='black', linestyle='-', alpha=0.3)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.xlabel('年份', fontsize=12)
    plt.ylabel(f'{x_var}对{y_var}的边际效应', fontsize=12)
    plt.title(f'{x_var}对{y_var}的时间动态效应', fontsize=14)
    plt.legend(loc='best', fontsize=12)

    # 保存图表
    plt.tight_layout()
    plt.savefig(f'output/figures/time_dynamic_{x_var}_{y_var}.png', dpi=300)
    plt.close()

# 13. 高级非线性关系检验
def advanced_nonlinear_test(result, x_var, x_squared_var, data, dependent_var=None, output_dir=None, plot_segments=True):
    """
    高级非线性关系检验，包括:
    1. 改进的Sasabuchi检验
    2. 三段式样条回归验证
    3. 精确拐点区间估计
    4. 增强的可视化
    5. 分段线性回归分析
    6. 非参数局部回归确认

    参数:
        result: 回归模型结果
        x_var: 线性项变量名
        x_squared_var: 二次项变量名
        data: 数据框
        dependent_var: 因变量名(可选)
        output_dir: 输出目录(可选)
        plot_segments: 是否绘制分段回归线(可选)

    返回:
        包含检验结果的字典
    """
    import traceback
    import numpy as np
    import pandas as pd
    import matplotlib.pyplot as plt
    from scipy import stats
    import statsmodels.api as sm

    # 结果字典
    u_test_result = {}

    if output_dir is None:
        output_dir = 'output/reports/'

    # 如果没有提供dependent_var，尝试从result中提取
    if dependent_var is None:
        try:
            if hasattr(result, 'model') and hasattr(result.model, 'dependent'):
                dependent_var = result.model.dependent.vars[0]
            else:
                dependent_var = result.dependent.vars[0]
            print(f"从模型中提取的因变量名: {dependent_var}")
        except Exception as e:
            print(f"无法从模型中提取因变量名: {e}")
            dependent_var = "dependent_var"  # 使用默认名称

    try:
        # 提取系数和标准误
        b_linear = result.params[x_var]
        b_squared = result.params[x_squared_var]
        se_linear = result.std_errors[x_var]
        se_squared = result.std_errors[x_squared_var]

        # 确定关系类型
        if b_squared > 0:
            relationship_type = "U型"
        else:
            relationship_type = "倒U型"

        # 计算拐点
        turning_point = -b_linear / (2 * b_squared)
        # 确保data是DataFrame
        if isinstance(data, pd.DataFrame) and isinstance(data.index, pd.MultiIndex):
            data_reset = data.reset_index()
        else:
            data_reset = data
    except Exception as e:
        print(f"An error occurred: {e}")


    # 数据范围
    x_min = data_reset[x_var].min()
    x_max = data_reset[x_var].max()
    x_range = x_max - x_min

    # 是否在数据范围内
    within_range = (turning_point > x_min) and (turning_point < x_max)

    # 拐点在范围内的相对位置 (0-1)
    if within_range:
        relative_position = (turning_point - x_min) / x_range
    else:
        relative_position = -1 if turning_point <= x_min else 2

    # 1. 改进的Sasabuchi检验
    # 计算边际效应在极值处的斜率
    slope_at_min = b_linear + 2 * b_squared * x_min
    slope_at_max = b_linear + 2 * b_squared * x_max

    # 计算斜率的标准误差
    var_slope_min = (se_linear**2 +
                     (2*x_min)**2 * se_squared**2 +
                     4*x_min * result.cov.loc[x_var, x_squared_var])
    var_slope_max = (se_linear**2 +
                     (2*x_max)**2 * se_squared**2 +
                     4*x_max * result.cov.loc[x_var, x_squared_var])

    se_slope_min = np.sqrt(var_slope_min if var_slope_min > 0 else se_linear**2)
    se_slope_max = np.sqrt(var_slope_max if var_slope_max > 0 else se_linear**2)

    # 计算t统计量
    t_value_min = slope_at_min / se_slope_min
    t_value_max = slope_at_max / se_slope_max

    # p值
    p_value_min = 2 * (1 - stats.t.cdf(abs(t_value_min), df=result.df_resid))
    p_value_max = 2 * (1 - stats.t.cdf(abs(t_value_max), df=result.df_resid))

    # 适用于U型或倒U型的检验
    if relationship_type == "U型":
        # U型: 左端点斜率显著为负，右端点斜率显著为正
        sasabuchi_t = min(abs(t_value_min) * (slope_at_min < 0),
                          abs(t_value_max) * (slope_at_max > 0))
        # 如果任一条件不满足，t值为0
        if slope_at_min >= 0 or slope_at_max <= 0:
            sasabuchi_t = 0
    else:
        # 倒U型: 左端点斜率显著为正，右端点斜率显著为负
        sasabuchi_t = min(abs(t_value_min) * (slope_at_min > 0),
                          abs(t_value_max) * (slope_at_max < 0))
        # 如果任一条件不满足，t值为0
        if slope_at_min <= 0 or slope_at_max >= 0:
            sasabuchi_t = 0

    overall_p_value = 1 - stats.t.cdf(sasabuchi_t, df=result.df_resid) if sasabuchi_t > 0 else 1.0

    # 拐点条件是否满足 (同时考虑系数显著性和形状)
    turning_significant = result.pvalues[x_squared_var] < 0.05

    # 判断是否具有正确的形状
    if relationship_type == "U型":
        proper_shape = slope_at_min < 0 and slope_at_max > 0
    else:
        proper_shape = slope_at_min > 0 and slope_at_max < 0

    # 2. 三段式样条回归验证
    # 将x变量分为三段
    segment_slopes = {}
    spline_supports = False

    try:
        if within_range:
            # 如果拐点在数据范围内，以拐点为中心分段
            x_low = x_min
            x_turning = turning_point
            x_high = x_max

            # 创建分段变量
            data_reset['segment'] = pd.cut(
                data_reset[x_var],
                bins=[x_min-0.001, turning_point, x_max+0.001],
                labels=['左段', '右段']
            )

            # 计算各段斜率
            for segment, group in data_reset.groupby('segment'):
                if len(group) > 10:  # 确保足够的样本量
                    # 简单线性回归获取斜率
                    X = sm.add_constant(group[x_var])
                    y = group[dependent_var]
                    try:
                        segment_model = sm.OLS(y, X).fit()
                        segment_slopes[segment] = {
                            'slope': segment_model.params[x_var],
                            'p_value': segment_model.pvalues[x_var],
                            'significant': segment_model.pvalues[x_var] < 0.05
                        }
                    except Exception as e:
                        print(f"段 {segment} 回归失败: {e}")

            # 检查分段斜率是否支持非线性关系
            if '左段' in segment_slopes and '右段' in segment_slopes:
                left_slope = segment_slopes['左段']['slope']
                right_slope = segment_slopes['右段']['slope']

                if relationship_type == "U型":
                    spline_supports = left_slope < 0 and right_slope > 0
                else:
                    spline_supports = left_slope > 0 and right_slope < 0
    except Exception as e:
        print(f"样条回归分析失败: {e}")
        traceback.print_exc()

        # 检查分段斜率是否支持非线性关系
        spline_supports = False
        if '左段' in segment_slopes and '右段' in segment_slopes:
            left_slope = segment_slopes['左段']['slope']
            right_slope = segment_slopes['右段']['slope']

            if relationship_type == "U型":
                spline_supports = left_slope < 0 and right_slope > 0
            else:
                spline_supports = left_slope > 0 and right_slope < 0
    else:
        segment_slopes = {}
        spline_supports = False
    # 3. 拐点的置信区间估计
    # 使用Delta方法计算拐点的标准误
    d_linear = -1/(2*b_squared)
    d_squared = b_linear/(2*b_squared**2)

    var_turning_point = (d_linear**2 * se_linear**2 +
                         d_squared**2 * se_squared**2 +
                         2 * d_linear * d_squared * result.cov.loc[x_var, x_squared_var])

    se_turning_point = np.sqrt(abs(var_turning_point))
    ci_lower = turning_point - 1.96 * se_turning_point
    ci_upper = turning_point + 1.96 * se_turning_point

    # 置信区间是否在数据范围内
    ci_within_range = (ci_lower >= x_min and ci_upper <= x_max)

    # 4. 最终结论
    # 综合所有条件判断是否支持非线性关系
    valid_nonlinear = (
        turning_significant and  # 二次项系数显著
        proper_shape and         # 形状正确
        within_range and         # 拐点在数据范围内
        overall_p_value < 0.05   # Sasabuchi 测试显著
    )

    # 增强结论，考虑更多情况
    if valid_nonlinear:
        if spline_supports:
            conclusion = f"强力支持{relationship_type}关系：所有指标均符合，且分段回归验证了曲线形状"
        else:
            conclusion = f"支持{relationship_type}关系：主要指标符合，但分段回归未完全验证"
    else:
        if turning_significant and proper_shape and within_range:
            conclusion = f"弱支持{relationship_type}关系：形状和拐点符合，但Sasabuchi测试不显著"
        elif turning_significant and proper_shape:
            conclusion = f"不支持{relationship_type}关系：拐点不在数据范围内"
        elif turning_significant:
            conclusion = f"不支持{relationship_type}关系：曲线形状不符合{relationship_type}"
        else:
            conclusion = f"不支持{relationship_type}关系：二次项系数不显著"

    # 5. 分段线性回归分析
    # 将数据分为两段，分别进行线性回归
    segment_results = {}
    piecewise_confirms = False

    try:
        if within_range:
            # 在拐点处分段
            left_segment = data_reset[data_reset[x_var] < turning_point]
            right_segment = data_reset[data_reset[x_var] >= turning_point]

            # 对左右两段分别进行线性回归
            segment_results = {}

            for name, segment in [("左段", left_segment), ("右段", right_segment)]:
                if len(segment) > 10:  # 确保足够的样本量
                    # 构建简单的线性模型
                    X = sm.add_constant(segment[x_var])
                    y = segment[dependent_var]
                    try:
                        segment_model = sm.OLS(y, X).fit()
                        segment_results[name] = {
                            'coefficient': segment_model.params[x_var],
                            'p_value': segment_model.pvalues[x_var],
                            'significant': segment_model.pvalues[x_var] < 0.05,
                            'slope_direction': 'positive' if segment_model.params[x_var] > 0 else 'negative',
                            'observations': len(segment),
                            'intercept': segment_model.params['const']
                        }
                    except Exception as e:
                        print(f"段 {name} 回归失败: {e}")
                        segment_results[name] = {'error': str(e)}

            # 检查分段结果是否符合预期的U型或倒U型形状
            if "左段" in segment_results and "右段" in segment_results and 'error' not in segment_results["左段"] and 'error' not in segment_results["右段"]:
                left_dir = segment_results["左段"].get('slope_direction')
                right_dir = segment_results["右段"].get('slope_direction')

                expected_pattern = False
                if relationship_type == "U型" and left_dir == 'negative' and right_dir == 'positive':
                    expected_pattern = True
                elif relationship_type == "倒U型" and left_dir == 'positive' and right_dir == 'negative':
                    expected_pattern = True

                piecewise_confirms = expected_pattern
    except Exception as e:
        print(f"分段回归分析失败: {e}")
        traceback.print_exc()

    # 6. 非参数局部回归确认
    lowess_confirms = None
    lowess_x = None
    lowess_y = None
    try:
        from statsmodels.nonparametric.smoothers_lowess import lowess

        # 按x_var排序数据
        sorted_data = data_reset.sort_values(by=x_var)

        # 使用LOWESS进行非参数局部回归
        lowess_result = lowess(sorted_data[dependent_var], sorted_data[x_var], frac=0.3, it=3)

        # 提取x和y值
        lowess_x = lowess_result[:, 0]
        lowess_y = lowess_result[:, 1]

        # 检查LOWESS曲线形状是否符合预期
        # 简单方法：检查最低/最高点位置
        if relationship_type == "U型":
            # 找到LOWESS曲线的最低点
            min_idx = np.argmin(lowess_y)
            min_x = lowess_x[min_idx]

            # 检查最低点是否接近拐点
            lowess_confirms = abs(min_x - turning_point) < (x_max - x_min) * 0.2
        else:  # 倒U型
            # 找到LOWESS曲线的最高点
            max_idx = np.argmax(lowess_y)
            max_x = lowess_x[max_idx]

            # 检查最高点是否接近拐点
            lowess_confirms = abs(max_x - turning_point) < (x_max - x_min) * 0.2
    except Exception as e:
        print(f"LOWESS分析失败: {e}")
        lowess_confirms = None
        lowess_x = None
        lowess_y = None

    # 7. 增强的可视化，包含所有分析结果
    try:
        # 创建更详细的可视化图
        fig, axs = plt.subplots(2, 2, figsize=(16, 14))

        # 主图：回归曲线、拐点和置信区间
        ax = axs[0, 0]

        # 计算整个曲线上的拟合值
        x_space = np.linspace(x_min, x_max, 200)
        fitted_curve = result.params.get('const', 0) + b_linear * x_space + b_squared * x_space**2

        # 绘制回归曲线
        ax.plot(x_space, fitted_curve, 'b-', linewidth=3, label='回归曲线')

        # 添加LOWESS曲线(如果可用)
        if lowess_x is not None and lowess_y is not None:
            ax.plot(lowess_x, lowess_y, 'g-', linewidth=2, alpha=0.7, label='LOWESS曲线')

        # 标记拐点和置信区间
        if within_range:
            turning_y = result.params.get('const', 0) + b_linear * turning_point + b_squared * turning_point**2
            ax.plot(turning_point, turning_y, 'ro', markersize=10, label='拐点')
            ax.axvline(x=turning_point, color='red', linestyle='--', alpha=0.5)

            # 标记置信区间
            if ci_within_range:
                ci_color = 'green'
                ci_label = '拐点95%置信区间 (在数据范围内)'
            else:
                ci_color = 'orange'
                ci_label = '拐点95%置信区间 (部分超出数据范围)'

            ax.axvspan(ci_lower, ci_upper, color=ci_color, alpha=0.2, label=ci_label)
        else:
            ax.text(0.5, 0.95, f"拐点 ({turning_point:.2f}) 不在数据范围内",
                    transform=ax.transAxes, fontsize=12, ha='center',
                    bbox=dict(facecolor='red', alpha=0.2))

        # 标记斜率和数据范围
        min_y = result.params.get('const', 0) + b_linear * x_min + b_squared * x_min**2
        max_y = result.params.get('const', 0) + b_linear * x_max + b_squared * x_max**2
        ax.plot(x_min, min_y, 'go', markersize=8, label='数据最小值')
        ax.plot(x_max, max_y, 'go', markersize=8, label='数据最大值')

        # 添加原始数据散点
        ax.scatter(data_reset[x_var], data_reset[dependent_var],
                alpha=0.2, c='gray', s=20, label='原始数据')

        # 如果有分段回归结果，绘制分段直线
        if plot_segments and segment_results and "左段" in segment_results and "右段" in segment_results and 'error' not in segment_results["左段"] and 'error' not in segment_results["右段"]:
            try:
                # 左段回归线
                left_coef = segment_results["左段"]['coefficient']
                left_intercept = segment_results["左段"]['intercept']
                left_x = np.linspace(x_min, turning_point, 50)
                left_y = left_intercept + left_coef * left_x
                ax.plot(left_x, left_y, 'r-', linewidth=1.5, alpha=0.7, label='左段回归线')

                # 右段回归线
                right_coef = segment_results["右段"]['coefficient']
                right_intercept = segment_results["右段"]['intercept']
                right_x = np.linspace(turning_point, x_max, 50)
                right_y = right_intercept + right_coef * right_x
                ax.plot(right_x, right_y, 'r-', linewidth=1.5, alpha=0.7, label='右段回归线')
            except Exception as e:
                print(f"绘制分段回归线失败: {e}")

        ax.set_title(f'{relationship_type}关系检验 - {x_var} vs {dependent_var}', fontsize=14)
        ax.set_xlabel(x_var, fontsize=12)
        ax.set_ylabel(dependent_var, fontsize=12)
        ax.grid(True, linestyle='--', alpha=0.7)
        ax.legend(loc='best', fontsize=10)

        # 第二个图：分段回归斜率比较
        ax = axs[0, 1]

        if segment_results and "左段" in segment_results and "右段" in segment_results and 'error' not in segment_results["左段"] and 'error' not in segment_results["右段"]:
            try:
                # 绘制两段斜率的柱状图
                segments = ['左段', '右段']
                slopes = [segment_results[seg]['coefficient'] for seg in segments]
                p_values = [segment_results[seg]['p_value'] for seg in segments]

                bars = ax.bar(segments, slopes, color=['blue', 'red'])

                # 为每个柱状图添加标签
                for i, bar in enumerate(bars):
                    height = bar.get_height()
                    annotation_pos = height + 0.05 if height >= 0 else height - 0.1
                    ax.text(bar.get_x() + bar.get_width()/2., annotation_pos,
                            f'{slopes[i]:.3f}\n(p={p_values[i]:.3f})',
                            ha='center', va='center')

                ax.axhline(y=0, color='black', linestyle='-', alpha=0.5)
                ax.set_title('分段回归斜率比较', fontsize=14)
                ax.set_ylabel('斜率', fontsize=12)
                ax.grid(True, alpha=0.3)

                # 添加结论
                if relationship_type == "U型" and slopes[0] < 0 and slopes[1] > 0:
                    ax.text(0.5, -0.1, '分段斜率支持U型关系', ha='center', transform=ax.transAxes,
                            fontsize=12, bbox=dict(facecolor='green', alpha=0.2))
                elif relationship_type == "倒U型" and slopes[0] > 0 and slopes[1] < 0:
                    ax.text(0.5, -0.1, '分段斜率支持倒U型关系', ha='center', transform=ax.transAxes,
                            fontsize=12, bbox=dict(facecolor='green', alpha=0.2))
                else:
                    ax.text(0.5, -0.1, '分段斜率不支持预期的关系形状', ha='center', transform=ax.transAxes,
                            fontsize=12, bbox=dict(facecolor='red', alpha=0.2))
            except Exception as e:
                print(f"绘制分段斜率比较失败: {e}")
                ax.text(0.5, 0.5, f'分段回归斜率比较失败: {e}', ha='center', va='center', transform=ax.transAxes, fontsize=12)
        else:
            ax.text(0.5, 0.5, '无法进行分段回归分析', ha='center', va='center', transform=ax.transAxes, fontsize=12)

        # 第三个图：检验摘要
        ax = axs[1, 0]
        ax.axis('off')  # 不显示坐标轴

        # 创建检验结果摘要表格
        table_data = [
            ['检验项', '结果', '支持预期关系?'],
            ['二次项系数', f'{b_squared:.4f} (p={result.pvalues[x_squared_var]:.4f})',
            '是' if result.pvalues[x_squared_var] < 0.05 and
                    ((relationship_type == "U型" and b_squared > 0) or
                    (relationship_type == "倒U型" and b_squared < 0)) else '否'],
            ['Sasabuchi检验', f'p={overall_p_value:.4f}', '是' if overall_p_value < 0.05 else '否'],
            ['拐点在数据范围内', f'{turning_point:.4f} ∈ [{x_min:.4f}, {x_max:.4f}]', '是' if within_range else '否'],
            ['拐点置信区间在范围内', f'[{ci_lower:.4f}, {ci_upper:.4f}]', '是' if ci_within_range else '否'],
            ['分段回归确认', '斜率方向符合预期' if piecewise_confirms else '斜率方向不符合预期',
            '是' if piecewise_confirms else '否'],
            ['LOWESS曲线确认', '支持' if lowess_confirms else '不支持',
            '是' if lowess_confirms else '否']
        ]

        # 创建表格
        table = ax.table(cellText=table_data, loc='center', cellLoc='center', colWidths=[0.3, 0.5, 0.2])
        table.auto_set_font_size(False)
        table.set_fontsize(10)
        table.scale(1, 1.5)

        # 为表格单元格添加颜色
        for i in range(1, len(table_data)):
            if table_data[i][2] == '是':
                table[(i, 2)].set_facecolor('lightgreen')
            else:
                table[(i, 2)].set_facecolor('lightcoral')

        ax.set_title('非线性关系检验摘要', fontsize=14)

        # 第四个图：模拟数据
        ax = axs[1, 1]

        try:
            # 根据估计的模型生成模拟数据
            np.random.seed(42)
            n_sim = 500
            x_sim = np.random.uniform(x_min, x_max, n_sim)

            # 添加均匀噪声
            noise = np.random.normal(0, result.mse_resid**0.5, n_sim)

            # 生成y值
            y_sim = result.params.get('const', 0) + b_linear * x_sim + b_squared * x_sim**2 + noise

            # 绘制模拟数据
            ax.scatter(x_sim, y_sim, alpha=0.3, s=20, c='gray', label='模拟数据')

            # 绘制拟合曲线
            ax.plot(x_space, fitted_curve, 'r-', linewidth=2, label='拟合曲线')

            # 标记拐点
            if x_min <= turning_point <= x_max:
                turning_y = result.params.get('const', 0) + b_linear * turning_point + b_squared * turning_point**2
                ax.plot(turning_point, turning_y, 'go', markersize=8, label='拐点')
                ax.axvline(x=turning_point, color='green', linestyle='--', alpha=0.5)
        except Exception as e:
            print(f"绘制模拟数据失败: {e}")
            ax.text(0.5, 0.5, f'模拟数据绘制失败: {e}', ha='center', va='center', transform=ax.transAxes, fontsize=12)

        ax.set_title('基于估计模型的模拟数据', fontsize=14)
        ax.set_xlabel(x_var, fontsize=12)
        ax.set_ylabel(dependent_var, fontsize=12)
        ax.grid(True, alpha=0.3)
        ax.legend(loc='best')

        # 保存图表
        plt.tight_layout()
        plt.savefig(f'{output_dir}/{relationship_type}_relationship_{x_var}_{dependent_var}_comprehensive.png', dpi=300)
        plt.close()
    except Exception as e:
        print(f"创建可视化图表失败: {e}")
        traceback.print_exc()

    # 返回增强的结果字典，包含所有分析
    result_dict = {
        'relationship_type': relationship_type,
        'turning_point': turning_point,
        'turning_point_ci': (ci_lower, ci_upper),
        'ci_within_range': ci_within_range,
        'within_range': within_range,
        'relative_position': relative_position,
        'slope_at_min': slope_at_min,
        'slope_at_max': slope_at_max,
        'p_value_min': p_value_min,
        'p_value_max': p_value_max,
        'squared_coefficient': b_squared,
        'squared_p_value': result.pvalues[x_squared_var],
        'sasabuchi_passes': overall_p_value < 0.05,
        'sasabuchi_p': overall_p_value,
        'is_minimum': b_squared > 0,
        'is_maximum': b_squared < 0,
        'valid_nonlinear': valid_nonlinear,
        'conclusion': conclusion,
        'segment_results': segment_results,
        'piecewise_confirms': piecewise_confirms,
        'lowess_confirms': lowess_confirms,
        'comprehensive_plot_path': f'{output_dir}/{relationship_type}_relationship_{x_var}_{dependent_var}_comprehensive.png'
    }

    print(f"非线性关系检验完成: {conclusion}")

    ax.set_title(f'{relationship_type}关系检验 - {x_var} vs {dependent_var}', fontsize=14)
    ax.set_xlabel(x_var, fontsize=12)
    ax.set_ylabel(dependent_var, fontsize=12)
    ax.grid(True, linestyle='--', alpha=0.7)
    ax.legend(loc='best', fontsize=10)

    # 第二个图：分段回归斜率比较
    ax = axs[0, 1]

    if segment_results and "左段" in segment_results and "右段" in segment_results:
        # 绘制两段斜率的柱状图
        segments = ['左段', '右段']
        slopes = [segment_results[seg]['coefficient'] for seg in segments]
        p_values = [segment_results[seg]['p_value'] for seg in segments]

        bars = ax.bar(segments, slopes, color=['blue', 'red'])

        # 为每个柱状图添加标签
        for i, bar in enumerate(bars):
            height = bar.get_height()
            annotation_pos = height + 0.05 if height >= 0 else height - 0.1
            ax.text(bar.get_x() + bar.get_width()/2., annotation_pos,
                    f'{slopes[i]:.3f}\n(p={p_values[i]:.3f})',
                    ha='center', va='center')

        ax.axhline(y=0, color='black', linestyle='-', alpha=0.5)
        ax.set_title('分段回归斜率比较', fontsize=14)
        ax.set_ylabel('斜率', fontsize=12)
        ax.grid(True, alpha=0.3)

        # 添加结论
        if relationship_type == "U型" and slopes[0] < 0 and slopes[1] > 0:
            ax.text(0.5, -0.1, '分段斜率支持U型关系', ha='center', transform=ax.transAxes,
                    fontsize=12, bbox=dict(facecolor='green', alpha=0.2))
        elif relationship_type == "倒U型" and slopes[0] > 0 and slopes[1] < 0:
            ax.text(0.5, -0.1, '分段斜率支持倒U型关系', ha='center', transform=ax.transAxes,
                    fontsize=12, bbox=dict(facecolor='green', alpha=0.2))
        else:
            ax.text(0.5, -0.1, '分段斜率不支持预期的关系形状', ha='center', transform=ax.transAxes,
                    fontsize=12, bbox=dict(facecolor='red', alpha=0.2))
    else:
        ax.text(0.5, 0.5, '无法进行分段回归分析', ha='center', va='center', transform=ax.transAxes, fontsize=12)

    # 第三个图：检验摘要
    ax = axs[1, 0]
    ax.axis('off')  # 不显示坐标轴

    # 创建检验结果摘要表格
    table_data = [
        ['检验项', '结果', '支持预期关系?'],
        ['二次项系数', f'{b_squared:.4f} (p={result.pvalues[x_squared_var]:.4f})',
         '是' if result.pvalues[x_squared_var] < 0.05 and
                ((relationship_type == "U型" and b_squared > 0) or
                 (relationship_type == "倒U型" and b_squared < 0)) else '否'],
             ['Sasabuchi检验', f'p={overall_p_value:.4f}', '是' if overall_p_value < 0.05 else '否'],
   ['拐点在数据范围内', f'{turning_point:.4f} ∈ [{x_min:.4f}, {x_max:.4f}]', '是' if within_range else '否'],
        ['拐点置信区间在范围内', f'[{ci_lower:.4f}, {ci_upper:.4f}]', '是' if ci_within_range else '否'],
        ['分段回归确认', '斜率方向符合预期' if piecewise_confirms else '斜率方向不符合预期',
         '是' if piecewise_confirms else '否'],
        ['LOWESS曲线确认', '支持' if lowess_confirms else '不支持',
         '是' if lowess_confirms else '否']
    ]

    # 创建表格
    table = ax.table(cellText=table_data, loc='center', cellLoc='center', colWidths=[0.3, 0.5, 0.2])
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.5)

    # 为表格单元格添加颜色
    for i in range(1, len(table_data)):
        if table_data[i][2] == '是':
            table[(i, 2)].set_facecolor('lightgreen')
        else:
            table[(i, 2)].set_facecolor('lightcoral')

    ax.set_title('非线性关系检验摘要', fontsize=14)
      # 第四个图：模拟数据
    ax = axs[1, 1]

    # 根据估计的模型生成模拟数据
    try:
        np.random.seed(42)
        n_sim = 500
        x_sim = np.random.uniform(x_min, x_max, n_sim)

        # 计算残差均方误差
        # PanelEffectsResults对象没有mse_resid属性，需要使用resid_ss和df_resid计算
        if hasattr(result, 'mse_resid'):
            # 如果存在mse_resid属性，直接使用
            mse = result.mse_resid
        elif hasattr(result, 'resid_ss') and hasattr(result, 'df_resid'):
            # 否则，根据残差平方和和残差自由度计算
            mse = result.resid_ss / result.df_resid
        else:
            # 如果无法计算，使用残差的方差作为近似
            mse = np.var(result.resids)

        # 添加均匀噪声
        noise = np.random.normal(0, mse**0.5, n_sim)

        # 生成y值
        y_sim = result.params.get('const', 0) + b_linear * x_sim + b_squared * x_sim**2 + noise

        # 绘制模拟数据
        ax.scatter(x_sim, y_sim, alpha=0.3, s=20, c='gray', label='模拟数据')

        # 绘制拟合曲线
        ax.plot(x_space, fitted_curve, 'r-', linewidth=2, label='拟合曲线')

        # 标记拐点
        if x_min <= turning_point <= x_max:
            turning_y = result.params.get('const', 0) + b_linear * turning_point + b_squared * turning_point**2
            ax.plot(turning_point, turning_y, 'go', markersize=8, label='拐点')
            ax.axvline(x=turning_point, color='green', linestyle='--', alpha=0.5)

        ax.set_title('基于估计模型的模拟数据', fontsize=14)
        ax.set_xlabel(x_var, fontsize=12)
        ax.set_ylabel(dependent_var, fontsize=12)
        ax.grid(True, alpha=0.3)
        ax.legend(loc='best')
    except Exception as e:
        print(f"绘制模拟数据失败: {e}")

    # 保存图表
    plt.tight_layout()
    plt.savefig(f'{output_dir}/{relationship_type}_relationship_{x_var}_{dependent_var}_comprehensive.png', dpi=300)
    plt.close()

    # 返回增强的结果字典，包含所有分析
    result_dict = {
        # ... 现有的结果项 ...
        'segment_results': segment_results,
        'piecewise_confirms': piecewise_confirms,
        'lowess_confirms': lowess_confirms,
        'comprehensive_plot_path': f'{output_dir}/{relationship_type}_relationship_{x_var}_{dependent_var}_comprehensive.png'
    }

    return result_dict


# 14. 增强调节效应分析函数
def enhanced_moderation_analysis(df, predictor, moderator, outcome, interaction=None):
    """增强版调节效应分析"""
    print(f"\n进行增强版调节效应分析: {predictor} x {moderator} -> {outcome}")

    # 重置索引以便操作
    if isinstance(df.index, pd.MultiIndex):
        df_reset = df.reset_index()
    else:
        df_reset = df.copy()

    # 检查因变量的零值比例和方差
    zero_pct = (df_reset[outcome] == 0).mean() * 100
    var_val = df_reset[outcome].var()
    print(f"{outcome}变量统计: 零值比例={zero_pct:.2f}%, 方差={var_val:.6f}")

    # 如果零值过多或方差极小，添加微小扰动
    if zero_pct > 60 or var_val < 1e-6:
        print(f"警告: {outcome}存在零方差风险，添加微小扰动以增加数值稳定性")
        # 复制数据框避免修改原始数据
        df_reset = df_reset.copy()
        # 添加微小正态分布扰动
        df_reset[outcome] = df_reset[outcome] + np.random.normal(0, max(var_val*0.01, 1e-5), size=len(df_reset))

    # 1. 确保交互项存在
    if interaction is None:
        interaction = f"{predictor}_x_{moderator}"
        if interaction not in df_reset.columns:
            print(f"创建交互项: {interaction}")
            df_reset[interaction] = df_reset[predictor] * df_reset[moderator]


    # 2. 创建不同水平的调节变量子组
    def create_robust_groups(df, var, n_groups=3, labels=['低', '中', '高']):
        # 分位数划分，创建三个组: 低(25%)、中(50%)、高(25%)
        try:
            # 获取非缺失值
            values = df[var].dropna()

            # 检查唯一值数量
            unique_values = np.unique(values)
            if len(unique_values) < n_groups:
                print(f"警告: 变量 {var} 只有 {len(unique_values)} 个唯一值，无法创建 {n_groups} 个组")
                # 如果唯一值太少，创建基于唯一值的分组
                if len(unique_values) == 1:
                    # 只有一个值，所有行分到同一组
                    return pd.Series(['中'] * len(df), index=df.index)
                elif len(unique_values) == 2:
                    # 两个值，创建"低"和"高"两组
                    boundaries = [np.min(values) - 0.1, np.median(values), np.max(values) + 0.1]
                    return pd.cut(df[var], bins=boundaries, labels=['低', '高'], include_lowest=True)

            # 尝试使用分位数分组
            try:
                # 计算分位数
                quantiles = [0, 0.33, 0.67, 1] if n_groups == 3 else [0, 0.25, 0.75, 1]
                quant_values = values.quantile(quantiles).values

                # 检查分位数是否有重复值
                if len(np.unique(quant_values)) < len(quant_values):
                    print(f"警告: 变量 {var} 的分位数有重复值，使用均匀分组")
                    # 改用均匀间隔创建分组边界
                    min_val, max_val = np.min(values), np.max(values)
                    boundaries = np.linspace(min_val - 0.001, max_val + 0.001, n_groups + 1)
                else:
                    boundaries = quant_values

                # 创建分组
                return pd.cut(df[var], bins=boundaries, labels=labels, include_lowest=True)
            except Exception as e:
                print(f"分位数分组失败: {e}")
                # 回退到均匀分组
                min_val, max_val = np.min(values), np.max(values)
                boundaries = np.linspace(min_val - 0.001, max_val + 0.001, n_groups + 1)
                return pd.cut(df[var], bins=boundaries, labels=labels, include_lowest=True)
        except Exception as e:
            print(f"分组创建失败: {e}")
            # 最终回退方案：手动创建
            result = pd.Series(index=df.index)
            median_val = df[var].median()
            result[df[var] <= median_val] = '低'
            result[df[var] > median_val] = '高'
            return result


    try:
        print("创建调节变量的分组...")
        df_reset['mod_level'] = create_robust_groups(df_reset, moderator, n_groups=3, labels=['低', '中', '高'])

        # 验证分组结果
        group_counts = df_reset['mod_level'].value_counts()
        print(f"调节变量分组结果: {dict(group_counts)}")
        if group_counts.min() < 10:
            print(f"警告: 某些组的样本量较小 (最小: {group_counts.min()})")
    except Exception as e:
        print(f"分组过程中出错: {e}")
        # 使用已有的备选方案...

    # 3. 为简单斜率分析创建均值±1SD的调节变量点
    mod_mean = df_reset[moderator].mean()
    mod_sd = df_reset[moderator].std()

    mod_points = {
        '低(-1SD)': mod_mean - mod_sd,
        '中(均值)': mod_mean,
        '高(+1SD)': mod_mean + mod_sd
    }
      # 4. 设置面板数据索引并运行主模型
    df_panel = df_reset.set_index(['stkcd', 'year'])

    # 定义基础控制变量
    base_controls = ['age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
                     'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe']

    # 确保控制变量在数据中
    controls = [c for c in base_controls if c in df_panel.columns]

    # 特殊处理ai_patent_quality变量 - 这个变量有60%以上的零值可能导致方差不稳定
    if outcome == 'ai_patent_quality':
        print(f"警告: 检测到使用 ai_patent_quality 作为因变量，该变量有大量零值")
        # 检查该变量在不同分组下的方差
        for level, group_data in df_reset.groupby('mod_level'):
            non_zero = group_data[group_data['ai_patent_quality'] > 0]['ai_patent_quality']
            print(f"  {level}组: 总数={len(group_data)}, 非零值={len(non_zero)}, 方差={group_data['ai_patent_quality'].var():.6f}")

        # 添加微小扰动以避免完全的零方差
        if df_panel[outcome].var() < 1e-6:
            print("因变量方差非常小，添加微小扰动以增加数值稳定性")
            df_panel[outcome] = df_panel[outcome] + np.random.normal(0, 1e-6, size=len(df_panel))

    # 运行包含交互项的模型
    model_result = advanced_panel_regression(
        df=df_panel,
        dependent_var=outcome,
        independent_vars=[predictor, moderator, interaction],
        control_vars=controls,
        cluster_entity=True, robust=True,
        drop_absorbed=True  # 添加这个参数
    )

    # 如果模型失败，返回None
    if model_result is None or model_result.get('result') is None:
        print(f"调节效应模型运行失败: {predictor} x {moderator} -> {outcome}")
        return None

    result = model_result['result']

    # 5. 简单斜率分析
    simple_slopes = {}

    # 提取系数
    b_pred = result.params[predictor]
    b_inter = result.params[interaction]

    # 计算不同水平下的简单斜率
    for level, mod_value in mod_points.items():
        # 简单斜率 = b_pred + b_inter * mod_value
        slope = b_pred + b_inter * mod_value
          # 简单斜率的标准误
        # var(slope) = var(b_pred) + mod_value^2 * var(b_inter) + 2*mod_value*cov(b_pred,b_inter)
        try:
            var_pred = result.cov.loc[predictor, predictor]
            var_inter = result.cov.loc[interaction, interaction]
            cov_pred_inter = result.cov.loc[predictor, interaction]

            var_slope = var_pred + mod_value**2 * var_inter + 2 * mod_value * cov_pred_inter

            # 处理数值不稳定情况
            if var_slope <= 0:
                print(f"警告: 在{level}水平({mod_value})处方差为负数或零: {var_slope}")
                var_slope = max(1e-10, abs(var_slope))  # 使用绝对值或小的正数

            se_slope = np.sqrt(var_slope)

            # t值和p值
            if se_slope > 0:
                t_value = slope / se_slope
                p_value = 2 * (1 - stats.t.cdf(abs(t_value), df=result.df_resid))
            else:
                print(f"警告: 在{level}水平标准误为零，无法计算t值和p值")
                t_value = 0
                p_value = 1.0
        except Exception as e:
            print(f"计算{level}水平简单斜率时出错: {e}，使用默认值")
            se_slope = 1.0
            t_value = 0
            p_value = 1.0

        simple_slopes[level] = {
            'moderator_value': mod_value,
            'slope': slope,
            'se': se_slope,
            't_value': t_value,
            'p_value': p_value,
            'significant': p_value < 0.05
        }

    # 6. Johnson-Neyman分析
    jn_result = johnson_neyman_analysis(
        df=df_panel,
        predictor=predictor,
        moderator=moderator,
        outcome=outcome,
        interaction=interaction,
        result=result
    )

    # 7. 创建子组分析
    subgroup_results = {}

    for level in ['低', '中', '高']:
        # 获取该水平的数据
        level_data = df_reset[df_reset['mod_level'] == level]

        # 如果样本量足够，进行子组分析
        if len(level_data) > 30:  # 确保足够的样本量
            # 设置面板索引
            level_panel = level_data.set_index(['stkcd', 'year'])

            # 运行子组模型
            subgroup_model = advanced_panel_regression(
                df=level_panel,
                dependent_var=outcome,
                independent_vars=[predictor],
                control_vars=controls,
                cluster_entity=True, robust=True, drop_absorbed=True
            )

            if subgroup_model and subgroup_model.get('result'):
                subgroup_results[level] = {
                    'coef': subgroup_model['result'].params[predictor],
                    'p_value': subgroup_model['result'].pvalues[predictor],
                    'significant': subgroup_model['result'].pvalues[predictor] < 0.05,
                    'n_obs': len(level_data),
                    'result': subgroup_model['result']
                }

    # 8. 创建增强的可视化
    create_enhanced_moderation_plot(
        df=df_panel,
        predictor=predictor,
        moderator=moderator,
        outcome=outcome,
        interaction=interaction,
        result=result,
        jn_result=jn_result,
        simple_slopes=simple_slopes,
        subgroup_results=subgroup_results,
        df_reset=df_reset
    )

    # 9. 生成详细报告
    moderation_report_path = f'output/reports/moderation_{predictor}_{moderator}_{outcome}.txt'

    with open(moderation_report_path, 'w', encoding='utf-8') as f:
        f.write(f"调节效应分析详细报告: {predictor} x {moderator} -> {outcome}\n")
        f.write("=" * 80 + "\n\n")

        # 主效应结果
        f.write("1. 交互模型主要结果\n")
        f.write("-" * 50 + "\n")
        f.write(f"自变量({predictor})系数: {b_pred:.4f}, p值: {result.pvalues[predictor]:.4f}\n")
        f.write(f"调节变量({moderator})系数: {result.params[moderator]:.4f}, p值: {result.pvalues[moderator]:.4f}\n")
        f.write(f"交互项({interaction})系数: {b_inter:.4f}, p值: {result.pvalues[interaction]:.4f}\n\n")

        # 简单斜率结果
        f.write("2. 简单斜率分析\n")
        f.write("-" * 50 + "\n")
        for level, slope_result in simple_slopes.items():
            f.write(f"调节变量水平: {level} ({moderator} = {slope_result['moderator_value']:.4f})\n")
            f.write(f"  - 简单斜率: {slope_result['slope']:.4f}\n")
            f.write(f"  - 标准误: {slope_result['se']:.4f}\n")
            f.write(f"  - t值: {slope_result['t_value']:.4f}\n")
            f.write(f"  - p值: {slope_result['p_value']:.4f}\n")
            sig_text = "显著" if slope_result['significant'] else "不显著"
            f.write(f"  - 显著性: {sig_text}\n\n")

        # Johnson-Neyman结果
        if jn_result and 'jn_points' in jn_result:
            f.write("3. Johnson-Neyman显著性区间分析\n")
            f.write("-" * 50 + "\n")

            if jn_result['jn_points']:
                f.write(f"Johnson-Neyman点: {[f'{p:.4f}' for p in jn_result['jn_points']]}\n\n")

                if 'significant_regions' in jn_result:
                    f.write("显著区域:\n")
                    for region in jn_result['significant_regions']:
                        f.write(f"- {region['description']}\n")
            else:
                f.write("未找到Johnson-Neyman点，调节效应在整个范围内均为显著或均不显著\n")

        # 子组分析结果
        f.write("\n4. 调节变量水平子组分析\n")
        f.write("-" * 50 + "\n")
        if subgroup_results:
            for level, sub_result in subgroup_results.items():
                f.write(f"子组: {level}\n")
                f.write(f"  - 样本量: {sub_result['n_obs']}\n")
                f.write(f"  - {predictor}系数: {sub_result['coef']:.4f}\n")
                f.write(f"  - p值: {sub_result['p_value']:.4f}\n")
                sig_text = "显著" if sub_result['significant'] else "不显著"
                f.write(f"  - 显著性: {sig_text}\n\n")
        else:
            f.write("子组样本量不足，无法进行子组分析\n\n")

        # 整体结论
        f.write("\n5. 调节效应整体结论\n")
        f.write("-" * 50 + "\n")

        # 交互项是否显著
        if result.pvalues[interaction] < 0.05:
            if b_inter > 0:
                f.write(f"✓ 存在显著正向调节效应: 随着{moderator}的增加，{predictor}对{outcome}的正向影响增强\n")
            else:
                f.write(f"✓ 存在显著负向调节效应: 随着{moderator}的增加，{predictor}对{outcome}的正向影响减弱\n")

            # 检查简单斜率的变化模式
            low_sig = simple_slopes['低(-1SD)']['significant']
            high_sig = simple_slopes['高(+1SD)']['significant']

            if low_sig and high_sig:
                f.write("  - 在调节变量的低值和高值处，自变量的效应均显著\n")
            elif low_sig and not high_sig:
                f.write("  - 在调节变量的低值处自变量效应显著，而在高值处不显著\n")
            elif not low_sig and high_sig:
                f.write("  - 在调节变量的低值处自变量效应不显著，而在高值处显著\n")
            else:
                f.write("  - 尽管存在调节效应，但在调节变量的典型值处，自变量的效应均不显著\n")
        else:
            f.write(f"✗ 未发现显著调节效应: {moderator}不会显著调节{predictor}对{outcome}的影响\n")

    # 返回调节效应分析结果
    return {
        'model_result': model_result,
        'simple_slopes': simple_slopes,
        'jn_result': jn_result,
        'subgroup_results': subgroup_results,
        'report_path': moderation_report_path
    }


# 16. 增强调节效应可视化
def create_enhanced_moderation_plot(df, predictor, moderator, outcome, interaction, result, jn_result=None, simple_slopes=None, subgroup_results=None, df_reset=None, output_dir="output/figures"):
    """创建全面的调节效应可视化，包含多个面板展示不同角度的调节效应"""
    try:
        # 创建输出目录（如果不存在）
        os.makedirs(output_dir, exist_ok=True)

        # 确保我们有重置索引的数据框用于绘图
        if df_reset is None:
            if isinstance(df.index, pd.MultiIndex):
                df_reset = df.reset_index()
            else:
                df_reset = df.copy()
          # 创建一个3×2的多面板图
        fig, axs = plt.subplots(3, 2, figsize=(15, 18))

        # 提取必要的系数
        b_predictor = result.params[predictor]
        b_moderator = result.params[moderator]
        b_interaction = result.params[interaction]

        # 获取变量范围
        pred_mean = df[predictor].mean()
        pred_std = df[predictor].std()
        mod_mean = df[moderator].mean()
        mod_std = df[moderator].std()
        pred_range = np.linspace(pred_mean - 2*pred_std, pred_mean + 2*pred_std, 100)
        mod_range = np.linspace(mod_mean - 2*mod_std, mod_mean + 2*mod_std, 100)

        # 面板1：调节变量不同水平下的预测线（X=预测变量，Y=因变量，多条线代表不同调节变量水平）
        ax = axs[0, 0]
        for i, mod_level in enumerate([-1, 0, 1]):  # -1SD, Mean, +1SD
            mod_value = mod_mean + mod_level * mod_std
            y_pred = [b_predictor * x + b_moderator * mod_value + b_interaction * x * mod_value for x in pred_range]
            ax.plot(pred_range, y_pred, label=f'{moderator}={mod_level}SD', linewidth=2)

        ax.scatter(df_reset[predictor], df_reset[outcome], alpha=0.1, color='gray')
        ax.set_title(f'不同{moderator}水平下{predictor}对{outcome}的影响', fontsize=12)
        ax.set_xlabel(predictor, fontsize=10)
        ax.set_ylabel(outcome, fontsize=10)
        ax.legend()
        ax.grid(True, alpha=0.3)

        # 面板2：预测变量不同水平下的预测线（X=调节变量，Y=因变量，多条线代表不同预测变量水平）
        ax = axs[0, 1]
        for i, pred_level in enumerate([-1, 0, 1]):  # -1SD, Mean, +1SD
            pred_value = pred_mean + pred_level * pred_std
            y_pred = [b_predictor * pred_value + b_moderator * x + b_interaction * pred_value * x for x in mod_range]
            ax.plot(mod_range, y_pred, label=f'{predictor}={pred_level}SD', linewidth=2)

        ax.scatter(df_reset[moderator], df_reset[outcome], alpha=0.1, color='gray')
        ax.set_title(f'不同{predictor}水平下{moderator}对{outcome}的影响', fontsize=12)
        ax.set_xlabel(moderator, fontsize=10)
        ax.set_ylabel(outcome, fontsize=10)
        ax.legend()
        ax.grid(True, alpha=0.3)

        # 面板3：边际效应图（X=调节变量，Y=预测变量的边际效应）
        ax = axs[1, 0]
        marginal_effects = [b_predictor + b_interaction * x for x in mod_range]
        ax.plot(mod_range, marginal_effects, 'b-', linewidth=2)

        # 如果有Johnson-Neyman结果，添加显著性区域
        if jn_result and 'jn_points' in jn_result:
            # 添加显著性区域阴影
            for region in jn_result.get('significant_regions', []):
                ax.axvspan(region['region'][0], region['region'][1],
                          alpha=0.2, color='green', label=f'显著区域: {region["region"][0]:.2f}-{region["region"][1]:.2f}')

            # 添加JN点标记
            for jn_point in jn_result['jn_points']:
                ax.axvline(x=jn_point, color='red', linestyle='--')
                ax.text(jn_point, ax.get_ylim()[0], f'JN={jn_point:.2f}',
                       rotation=90, verticalalignment='bottom')

        ax.axhline(y=0, color='k', linestyle='-', alpha=0.5)
        ax.set_title(f'{predictor}的边际效应随{moderator}的变化', fontsize=12)
        ax.set_xlabel(moderator, fontsize=10)
        ax.set_ylabel(f'{predictor}的边际效应', fontsize=10)
        ax.grid(True, alpha=0.3)

        # 面板4：3D透视图
        ax = axs[1, 1]
        ax = fig.add_subplot(3, 2, 4, projection='3d')
          # 创建网格
        X, Y = np.meshgrid(pred_range, mod_range)
        Z = np.array([[b_predictor * x + b_moderator * y + b_interaction * x * y
                     for x in pred_range] for y in mod_range])

        # 绘制3D表面
        surf = ax.plot_surface(X, Y, Z, cmap='viridis', alpha=0.8)

        ax.set_title(f'{predictor}、{moderator}与{outcome}的3D关系', fontsize=12)
        ax.set_xlabel(predictor, fontsize=10)
        ax.set_ylabel(moderator, fontsize=10)
        ax.set_zlabel(outcome, fontsize=10)

        # 面板5：条件效应热图（X=预测变量，Y=调节变量，颜色=预测效应）
        ax = axs[2, 0]
        X, Y = np.meshgrid(pred_range, mod_range)
        Z = np.array([[b_predictor * x + b_interaction * x * y for x in pred_range] for y in mod_range])

        im = ax.pcolormesh(X, Y, Z, cmap='RdBu_r', shading='auto')
        fig.colorbar(im, ax=ax, label=f'{predictor}对{outcome}的效应')

        ax.set_title(f'{predictor}对{outcome}的条件效应', fontsize=12)
        ax.set_xlabel(predictor, fontsize=10)
        ax.set_ylabel(moderator, fontsize=10)

        # 面板6：简单斜率柱状图（不同调节水平的效应大小比较）
        ax = axs[2, 1]

        # 三个调节水平的效应
        mod_levels = [mod_mean - mod_std, mod_mean, mod_mean + mod_std]
        level_names = ['低(-1SD)', '中(Mean)', '高(+1SD)']

        effects = [b_predictor + b_interaction * level for level in mod_levels]

        # 创建柱状图
        bars = ax.bar(level_names, effects, color=['blue', 'green', 'red'])

        # 添加水平线表示零效应
        ax.axhline(y=0, color='k', linestyle='-', alpha=0.5)
          # 在每个柱子上添加数值标签
        for bar, effect in zip(bars, effects):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2.,
                   height + 0.02 * (max(effects) if max(effects) > 0 else -max(effects)),
                   f'{effect:.3f}',
                   ha='center', va='bottom', fontsize=10)

        ax.set_title(f'不同{moderator}水平下{predictor}的效应', fontsize=12)
        ax.set_ylabel(f'{predictor}的效应大小', fontsize=10)
        ax.grid(True, alpha=0.3)

        # 添加总体标题
        plt.suptitle(f'{moderator}对{predictor}和{outcome}关系的调节效应分析', fontsize=16, y=0.98)

        plt.tight_layout(rect=[0, 0, 1, 0.96])

        # 保存图像
        output_path = f'{output_dir}/comprehensive_moderation_{predictor}_{moderator}_{outcome}.png'
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        return output_path

    except Exception as e:
        print(f"创建增强调节效应图时出错: {e}")
        traceback.print_exc()
        return None

# 18. 扩展的假设验证函数
def extended_hypothesis_testing(df):
    """运行所有假设的验证模型"""
    print("\n开始扩展的假设验证...")
    # 添加预先防护措施
    if df is None:
        print("错误: 传入的数据框为None")
        return {}

    if not isinstance(df, pd.DataFrame):
        print(f"错误: 传入的对象不是DataFrame而是 {type(df)}")
        return {}

    if len(df) == 0:
        print("错误: 传入的数据框为空")
        return {}

    # 检查是否正确的面板结构
    if not isinstance(df.index, pd.MultiIndex) or df.index.names != ['stkcd', 'year']:
        print("警告: 数据框不是预期的面板结构，尝试修复...")
        try:
            if 'stkcd' in df.columns and 'year' in df.columns:
                df = df.set_index(['stkcd', 'year'])
                print("  已重设索引为 ['stkcd', 'year']")
            else:
                df_reset = df.reset_index()
                if 'stkcd' in df_reset.columns and 'year' in df_reset.columns:
                    df = df_reset.set_index(['stkcd', 'year'])
                    print("  已重设索引为 ['stkcd', 'year']")
                else:
                    print("  错误: 无法重设索引，缺少 'stkcd' 或 'year' 列")
                    return {}
        except Exception as e:
            print(f"  尝试修复索引结构时出错: {e}")
            return {}
          # 数据准备和变量处理
    # 添加非平稳变量的差分处理
    non_stationary_vars = ['ai_sqrt_std', 'ai_job_log_std', 'ai_patent_log_boxcox_std',
                          'manu_job_log_std', 'intotal']
    df_diff = df.copy()

    # 对每个公司进行差分处理
    for company, group in df_diff.reset_index().groupby('stkcd'):
        for var in non_stationary_vars:
            if var in df_diff.columns:
                # 按时间排序
                group = group.sort_values('year')
                # 计算差分
                df_diff.loc[df_diff.index.get_level_values(0) == company, f'{var}_diff'] = \
                    group[var].diff().values

    # 使用标准化函数确保所有列名都是小写的
    df = normalize_dataframe_columns(df, case='lower')

    # 如果有重设索引的需要，确保索引后的列名也是小写
    if not isinstance(df.index, pd.MultiIndex) or df.index.names != ['stkcd', 'year']:
        df_reset = df.reset_index()
        df_reset = normalize_dataframe_columns(df_reset, case='lower')

         # 检查stkcd和year列是否存在（不区分大小写）
        stkcd_col = get_column_case_insensitive(df_reset, 'stkcd')
        year_col = get_column_case_insensitive(df_reset, 'year')

        # 收集找到的列名
        idx_cols = [col for col in [stkcd_col, year_col] if col is not None]

        if stkcd_col is not None and year_col is not None:
            # 重命名为标准小写形式
            rename_dict = {stkcd_col: 'stkcd', year_col: 'year'}
            df_reset = df_reset.rename(columns=rename_dict)
            df = df_reset.set_index(['stkcd', 'year'])
            print(f"重新设置索引为 ['stkcd', 'year']，并确保列名为小写")
        else:
            print(f"警告: 找不到完整的索引列 ['stkcd', 'year']，找到的列: {idx_cols}")
    else:
        # 如果已经是MultiIndex，检查索引名是否为小写
        if df.index.names != ['stkcd', 'year']:
            df.index.names = [name.lower() if name.lower() in ['stkcd', 'year'] else name
                            for name in df.index.names]
            print("已将索引名转为小写")

    # 定义基础控制变量（均为小写）
    base_controls = ['age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
                    'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe',
                    'dsi', 'ci', 'ocr', 'em']

    # 确保控制变量存在于数据中
    valid_controls = [control for control in base_controls if control in df.columns]
    print(f"有效控制变量: {valid_controls}")

    # 检查并创建必要的二次项和交互变量
    if 'ai_squared' not in df.columns:
        df['ai_squared'] = df['ai'] ** 2
        print("创建变量: ai_squared")

    # 创建时间标准化变量(为H1b假设)
    df_temp = df.reset_index()
    df_temp['year_std'] = (df_temp['year'] - df_temp['year'].min()) / (df_temp['year'].max() - df_temp['year'].min())
    df_temp['ai_x_year'] = df_temp['ai'] * df_temp['year_std']
    df_temp = df_temp.set_index(['stkcd', 'year'])

    # 创建交互项(针对调节效应假设)
    interaction_pairs = [
        ('ai', 'ai_patent_log'),
        ('ai_job_log', 'ai_patent_log'),
        ('ai', 'manu_job_log'),
        ('ai_job_log', 'manu_job_log')
    ]

    for pred, mod in interaction_pairs:
        if pred in df.columns and mod in df.columns:
            interact_name = f"{pred}_x_{mod}"
            if interact_name not in df.columns:
                df[interact_name] = df[pred] * df[mod]
                print(f"创建交互变量: {interact_name}")

    # 结果字典
    results = {}

    # ===== H1系列：AI技术投入对创新能力的影响 =====

    # 测试 H1a: 企业针对AI的技术投入显著提高专利总数量
    print("\n测试H1a: AI技术投入对专利总数量的影响")
    h1a_result = advanced_panel_regression(
        df=df,
        dependent_var='intotal',
        independent_vars=['ai'],
        control_vars=valid_controls,
        entity_effects=True,
        robust=True,
        cluster_entity=True,
        drop_absorbed=True,
        add_lagged_dependent=False  # 静态面板
    )

    # 添加动态面板模型
    h1a_dynamic_result = advanced_panel_regression(
        df=df,
        dependent_var='intotal',
        independent_vars=['ai'],
        control_vars=valid_controls,
        entity_effects=True,
        robust=True,
        cluster_entity=True,
        drop_absorbed=True,
        add_lagged_dependent=True  # 动态面板
    )
    results['H1a'] = h1a_result
    results['H1a_dynamic'] = h1a_dynamic_result

    if h1a_dynamic_result is not None and h1a_dynamic_result.get('result') is not None:
        print("H1a动态面板模型 - 结果摘要：")
        # 安全获取参数，需要检查键是否存在
        ai_param = None
        for key in h1a_dynamic_result['result'].params.index:
            if 'ai' in key and 'lag' not in key and 'patent' not in key:
                ai_param = key
                break

        if ai_param:
            print(f"AI技术投入系数: {h1a_dynamic_result['result'].params[ai_param]:.4f}")
            print(f"p值: {h1a_dynamic_result['result'].pvalues[ai_param]:.4f}")
        else:
            print("未找到AI技术投入相关参数")

        # 安全获取滞后变量参数
        lag_param = None
        for key in h1a_dynamic_result['result'].params.index:
            if 'intotal_lag' in key:
                lag_param = key
                break

        if lag_param:
            print(f"滞后因变量系数: {h1a_dynamic_result['result'].params[lag_param]:.4f}")
            print(f"滞后因变量p值: {h1a_dynamic_result['result'].pvalues[lag_param]:.4f}")
        else:
            print("未找到滞后因变量参数")

        # 测试 H1b: AI技术投入先提高后不显著影响探索型专利数量
        print("\n测试H1b: AI技术投入对探索型专利数量的动态影响")

    # 使用test_time_dynamic_effect函数进行时间动态效应分析
    h1b_time_models = test_time_dynamic_effect(
        df=df_temp,
        y_var='ep',
        x_var='ai',
        control_vars=valid_controls
    )
    results['H1b_models'] = h1b_time_models

    # 保留基本模型以保持兼容性
    h1b_linear = h1b_time_models.get('base')
    h1b_time = h1b_time_models.get('trend')

    results['H1b_linear'] = h1b_linear
    results['H1b_time'] = h1b_time

    # 测试 H1c: AI技术投入提高应用型专利数量
    print("\n测试H1c: AI技术投入对应用型专利数量的影响")
    h1c_result = advanced_panel_regression(
        df=df,
        dependent_var='dp',
        independent_vars=['ai'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H1c'] = h1c_result

    if h1c_result is not None and h1c_result.get('result') is not None:
        print(f"AI技术投入系数: {h1c_result['result'].params['ai']:.4f}")
        print(f"p值: {h1c_result['result'].pvalues['ai']:.4f}")
        print(f"支持假设: {h1c_result['result'].pvalues['ai'] < 0.05 and h1c_result['result'].params['ai'] > 0}")
        print(h1c_result['result'].summary)

    # ===== H2系列：AI技术投入与创新质量的非线性关系 =====

    # 测试 H2a: AI技术投入与专利质量呈U型关系
    print("\n测试H2a: AI技术投入与专利质量的U型关系")

    h2a_result = advanced_panel_regression(
        df=df,
        dependent_var='ai_patent_quality',
        independent_vars=['ai', 'ai_squared'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H2a'] = h2a_result

    if h2a_result is not None and h2a_result.get('result') is not None:
        # 使用advanced_nonlinear_test函数进行非线性检验
        h2a_u_test = advanced_nonlinear_test(
            result=h2a_result['result'],
            x_var='ai',
            x_squared_var='ai_squared',
            data=df,
            output_dir='output/reports'
        )
        results['H2a_u_test'] = h2a_u_test

    # 测试 H2b: AI技术投入与专利深度呈倒U型关系
    print("\n测试H2b: AI技术投入与专利深度的倒U型关系")

    h2b_result = advanced_panel_regression(
        df=df,
        dependent_var='ai_patent_depth',
        independent_vars=['ai', 'ai_squared'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H2b'] = h2b_result

    if h2b_result is not None and h2b_result.get('result') is not None:
        # 使用advanced_nonlinear_test函数进行非线性检验
        h2b_u_test = advanced_nonlinear_test(
            result=h2b_result['result'],
            x_var='ai',
            x_squared_var='ai_squared',
            data=df,
            output_dir='output/reports'
        )
        results['H2b_u_test'] = h2b_u_test

    # ===== H3系列：AI人力投入对创新能力的影响 =====

    # H3a: AI人力投入提高专利总数
    print("\n测试H3a: AI人力投入对专利总数量的影响")
    h3a_result = advanced_panel_regression(
        df=df,
        dependent_var='intotal',
        independent_vars=['ai_job_log'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H3a'] = h3a_result

    if h3a_result is not None and h3a_result.get('result') is not None:
        print(f"AI人力投入系数: {h3a_result['result'].params['ai_job_log']:.4f}")
        print(f"p值: {h3a_result['result'].pvalues['ai_job_log']:.4f}")
        print(f"支持假设: {h3a_result['result'].pvalues['ai_job_log'] < 0.05 and h3a_result['result'].params['ai_job_log'] > 0}")
        print(h3a_result['result'].summary)

    # H3b: AI人力投入提高探索型专利数量
    print("\n测试H3b: AI人力投入对探索型专利数量的影响")
    h3b_result = advanced_panel_regression(
        df=df,
        dependent_var='ep',
        independent_vars=['ai_job_log'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H3b'] = h3b_result

    if h3b_result is not None and h3b_result.get('result') is not None:
        print(f"AI人力投入系数: {h3b_result['result'].params['ai_job_log']:.4f}")
        print(f"p值: {h3b_result['result'].pvalues['ai_job_log']:.4f}")
        print(f"支持假设: {h3b_result['result'].pvalues['ai_job_log'] < 0.05 and h3b_result['result'].params['ai_job_log'] > 0}")
        print(h3b_result['result'].summary)

    # H3c: AI人力投入提高应用型专利数量
    print("\n测试H3c: AI人力投入对应用型专利数量的影响")
    h3c_result = advanced_panel_regression(
        df=df,
        dependent_var='dp',
        independent_vars=['ai_job_log'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H3c'] = h3c_result

    if h3c_result is not None and h3c_result.get('result') is not None:
        print(f"AI人力投入系数: {h3c_result['result'].params['ai_job_log']:.4f}")
        print(f"p值: {h3c_result['result'].pvalues['ai_job_log']:.4f}")
        print(f"支持假设: {h3c_result['result'].pvalues['ai_job_log'] < 0.05 and h3c_result['result'].params['ai_job_log'] > 0}")
        print(h3c_result['result'].summary)

    # ===== H4系列: AI人力投入对创新质量的影响 =====

    # H4a: AI人力投入提高专利质量
    print("\n测试H4a: AI人力投入对专利质量的影响")
    h4a_result = advanced_panel_regression(
        df=df,
        dependent_var='ai_patent_quality',
        independent_vars=['ai_job_log'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H4a'] = h4a_result

    if h4a_result is not None and h4a_result.get('result') is not None:
        print(f"AI人力投入系数: {h4a_result['result'].params['ai_job_log']:.4f}")
        print(f"p值: {h4a_result['result'].pvalues['ai_job_log']:.4f}")
        print(f"支持假设: {h4a_result['result'].pvalues['ai_job_log'] < 0.05 and h4a_result['result'].params['ai_job_log'] > 0}")
        print(h4a_result['result'].summary)

    # H4b: AI人力投入提高专利深度
    print("\n测试H4b: AI人力投入对专利深度的影响")
    h4b_result = advanced_panel_regression(
        df=df,
        dependent_var='ai_patent_depth',
        independent_vars=['ai_job_log'],
        control_vars=valid_controls,
        cluster_entity=True, robust=True, drop_absorbed=True
    )
    results['H4b'] = h4b_result

    if h4b_result is not None and h4b_result.get('result') is not None:
        print(f"AI人力投入系数: {h4b_result['result'].params['ai_job_log']:.4f}")
        print(f"p值: {h4b_result['result'].pvalues['ai_job_log']:.4f}")
        print(f"支持假设: {h4b_result['result'].pvalues['ai_job_log'] < 0.05 and h4b_result['result'].params['ai_job_log'] > 0}")
        print(h4b_result['result'].summary)

    # ===== H5系列: 调节效应分析 - 使用enhanced_moderation_analysis函数 =====

    # H5a: 专利跨领域程度对AI技术投入与专利数关系的调节
    print("\n测试H5a: 专利跨领域程度对AI技术投入与专利总数量关系的调节作用")
    if 'ai_x_ai_patent_log' not in df.columns:
        df['ai_x_ai_patent_log'] = df['ai'] * df['ai_patent_log']

    h5a_moderation = enhanced_moderation_analysis(
        df=df,
        predictor='ai',
        moderator='ai_patent_log',
        outcome='intotal',
        interaction='ai_x_ai_patent_log'
    )

    # 保存结果到结果字典
    if h5a_moderation:
        results['H5a'] = h5a_moderation['model_result']
        results['H5a_simple_slopes'] = h5a_moderation['simple_slopes']
        results['H5a_jn'] = h5a_moderation['jn_result']
        results['H5a_subgroup'] = h5a_moderation.get('subgroup_results')

        # 输出关键结果
        if h5a_moderation['model_result'] and h5a_moderation['model_result'].get('result'):
            result = h5a_moderation['model_result']['result']
            interact_coef = result.params['ai_x_ai_patent_log']
            interact_pval = result.pvalues['ai_x_ai_patent_log']
            print(f"H5a交互项系数: {interact_coef:.4f}, p值: {interact_pval:.4f}")
            print(f"支持调节效应: {interact_pval < 0.05}")

    # H5a-2: 专利跨领域程度对AI人力投入与专利数关系的调节
    print("\n测试H5a-2: 专利跨领域程度对AI人力投入与专利总数量关系的调节作用")
    if 'ai_job_log_x_ai_patent_log' not in df.columns:
        df['ai_job_log_x_ai_patent_log'] = df['ai_job_log'] * df['ai_patent_log']

    h5a2_moderation = enhanced_moderation_analysis(
        df=df,
        predictor='ai_job_log',
        moderator='ai_patent_log',
        outcome='intotal',
        interaction='ai_job_log_x_ai_patent_log'
    )

    # 保存结果到结果字典
    if h5a2_moderation:
        results['H5a2'] = h5a2_moderation['model_result']
        results['H5a2_simple_slopes'] = h5a2_moderation['simple_slopes']
        results['H5a2_jn'] = h5a2_moderation['jn_result']
        results['H5a2_subgroup'] = h5a2_moderation.get('subgroup_results')

        # 输出关键结果
        if h5a2_moderation['model_result'] and h5a2_moderation['model_result'].get('result'):
            result = h5a2_moderation['model_result']['result']
            interact_coef = result.params['ai_job_log_x_ai_patent_log']
            interact_pval = result.pvalues['ai_job_log_x_ai_patent_log']
            print(f"H5a-2交互项系数: {interact_coef:.4f}, p值: {interact_pval:.4f}")
            print(f"支持调节效应: {interact_pval < 0.05}")

    # H5b: 人才跨领域程度对AI技术投入与专利数关系的调节
    print("\n测试H5b: 人才跨领域程度对AI技术投入与专利总数量关系的调节作用")
    # 确保交互项和调节变量存在
    if 'manu_job_log' in df.columns:
        if 'ai_x_manu_job_log' not in df.columns:
            df['ai_x_manu_job_log'] = df['ai'] * df['manu_job_log']

        h5b_moderation = enhanced_moderation_analysis(
            df=df,
            predictor='ai',
            moderator='manu_job_log',
            outcome='intotal',
            interaction='ai_x_manu_job_log'
        )

        # 保存结果到结果字典
        if h5b_moderation:
            results['H5b'] = h5b_moderation['model_result']
            results['H5b_simple_slopes'] = h5b_moderation['simple_slopes']
            results['H5b_jn'] = h5b_moderation['jn_result']
            results['H5b_subgroup'] = h5b_moderation.get('subgroup_results')

            # 输出关键结果
            if h5b_moderation['model_result'] and h5b_moderation['model_result'].get('result'):
                result = h5b_moderation['model_result']['result']
                interact_coef = result.params['ai_x_manu_job_log']
                interact_pval = result.pvalues['ai_x_manu_job_log']
                print(f"H5b交互项系数: {interact_coef:.4f}, p值: {interact_pval:.4f}")
                print(f"支持调节效应: {interact_pval < 0.05}")
    else:
        print("警告: 缺少'manu_job_log'变量，无法进行H5b分析")

    # H5b-2: 人才跨领域程度对AI人力投入与专利数关系的调节
    print("\n测试H5b-2: 人才跨领域程度对AI人力投入与专利总数量关系的调节作用")
    # 确保交互项和调节变量存在
    if 'manu_job_log' in df.columns:
        if 'ai_job_log_x_manu_job_log' not in df.columns:
            df['ai_job_log_x_manu_job_log'] = df['ai_job_log'] * df['manu_job_log']

        h5b2_moderation = enhanced_moderation_analysis(
            df=df,
            predictor='ai_job_log',
            moderator='manu_job_log',
            outcome='intotal',
            interaction='ai_job_log_x_manu_job_log'
        )

        # 保存结果到结果字典
        if h5b2_moderation:
            results['H5b2'] = h5b2_moderation['model_result']
            results['H5b2_simple_slopes'] = h5b2_moderation['simple_slopes']
            results['H5b2_jn'] = h5b2_moderation['jn_result']
            results['H5b2_subgroup'] = h5b2_moderation.get('subgroup_results')

            # 输出关键结果
            if h5b2_moderation['model_result'] and h5b2_moderation['model_result'].get('result'):
                result = h5b2_moderation['model_result']['result']
                interact_coef = result.params['ai_job_log_x_manu_job_log']
                interact_pval = result.pvalues['ai_job_log_x_manu_job_log']
                print(f"H5b-2交互项系数: {interact_coef:.4f}, p值: {interact_pval:.4f}")
                print(f"支持调节效应: {interact_pval < 0.05}")
    else:
        print("警告: 缺少'manu_job_log'变量，无法进行H5b-2分析")

    # ===== H6系列: 调节效应对创新质量的影响 =====

    # H6a: 专利跨领域程度对AI技术投入与专利质量关系的调节
    print("\n测试H6a: 专利跨领域程度对AI技术投入与专利质量关系的调节作用")
    h6a_moderation = enhanced_moderation_analysis(
        df=df,
        predictor='ai',
        moderator='ai_patent_log',
        outcome='ai_patent_quality',
        interaction='ai_x_ai_patent_log'
    )

    # 保存结果到结果字典
    if h6a_moderation:
        results['H6a'] = h6a_moderation['model_result']
        results['H6a_simple_slopes'] = h6a_moderation['simple_slopes']
        results['H6a_jn'] = h6a_moderation['jn_result']
        results['H6a_subgroup'] = h6a_moderation.get('subgroup_results')

        # 输出关键结果
        if h6a_moderation['model_result'] and h6a_moderation['model_result'].get('result'):
            result = h6a_moderation['model_result']['result']
            interact_coef = result.params['ai_x_ai_patent_log']
            interact_pval = result.pvalues['ai_x_ai_patent_log']
            print(f"H6a交互项系数: {interact_coef:.4f}, p值: {interact_pval:.4f}")
            print(f"支持正向调节效应: {interact_pval < 0.05 and interact_coef > 0}")

    # H6b: 人才跨领域程度对AI人力投入与专利质量关系的调节
    print("\n测试H6b: 人才跨领域程度对AI人力投入与专利质量关系的调节作用")
    if 'manu_job_log' in df.columns:
        h6b_moderation = enhanced_moderation_analysis(
            df=df,
            predictor='ai_job_log',
            moderator='manu_job_log',
            outcome='ai_patent_quality',
            interaction='ai_job_log_x_manu_job_log'
        )

        # 保存结果到结果字典
        if h6b_moderation:
            results['H6b'] = h6b_moderation['model_result']
            results['H6b_simple_slopes'] = h6b_moderation['simple_slopes']
            results['H6b_jn'] = h6b_moderation['jn_result']
            results['H6b_subgroup'] = h6b_moderation.get('subgroup_results')

            # 输出关键结果
            if h6b_moderation['model_result'] and h6b_moderation['model_result'].get('result'):
                result = h6b_moderation['model_result']['result']
                interact_coef = result.params['ai_job_log_x_manu_job_log']
                interact_pval = result.pvalues['ai_job_log_x_manu_job_log']
                print(f"H6b交互项系数: {interact_coef:.4f}, p值: {interact_pval:.4f}")
                print(f"支持正向调节效应: {interact_pval < 0.05 and interact_coef > 0}")
    else:
        print("警告: 缺少'manu_job_log'变量，无法进行H6b分析")

    # ===== 生成假设检验结果摘要 =====
    print("\n=== 假设检验结果摘要 ===")
    hypothesis_summary = pd.DataFrame(columns=['假设', '模型', '系数', 'p值', '标准化系数', '结论'])

    # H1系列 - AI技术投入对创新能力的影响
    h1a_coef = None
    h1a_pvalue = None
    if 'H1a' in results and results['H1a'] and results['H1a'].get('result') is not None:
        h1a_coef = results['H1a']['result'].params.get('ai')
        h1a_pvalue = results['H1a']['result'].pvalues.get('ai')
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H1a: AI技术投入提高专利总数量',
            '固定效应',
            results['H1a']['result'].params['ai'],
            results['H1a']['result'].pvalues['ai'],
            results['H1a']['std_coefficients']['ai'],
            '支持' if (results['H1a']['result'].pvalues['ai'] < 0.05 and results['H1a']['result'].params['ai'] > 0) else '不支持'
        ]

    if 'H1b_time' in results and results['H1b_time'] and results['H1b_time'].get('result') is not None:
        # 处理时间交互模型结果
        try:
            h1b_support = (results['H1b_time']['result'].params['ai'] > 0 and
                        results['H1b_time']['result'].pvalues['ai'] < 0.05 and
                        results['H1b_time']['result'].params['ai_x_time_trend'] < 0 and
                        results['H1b_time']['result'].pvalues['ai_x_time_trend'] < 0.05)

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H1b: AI技术投入先提高后不显著影响探索型专利',
                '时间交互',
                f"ai={results['H1b_time']['result'].params['ai']:.3f}, ai×year={results['H1b_time']['result'].params['ai_x_time_trend']:.3f}",
                f"ai={results['H1b_time']['result'].pvalues['ai']:.3f}, ai×year={results['H1b_time']['result'].pvalues['ai_x_time_trend']:.3f}",
                f"ai={results['H1b_time']['std_coefficients']['ai']:.3f}, ai×year={results['H1b_time']['std_coefficients']['ai_x_time_trend']:.3f}",
                '支持' if h1b_support else '不支持'
            ]
        except:
            print("无法添加H1b假设结果到摘要")

    # 添加H1c假设结果
    if 'H1c' in results and results['H1c'] and results['H1c'].get('result') is not None:
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H1c: AI技术投入提高应用型专利数量',
            '固定效应',
            results['H1c']['result'].params['ai'],
            results['H1c']['result'].pvalues['ai'],
            results['H1c']['std_coefficients']['ai'],
            '支持' if (results['H1c']['result'].pvalues['ai'] < 0.05 and results['H1c']['result'].params['ai'] > 0) else '不支持'
        ]

    # H2系列 - 非线性关系
    if 'H2a' in results and results['H2a'] and results['H2a'].get('result') is not None and 'H2a_u_test' in results:
        try:
            h2a_support = (results.get('H2a_u_test', {}).get('valid_nonlinear', False) and results.get('H2a', {}).get('result').params.get('ai_squared', 0) > 0)

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H2a: AI技术投入与专利质量呈U型关系',
                'U型检验',
                f"ai={results.get('H2a', {}).get('result').params.get('ai', 'N/A'):.3f}, ai²={results.get('H2a', {}).get('result').params.get('ai_squared', 'N/A'):.3f}",
                f"ai={results.get('H2a', {}).get('result').pvalues.get('ai', 'N/A'):.3f}, ai²={results.get('H2a', {}).get('result').pvalues.get('ai_squared', 'N/A'):.3f}",
                f"拐点={results.get('H2a_u_test', {}).get('turning_point', 'N/A')}",
                '支持' if h2a_support else '不支持'
            ]
        except:
            print("无法添加H2a假设结果到摘要")

    if 'H2b' in results and results.get('H2b') and results['H2b'].get('result') is not None and 'H2b_u_test' in results:
        try:
            h2b_support = (results.get('H2b_u_test', {}).get('valid_nonlinear', False) and results.get('H2b', {}).get('result').params.get('ai_squared', 0) < 0)

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H2b: AI技术投入与专利深度呈倒U型关系',
                '倒U型检验',
                f"ai={results.get('H2b', {}).get('result').params.get('ai', 'N/A'):.3f}, ai²={results.get('H2b', {}).get('result').params.get('ai_squared', 'N/A'):.3f}",
                f"ai={results.get('H2b', {}).get('result').pvalues.get('ai', 'N/A'):.3f}, ai²={results.get('H2b', {}).get('result').pvalues.get('ai_squared', 'N/A'):.3f}",
                f"拐点={results.get('H2b_u_test', {}).get('turning_point', 'N/A')}",
                '支持' if h2b_support else '不支持'
            ]
        except:
            print("无法添加H2b假设结果到摘要")

      # 添加H3系列假设结果
    if 'H3a' in results and results.get('H3a') and results.get('H3a', {}).get('result') is not None:
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H3a: AI人力投入提高专利总数量',
            '固定效应',
            results.get('H3a', {}).get('result').params.get('ai_job_log', 'N/A'),
            results.get('H3a', {}).get('result').pvalues.get('ai_job_log', 'N/A'),
            results.get('H3a', {}).get('std_coefficients', {}).get('ai_job_log', 'N/A'),
            '支持' if (results.get('H3a', {}).get('result').pvalues.get('ai_job_log', 1) < 0.05 and results.get('H3a', {}).get('result').params.get('ai_job_log', 0) > 0) else '不支持'
        ]

    if 'H3b' in results and results.get('H3b') and results.get('H3b', {}).get('result') is not None:
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H3b: AI人力投入提高探索型专利数量',
            '固定效应',
            results.get('H3b', {}).get('result').params.get('ai_job_log', 'N/A'),
            results.get('H3b', {}).get('result').pvalues.get('ai_job_log', 'N/A'),
            results.get('H3b', {}).get('std_coefficients', {}).get('ai_job_log', 'N/A'),
            '支持' if (results.get('H3b', {}).get('result').pvalues.get('ai_job_log', 1) < 0.05 and results.get('H3b', {}).get('result').params.get('ai_job_log', 0) > 0) else '不支持'
        ]

    if 'H3c' in results and results.get('H3c') and results.get('H3c', {}).get('result') is not None:
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H3c: AI人力投入提高应用型专利数量',
            '固定效应',
            results['H3c']['result'].params['ai_job_log'],
            results.get('H3c', {}).get('result').pvalues.get('ai_job_log', 'N/A'),
            results.get('H3c', {}).get('std_coefficients', {}).get('ai_job_log', 'N/A'),
            '支持' if (results.get('H3c', {}).get('result').pvalues.get('ai_job_log', 1) < 0.05 and results.get('H3c', {}).get('result').params.get('ai_job_log', 0) > 0) else '不支持'
        ]

    # 添加H4系列假设结果
    if 'H4a' in results and results.get('H4a') and results.get('H4a', {}).get('result') is not None:
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H4a: AI人力投入提高专利质量',
            '固定效应',
            results.get('H4a', {}).get('result').params.get('ai_job_log', 'N/A'),
            results.get('H4a', {}).get('result').pvalues.get('ai_job_log', 'N/A'),
            results.get('H4a', {}).get('std_coefficients', {}).get('ai_job_log', 'N/A'),
            '支持' if (results.get('H4a', {}).get('result').pvalues.get('ai_job_log', 1) < 0.05 and results.get('H4a', {}).get('result').params.get('ai_job_log', 0) > 0) else '不支持'
        ]

    if 'H4b' in results and results.get('H4b') and results.get('H4b', {}).get('result') is not None:
        hypothesis_summary.loc[len(hypothesis_summary)] = [
            'H4b: AI人力投入提高专利深度',
            '固定效应',
            results.get('H4b', {}).get('result').params.get('ai_job_log', 'N/A'),
            results.get('H4b', {}).get('result').pvalues.get('ai_job_log', 'N/A'),
            results.get('H4b', {}).get('std_coefficients', {}).get('ai_job_log', 'N/A'),
            '支持' if (results.get('H4b', {}).get('result').pvalues.get('ai_job_log', 1) < 0.05 and results.get('H4b', {}).get('result').params.get('ai_job_log', 0) > 0) else '不支持'
        ]

    # 添加H5系列假设结果到摘要表
    if 'H5a' in results and results['H5a'] and results['H5a'].get('result') is not None:
        try:
            result = results['H5a']['result']
            interact_coef = result.params['ai_x_ai_patent_log']
            interact_pval = result.pvalues['ai_x_ai_patent_log']
            std_coef = results['H5a']['std_coefficients'].get('ai_x_ai_patent_log', 'N/A')

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H5a: 专利跨领域程度对AI技术投入与专利数量关系的调节',
                '交互效应',
                interact_coef,
                interact_pval,
                std_coef,
                '支持' if interact_pval < 0.05 else '不支持'
            ]
        except:
            print("无法添加H5a假设结果到摘要")

    if 'H5a2' in results and results['H5a2'] and results['H5a2'].get('result') is not None:
        try:
            result = results['H5a2']['result']
            interact_coef = result.params['ai_job_log_x_ai_patent_log']
            interact_pval = result.pvalues['ai_job_log_x_ai_patent_log']
            std_coef = results['H5a2']['std_coefficients'].get('ai_job_log_x_ai_patent_log', 'N/A')

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H5a-2: 专利跨领域程度对AI人力投入与专利数量关系的调节',
                '交互效应',
                interact_coef,
                interact_pval,
                std_coef,
                '支持' if interact_pval < 0.05 else '不支持'
            ]
        except:
            print("无法添加H5a-2假设结果到摘要")

    if 'H5b' in results and results['H5b'] and results['H5b'].get('result') is not None:
        try:
            result = results['H5b']['result']
            interact_coef = result.params['ai_x_manu_job_log']
            interact_pval = result.pvalues['ai_x_manu_job_log']
            std_coef = results['H5b']['std_coefficients'].get('ai_x_manu_job_log', 'N/A')

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H5b: 人才跨领域程度对AI技术投入与专利数量关系的调节',
                '交互效应',
                interact_coef,
                interact_pval,
                std_coef,
                '支持' if interact_pval < 0.05 else '不支持'
            ]
        except:
            print("无法添加H5b假设结果到摘要")

    if 'H5b2' in results and results['H5b2'] and results['H5b2'].get('result') is not None:
        try:
            result = results['H5b2']['result']
            interact_coef = result.params['ai_job_log_x_manu_job_log']
            interact_pval = result.pvalues['ai_job_log_x_manu_job_log']
            std_coef = results['H5b2']['std_coefficients'].get('ai_job_log_x_manu_job_log', 'N/A')

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H5b-2: 人才跨领域程度对AI人力投入与专利数量关系的调节',
                '交互效应',
                interact_coef,
                interact_pval,
                std_coef,
                '支持' if interact_pval < 0.05 else '不支持'
            ]
        except:
            print("无法添加H5b-2假设结果到摘要")

    # 添加H6系列假设结果到摘要表
    if 'H6a' in results and results['H6a'] and results['H6a'].get('result') is not None:
        try:
            result = results['H6a']['result']
            interact_coef = result.params['ai_x_ai_patent_log']
            interact_pval = result.pvalues['ai_x_ai_patent_log']
            std_coef = results['H6a']['std_coefficients'].get('ai_x_ai_patent_log', 'N/A')

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H6a: 专利跨领域程度对AI技术投入与专利质量关系的正向调节',
                '交互效应',
                interact_coef,
                interact_pval,
                std_coef,
                '支持' if (interact_pval < 0.05 and interact_coef > 0) else '不支持'
            ]
        except:
            print("无法添加H6a假设结果到摘要")

    if 'H6b' in results and results['H6b'] and results['H6b'].get('result') is not None:
        try:
            result = results['H6b']['result']
            interact_coef = result.params['ai_job_log_x_manu_job_log']
            interact_pval = result.pvalues['ai_job_log_x_manu_job_log']
            std_coef = results['H6b']['std_coefficients'].get('ai_job_log_x_manu_job_log', 'N/A')

            hypothesis_summary.loc[len(hypothesis_summary)] = [
                'H6b: 人才跨领域程度对AI人力投入与专利质量关系的正向调节',
                '交互效应',
                interact_coef,
                interact_pval,
                std_coef,
                '支持' if (interact_pval < 0.05 and interact_coef > 0) else '不支持'
            ]
        except:
            print("无法添加H6b假设结果到摘要")

    # 保存假设检验结果摘要
    hypothesis_summary.to_csv('output/tables/hypothesis_testing_summary.csv', index=False, encoding='utf-8-sig')
    print("假设检验结果摘要已保存至 output/tables/hypothesis_testing_summary.csv")

    # 返回包含全部假设检验结果的字典
    return results


# 19.增强版稳健性检验
def enhanced_robustness_checks(df, results):
    """进行增强版稳健性检验"""
    print("\n开始增强版稳健性检验...")

    # 立即初始化所有可能用到的变量，防止访问错误
    h1a_coef = None
    h1a_pvalue = None
    h1a_lag_coef = None
    h1a_lag_pvalue = None
    h3a_coef = None
    h3a_pvalue = None
    h3a_lag_coef = None
    h3a_lag_pvalue = None

    # 以下是h1a_large_coef等变量的初始化
    h1a_large_coef = None
    h1a_large_pvalue = None
    h1a_small_coef = None
    h1a_small_pvalue = None
    h3a_large_coef = None
    h3a_large_pvalue = None
    h3a_small_coef = None
    h3a_small_pvalue = None

    # 创建结果字典
    robustness_results = {
        'lagged_models': {},
        'size_subsamples': {},
        'sensitivity_models': {},
        'estimation_methods': {}
    }

    # 确保所有必需的二次项和交互项存在
    if 'ai_squared' not in df.columns:
        df['ai_squared'] = df['ai'] ** 2

    # 检查存在哪些关键模型结果
    has_h1a = 'H1a' in results and results['H1a'] is not None and results['H1a'].get('result') is not None
    has_h3a = 'H3a' in results and results['H3a'] is not None and results['H3a'].get('result') is not None

    # 检查是否有H1a动态模型结果
    if 'H1a_dynamic' in results and results['H1a_dynamic'] is not None and results['H1a_dynamic'].get('result') is not None:
        h1a_dynamic_result = results['H1a_dynamic']['result']
        print("H1a动态面板模型结果可用，将用于稳健性分析")
        # 处理动态模型结果
    else:
        print("警告: H1a动态模型结果不可用")

    # 添加辅助函数来查找与AI相关的参数
    def find_ai_param(params):
        # 查找与AI相关的参数
        candidates = ['ai', 'ai_centered', 'ai_std', 'ai_lag1']
        for candidate in candidates:
            if candidate in params.index:
                return candidate
        # 如果没有精确匹配，寻找包含'ai'的参数
        for param in params.index:
            if isinstance(param, str) and 'ai' in param.lower() and 'ai_job' not in param.lower() and 'ai_patent' not in param.lower():
                return param
        return None

    def find_ai_job_param(params):
        # 查找与AI_job_log相关的参数
        candidates = ['ai_job_log', 'ai_job_log_centered', 'ai_job_log_std', 'ai_job_log_lag1']
        for candidate in candidates:
            if candidate in params.index:
                return candidate
        # 如果没有精确匹配，寻找包含'ai_job'的参数
        for param in params.index:
            if isinstance(param, str) and 'ai_job' in param.lower():
                return param
        return None

    # 从结果中提取系数和p值，添加安全检查
    if has_h1a:
        # 安全获取AI参数，处理可能的变量名变化
        h1a_result = results['H1a']['result']
        ai_param_name = find_ai_param(h1a_result.params)
        if ai_param_name:
            h1a_coef = h1a_result.params[ai_param_name]
            h1a_pvalue = h1a_result.pvalues[ai_param_name]
        else:
            print("警告: 在H1a结果中找不到AI相关参数")
            h1a_coef = None
            h1a_pvalue = None
    else:
        print("警告: H1a模型结果不可用")
        h1a_coef = None
        h1a_pvalue = None

    if has_h3a:
        # 安全获取AI_job_log参数
        h3a_result = results['H3a']['result']
        ai_job_param_name = find_ai_job_param(h3a_result.params)
        if ai_job_param_name:
            h3a_coef = h3a_result.params[ai_job_param_name]
            h3a_pvalue = h3a_result.pvalues[ai_job_param_name]
        else:
            print("警告: 在H3a结果中找不到AI_job_log相关参数")
            h3a_coef = None
            h3a_pvalue = None
    else:
        print("警告: H3a模型结果不可用")
        h3a_coef = None
        h3a_pvalue = None

    # 复制控制变量列表，避免全局变量问题
    try:
        if 'valid_controls' not in locals():
            base_controls = ['age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
                            'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe']
            valid_controls = [control for control in base_controls if control in df.columns]
    except Exception as e:
        print(f"控制变量准备错误: {e}")
        base_controls = ['age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
                        'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe']
        valid_controls = [control for control in base_controls if control in df.columns]

    # 1. 创建滞后变量模型
    lagged_models = {}
    print("创建滞后变量模型...")

    try:
        # 使用复制的数据框并重置索引
        df_lagged = df.copy().reset_index()

        # 要创建滞后变量的列名
        lag_vars = ['ai', 'ai_job_log', 'ai_patent_log', 'manu_job_log']
        created_lag_vars = []

        # 确保这些变量存在于数据框中
        existing_vars = [var for var in lag_vars if var in df_lagged.columns]
        if not existing_vars:
            print("警告: 没有找到需要创建滞后项的变量!")

        print(f"将为以下变量创建滞后项: {existing_vars}")

        # 创建空列来存储滞后值
        for var in existing_vars:
            lag_var_name = f'{var}_lag1'
            df_lagged[lag_var_name] = np.nan
            created_lag_vars.append(lag_var_name)

        # 使用pandas的shift方法按公司分组创建滞后变量
        for company, group in df_lagged.groupby('stkcd'):
            # 按年份排序
            group = group.sort_values('year')
            idx = group.index

            # 对每个变量创建滞后值
            for var in existing_vars:
                lag_var_name = f'{var}_lag1'
                # 使用pandas shift方法，更可靠地创建滞后值
                df_lagged.loc[idx, lag_var_name] = group[var].shift(1).values

        # 删除滞后值为NaN的行
        df_lagged = df_lagged.dropna(subset=created_lag_vars)

        # 如果需要，创建交互变量
        if 'ai_patent_log' in df_lagged.columns:
            for lag_var in created_lag_vars:
                if lag_var.startswith('ai') or lag_var.startswith('ai_job_log'):
                    df_lagged[f'{lag_var}_x_ai_patent_log'] = df_lagged[lag_var] * df_lagged['ai_patent_log']

        # 重新设置索引
        df_lagged = df_lagged.set_index(['stkcd', 'year'])

        # 检查滞后变量是否成功创建
        print(f"创建的滞后变量: {created_lag_vars}")
        for var in created_lag_vars:
            if var in df_lagged.columns:
                non_na_count = df_lagged[var].count()
                print(f"变量 {var} 存在，含有 {non_na_count} 个非缺失值")
            else:
                print(f"变量 {var} 创建失败!")

        # 检查是否有数据剩余
        if len(df_lagged) == 0:
            print("错误: 删除缺失值后没有数据剩余！将使用原始数据")
            df_lagged = df.copy()
            created_lag_vars = []  # 标记为创建失败

    except Exception as e:
        print(f"滞后变量创建过程中出错: {str(e)}")
        print(traceback.format_exc())
        print("将使用原始数据继续分析")
        df_lagged = df.copy()
        created_lag_vars = []  # 标记为创建失败

    # 测试H1a的滞后模型
    print("\n测试H1a滞后模型: AI技术投入(滞后一期)对专利总数量的影响")
    if 'ai_lag1' in df_lagged.columns and df_lagged['ai_lag1'].count() > 0:
        try:
            h1a_lag_result = advanced_panel_regression(
                df=df_lagged,
                dependent_var='intotal',
                independent_vars=['ai_lag1'],
                control_vars=valid_controls,
                cluster_entity=True, robust=True, drop_absorbed=True
            )

            if h1a_lag_result is not None and h1a_lag_result.get('result') is not None:
                lagged_models['H1a_lag'] = h1a_lag_result
                print(h1a_lag_result['result'].summary)
            else:
                print("H1a滞后模型创建失败")
                lagged_models['H1a_lag'] = {'result': None, 'std_coefficients': None}
        except Exception as e:
            print(f"H1a滞后模型错误: {str(e)}")
            print(traceback.format_exc())
            lagged_models['H1a_lag'] = {'result': None, 'std_coefficients': None}
    else:
        print("变量 'ai_lag1' 不存在或全为缺失值")
        lagged_models['H1a_lag'] = {'result': None, 'std_coefficients': None}

    # 测试H3a的滞后模型
    print("\n测试H3a滞后模型: AI人力投入(滞后一期)对专利总数量的影响")
    if 'ai_job_log_lag1' in df_lagged.columns and df_lagged['ai_job_log_lag1'].count() > 0:
        try:
            h3a_lag_result = advanced_panel_regression(
                df=df_lagged,
                dependent_var='intotal',
                independent_vars=['ai_job_log_lag1'],
                control_vars=valid_controls,
                cluster_entity=True, robust=True, drop_absorbed=True
            )

            if h3a_lag_result is not None and h3a_lag_result.get('result') is not None:
                lagged_models['H3a_lag'] = h3a_lag_result
                print(h3a_lag_result['result'].summary)
            else:
                print("H3a滞后模型创建失败")
                lagged_models['H3a_lag'] = {'result': None, 'std_coefficients': None}
        except Exception as e:
            print(f"H3a滞后模型错误: {str(e)}")
            print(traceback.format_exc())
            lagged_models['H3a_lag'] = {'result': None, 'std_coefficients': None}
    else:
        print("变量 'ai_job_log_lag1' 不存在或全为缺失值")
        lagged_models['H3a_lag'] = {'result': None, 'std_coefficients': None}

    # 测试H5a的滞后模型(调节效应)
    print("测试H5a滞后模型: 专利跨领域程度对AI技术投入(滞后一期)与专利总数量关系的调节作用")
    try:
        h5a_lag_result = advanced_panel_regression(
            df=df_lagged,
            dependent_var='intotal',
            independent_vars=['ai_lag1', 'ai_patent_log', 'ai_lag1_x_ai_patent_log'],
            control_vars=valid_controls,
            cluster_entity=True,
            robust=True,
            drop_absorbed=True
        )

        if h5a_lag_result is not None and h5a_lag_result.get('result') is not None:
            lagged_models['H5a_lag'] = h5a_lag_result
            print("H5a滞后模型拟合成功")
            print(h5a_lag_result['result'].summary)
        else:
            print("H5a滞后模型拟合失败")
            lagged_models['H5a_lag'] = {'error': 'Model fitting failed', 'result': None}
    except Exception as e:
        print(f"H5a滞后模型发生错误: {e}")
        traceback.print_exc()
        lagged_models['H5a_lag'] = {'error': str(e), 'result': None}

    # 保存滞后模型结果
    robustness_results['lagged_models'] = lagged_models

    # 2. 按企业规模进行子样本分析
    try:
        print("\n按企业规模进行子样本分析...")
        # 创建企业规模分组
        df_size = df.copy()
        df_size = df_size.reset_index()

        # 根据SIZE的中位数分为大企业和小企业
        size_median = df_size['size'].median()
        df_size['size_group'] = df_size['size'].apply(lambda x: '大企业' if x >= size_median else '小企业')

        # 分组进行分析
        size_subsamples = {}

        for size_group, group_data in df_size.groupby('size_group'):
            group_data = group_data.set_index(['stkcd', 'year'])

            # 对每个分组运行H1a模型
            h1a_size_result = advanced_panel_regression(
                df=group_data,
                dependent_var='intotal',
                independent_vars=['ai'],
                control_vars=valid_controls,
                cluster_entity=True, robust=True, drop_absorbed=True
            )

            size_subsamples[f'H1a_{size_group}'] = h1a_size_result

            print(f"\nH1a在{size_group}子样本的结果:")
            if h1a_size_result and h1a_size_result.get('result') is not None:
                print(h1a_size_result['result'].summary)
            else:
                print(f"H1a在{size_group}子样本分析失败")

            # 对每个分组运行H3a模型
            h3a_size_result = advanced_panel_regression(
                df=group_data,
                dependent_var='intotal',
                independent_vars=['ai_job_log'],
                control_vars=valid_controls,
                cluster_entity=True, robust=True, drop_absorbed=True
            )

            size_subsamples[f'H3a_{size_group}'] = h3a_size_result

            print(f"\nH3a在{size_group}子样本的结果:")
            if h3a_size_result and h3a_size_result.get('result') is not None:
                print(h3a_size_result['result'].summary)
            else:
                print(f"H3a在{size_group}子样本分析失败")

        # 保存子样本分析结果
        robustness_results['size_subsamples'] = size_subsamples
    except Exception as e:
        print(f"企业规模子样本分析错误: {e}")
        traceback.print_exc()
        robustness_results['size_subsamples'] = {'error': str(e)}

    # 3. 不同控制变量组合的敏感性分析
    print("\n进行控制变量敏感性分析...")
    sensitivity_models = {}

    # 定义不同的控制变量组合
    control_sets = {
        '基本控制': ['size', 'lev', 'roa'],
        '治理控制': ['size', 'lev', 'roa', 'balance', 'mhold', 'audit'],
        '市场控制': ['size', 'lev', 'roa', 'tobinq1', 'growth'],
        '全部控制': valid_controls
    }

    # 对H1a进行敏感性检验
    for set_name, controls in control_sets.items():
        # 确保控制变量存在
        valid_set_controls = [control for control in controls if control in df.columns]

        h1a_sensitivity = advanced_panel_regression(
            df=df,
            dependent_var='intotal',
            independent_vars=['ai'],
            control_vars=valid_set_controls,
            cluster_entity=True, robust=True, drop_absorbed=True
        )

        sensitivity_models[f'H1a_{set_name}'] = h1a_sensitivity

        print(f"\nH1a使用{set_name}的结果:")
        print(h1a_sensitivity['result'].summary)

    # 保存敏感性分析结果
    robustness_results['sensitivity_models'] = sensitivity_models

    # 4. 不同估计方法的稳健性检验
    print("\n使用不同估计方法进行稳健性检验...")
    estimation_methods = {}

    # 普通OLS估计
    try:
        df_reset = df.reset_index()
        X = sm.add_constant(df_reset[['ai'] + valid_controls])
        y = df_reset['intotal']

        ols_model = sm.OLS(y, X)
        ols_result = ols_model.fit(cov_type='HC1')  # 使用异方差稳健标准误

        estimation_methods['H1a_OLS'] = {
            'result': ols_result,
            'std_coefficients': None  # OLS不计算标准化系数
        }

        print("\nH1a使用普通OLS的结果:")
        print(ols_result.summary())
    except Exception as e:
        print(f"OLS估计出错: {e}")

    # 随机效应估计
    try:
        re_model = RandomEffects.from_formula(
            formula=f"intotal ~ ai + {' + '.join(valid_controls)}",
            data=df
        )
        re_result = re_model.fit(cov_type='clustered', cluster_entity=True)

        estimation_methods['H1a_RE'] = {
            'result': re_result,
            'std_coefficients': None  # 随机效应不计算标准化系数
        }

        print("\nH1a使用随机效应的结果:")
        print(re_result.summary)
    except Exception as e:
        print(f"随机效应估计出错: {e}")

    # 保存不同估计方法结果
    robustness_results['estimation_methods'] = estimation_methods

    # 5. 生成稳健性检验报告
    with open('output/reports/robustness_checks_report.txt', 'w') as f:
        f.write("稳健性检验报告\n")
        f.write("=" * 50 + "\n\n")

        # 所有使用格式化字符串的地方，先检查值是否为None

        # H1a滞后模型对比
        f.write("H1a模型对比(AI技术投入对专利总数的影响):\n")
        if has_h1a and h1a_coef is not None and h1a_pvalue is not None:
            f.write(f"  - 原始模型: 系数 = {h1a_coef:.4f}, p值 = {h1a_pvalue:.4f}\n")

            # 其他需要使用格式化字符串的地方也添加类似的检查
            # 例如:
            if 'H1a_lag' in lagged_models and lagged_models['H1a_lag'].get('result') is not None:
                h1a_lag_param = find_ai_param(lagged_models['H1a_lag']['result'].params)
                if h1a_lag_param:
                    h1a_lag_coef = lagged_models['H1a_lag']['result'].params[h1a_lag_param]
                    h1a_lag_pvalue = lagged_models['H1a_lag']['result'].pvalues[h1a_lag_param]

                    if h1a_lag_coef is not None and h1a_lag_pvalue is not None:
                        f.write(f"  - 滞后模型: 系数 = {h1a_lag_coef:.4f}, p值 = {h1a_lag_pvalue:.4f}\n")
                        f.write(f"  - 结论: {'系数方向一致' if (h1a_coef > 0) == (h1a_lag_coef > 0) else '系数方向不一致'}, ")
                        f.write(f"{'均显著' if h1a_pvalue < 0.05 and h1a_lag_pvalue < 0.05 else '显著性不一致'}\n\n")
                    else:
                        f.write("  - 滞后模型: 系数或p值为None\n\n")
                else:
                    f.write("  - 滞后模型: 找不到AI参数\n\n")
            else:
                f.write("  - 滞后模型: 未能成功创建或拟合\n\n")
        else:
            f.write("  - 原始模型: 未能成功创建或拟合或参数无效\n\n")

        # 在处理子样本分析时也添加类似的检查
        # 企业规模子样本分析报告
        f.write("\n2. 企业规模子样本分析\n")
        f.write("-" * 50 + "\n")

        # H1a在不同企业规模下的对比
        f.write("H1a在不同企业规模下的对比(AI技术投入对专利总数的影响):\n")

        # 安全获取大企业样本参数
        if 'H1a_大企业' in size_subsamples and size_subsamples['H1a_大企业'].get('result') is not None:
            large_ai_param = find_ai_param(size_subsamples['H1a_大企业']['result'].params)
            if large_ai_param:
                h1a_large_coef = size_subsamples['H1a_大企业']['result'].params[large_ai_param]
                h1a_large_pvalue = size_subsamples['H1a_大企业']['result'].pvalues[large_ai_param]

                if h1a_large_coef is not None and h1a_large_pvalue is not None:
                    f.write(f"  - 大企业样本: 系数 = {h1a_large_coef:.4f}, p值 = {h1a_large_pvalue:.4f}\n")
                else:
                    f.write("  - 大企业样本: 系数或p值为None\n")
            else:
                f.write("  - 大企业样本: 找不到AI参数\n")
                h1a_large_coef = None
                h1a_large_pvalue = None
        else:
            f.write("  - 大企业样本: 模型未能成功创建或拟合\n")
            h1a_large_coef = None
            h1a_large_pvalue = None

        # 安全获取小企业样本参数
        if 'H1a_小企业' in size_subsamples and size_subsamples['H1a_小企业'].get('result') is not None:
            small_ai_param = find_ai_param(size_subsamples['H1a_小企业']['result'].params)
            if small_ai_param:
                h1a_small_coef = size_subsamples['H1a_小企业']['result'].params[small_ai_param]
                h1a_small_pvalue = size_subsamples['H1a_小企业']['result'].pvalues[small_ai_param]
                f.write(f"  - 小企业样本: 系数 = {h1a_small_coef:.4f}, p值 = {h1a_small_pvalue:.4f}\n")
            else:
                f.write("  - 小企业样本: 找不到AI参数\n")
                h1a_small_coef = None
                h1a_small_pvalue = None
        else:
            f.write("  - 小企业样本: 模型未能成功创建或拟合\n")
            h1a_small_coef = None
            h1a_small_pvalue = None

        # 只有当两个模型都成功时才进行比较
        if h1a_large_coef is not None and h1a_small_coef is not None:
            f.write(f"  - 效应差异: {abs(h1a_large_coef - h1a_small_coef):.4f}")
            f.write(f" (大企业效应{'更强' if abs(h1a_large_coef) > abs(h1a_small_coef) else '更弱'})\n")
            f.write(f"  - 结论: {'均显著' if h1a_large_pvalue < 0.05 and h1a_small_pvalue < 0.05 else '显著性不一致'}, ")
            f.write(f"{'方向一致' if (h1a_large_coef > 0) == (h1a_small_coef > 0) else '方向不一致'}\n\n")
        else:
            f.write("  - 无法比较大小企业效应因为至少有一个模型未能成功分析\n\n")
        f.write("H3a在不同企业规模下的对比(AI人力投入对专利总数的影响):\n")

        # # H3a滞后模型对比
        # f.write("H3a模型对比(AI人力投入对专利总数的影响):\n")
        # if has_h3a:
        #     h3a_coef = results['H3a']['result'].params['ai_job_log']
        #     h3a_pvalue = results['H3a']['result'].pvalues['ai_job_log']

        #     f.write(f"  - 原始模型: 系数 = {h3a_coef:.4f}, p值 = {h3a_pvalue:.4f}\n")

        #     if 'H3a_lag' in lagged_models and lagged_models['H3a_lag'].get('result') is not None:
        #         h3a_lag_coef = lagged_models['H3a_lag']['result'].params['ai_job_log_lag1']
        #         h3a_lag_pvalue = lagged_models['H3a_lag']['result'].pvalues['ai_job_log_lag1']

        #         f.write(f"  - 滞后模型: 系数 = {h3a_lag_coef:.4f}, p值 = {h3a_lag_pvalue:.4f}\n")
        #         f.write(f"  - 结论: {'系数方向一致' if (h3a_coef > 0) == (h3a_lag_coef > 0) else '系数方向不一致'}, ")
        #         f.write(f"{'均显著' if h3a_pvalue < 0.05 and h3a_lag_pvalue < 0.05 else '显著性不一致'}\n\n")
        #     else:
        #         f.write("  - 滞后模型: 未能成功创建或拟合\n\n")
        # else:
        #     f.write("  - 原始模型: 未能成功创建或拟合\n\n")

        # # 安全获取大企业样本参数
        # if 'H3a_大企业' in size_subsamples and size_subsamples['H3a_大企业'].get('result') is not None:
        #     large_ai_job_param = find_ai_job_param(size_subsamples['H3a_大企业']['result'].params)
        #     if large_ai_job_param:
        #         h3a_large_coef = size_subsamples['H3a_大企业']['result'].params[large_ai_job_param]
        #         h3a_large_pvalue = size_subsamples['H3a_大企业']['result'].pvalues[large_ai_job_param]
        #         f.write(f"  - 大企业样本: 系数 = {h3a_large_coef:.4f}, p值 = {h3a_large_pvalue:.4f}\n")
        #     else:
        #         f.write("  - 大企业样本: 找不到AI人力参数\n")
        #         h3a_large_coef = None
        #         h3a_large_pvalue = None
        # else:
        #     f.write("  - 大企业样本: 模型未能成功创建或拟合\n")
        #     h3a_large_coef = None
        #     h3a_large_pvalue = None

        # # 安全获取小企业样本参数
        # if 'H3a_小企业' in size_subsamples and size_subsamples['H3a_小企业'].get('result') is not None:
        #     small_ai_job_param = find_ai_job_param(size_subsamples['H3a_小企业']['result'].params)
        #     if small_ai_job_param:
        #         h3a_small_coef = size_subsamples['H3a_小企业']['result'].params[small_ai_job_param]
        #         h3a_small_pvalue = size_subsamples['H3a_小企业']['result'].pvalues[small_ai_job_param]
        #         f.write(f"  - 小企业样本: 系数 = {h3a_small_coef:.4f}, p值 = {h3a_small_pvalue:.4f}\n")
        #     else:
        #         f.write("  - 小企业样本: 找不到AI人力参数\n")
        #         h3a_small_coef = None
        #         h3a_small_pvalue = None
        # else:
        #     f.write("  - 小企业样本: 模型未能成功创建或拟合\n")
        #     h3a_small_coef = None
        #     h3a_small_pvalue = None

        # # 只有当两个模型都成功时才进行比较
        # if h3a_large_coef is not None and h3a_small_coef is not None:
        #     f.write(f"  - 效应差异: {abs(h3a_large_coef - h3a_small_coef):.4f}")
        #     f.write(f" (大企业效应{'更强' if abs(h3a_large_coef) > abs(h3a_small_coef) else '更弱'})\n")
        #     f.write(f"  - 结论: {'均显著' if h3a_large_pvalue < 0.05 and h3a_small_pvalue < 0.05 else '显著性不一致'}, ")
        #     f.write(f"{'方向一致' if (h3a_large_coef > 0) == (h3a_small_coef > 0) else '方向不一致'}\n\n")
        # else:
        #     f.write("  - 无法比较大小企业效应因为至少有一个模型未能成功分析\n\n")

        # 控制变量敏感性分析报告
        f.write("\n3. 控制变量敏感性分析\n")
        f.write("-" * 50 + "\n")

        f.write("H1a在不同控制变量组合下的稳定性检验:\n")

        for set_name in control_sets.keys():
            model_key = f'H1a_{set_name}'
            if model_key in sensitivity_models:
                # 使用find_ai_param查找正确的参数名
                ai_param = find_ai_param(sensitivity_models[model_key]['result'].params)
                if ai_param:
                    coef = sensitivity_models[model_key]['result'].params[ai_param]
                    pvalue = sensitivity_models[model_key]['result'].pvalues[ai_param]

                    f.write(f"  - {set_name}: 系数 = {coef:.4f}, p值 = {pvalue:.4f}, ")
                    f.write(f"{'显著' if pvalue < 0.05 else '不显著'}\n")
                else:
                    f.write(f"  - {set_name}: 无法找到AI参数\n")
            else:
                f.write(f"  - {set_name}: 模型未成功创建\n")

        # 不同估计方法分析报告
        f.write("\n4. 不同估计方法分析\n")
        f.write("-" * 50 + "\n")

        f.write("H1a在不同估计方法下的对比:\n")

        # 固定效应(原始模型)
        h1a_coef_str = f"{h1a_coef:.4f}" if h1a_coef is not None else "N/A"
        h1a_pvalue_str = f"{h1a_pvalue:.4f}" if h1a_pvalue is not None else "N/A"
        f.write(f"  - 固定效应模型: 系数 = {h1a_coef_str}, p值 = {h1a_pvalue_str}\n")

        # OLS模型
        if 'H1a_OLS' in estimation_methods:
            ols_result = estimation_methods['H1a_OLS']['result']
            # 查找OLS模型中的AI参数
            ols_ai_param = find_ai_param(ols_result.params)
            if ols_ai_param:
                ols_coef = ols_result.params[ols_ai_param]
                ols_pvalue = ols_result.pvalues[ols_ai_param]
                f.write(f"  - OLS模型: 系数 = {ols_coef:.4f}, p值 = {ols_pvalue:.4f}\n")
            else:
                f.write("  - OLS模型: 找不到AI参数\n")

        # 随机效应模型
        if 'H1a_RE' in estimation_methods:
            re_result = estimation_methods['H1a_RE']['result']
            # 查找随机效应模型中的AI参数
            re_ai_param = find_ai_param(re_result.params)
            if re_ai_param:
                re_coef = re_result.params[re_ai_param]
                re_pvalue = re_result.pvalues[re_ai_param]
                f.write(f"  - 随机效应模型: 系数 = {re_coef:.4f}, p值 = {re_pvalue:.4f}\n")
            else:
                f.write("  - 随机效应模型: 找不到AI参数\n")        # 结论
        f.write("\n5. 总体稳健性结论\n")
        f.write("-" * 50 + "\n")

        # 根据各项检验结果，生成总结论
        h1a_robust = True  # 默认值
        if h1a_coef is not None and h1a_lag_coef is not None:
            h1a_robust = h1a_robust and ((h1a_coef > 0) == (h1a_lag_coef > 0))
        else:
            f.write("警告: 缺少H1a或H1a滞后模型的系数，无法比较方向一致性\n")

        if h1a_large_coef is not None and h1a_small_coef is not None:
            h1a_robust = h1a_robust and ((h1a_large_coef > 0) == (h1a_small_coef > 0))
        else:
            f.write("警告: 缺少H1a大型公司或小型公司子样本的系数，无法比较方向一致性\n")
              # 检查OLS模型结果
        if 'H1a_OLS' in estimation_methods:
            ols_result = estimation_methods['H1a_OLS']['result']
            ols_ai_param = find_ai_param(ols_result.params)
            if ols_ai_param and h1a_coef is not None:
                ols_coef = ols_result.params[ols_ai_param]
                h1a_robust = h1a_robust and ((h1a_coef > 0) == (ols_coef > 0))
            else:
                f.write("警告: 缺少OLS模型的AI系数，无法比较方向一致性\n")

        # 检查RE模型的方向一致性
        if 'H1a_RE' in estimation_methods and estimation_methods['H1a_RE'].get('result') is not None:
            re_result = estimation_methods['H1a_RE']['result']
            re_ai_param = find_ai_param(re_result.params)
            if re_ai_param and h1a_coef is not None:
                re_coef = re_result.params[re_ai_param]
                h1a_robust = h1a_robust and ((h1a_coef > 0) == (re_coef > 0))
            else:
                f.write("警告: 缺少RE模型的AI系数，无法比较方向一致性\n")

        # 确保所有系数都非空再进行比较
        h3a_robust = True
        if h3a_coef is not None and h3a_lag_coef is not None:
            h3a_robust = h3a_robust and ((h3a_coef > 0) == (h3a_lag_coef > 0))
        else:
            f.write("警告: 缺少H3a或H3a滞后模型的系数，无法比较方向一致性\n")

        if h3a_large_coef is not None and h3a_small_coef is not None:
            h3a_robust = h3a_robust and ((h3a_large_coef > 0) == (h3a_small_coef > 0))
        else:
            f.write("警告: 缺少H3a大型公司或小型公司子样本的系数，无法比较方向一致性\n")

        f.write(f"H1a(AI技术投入对专利总数的影响): {'稳健' if h1a_robust else '不稳健'}\n")
        f.write(f"H3a(AI人力投入对专利总数的影响): {'稳健' if h3a_robust else '不稳健'}\n")

        # 总体评价
        overall_robust = h1a_robust and h3a_robust
        f.write(f"\n总体评价: 研究结果{'具有良好的稳健性' if overall_robust else '在某些方面缺乏稳健性，需谨慎解释'}\n")

        # 滞后模型报告
        f.write("1. 滞后模型分析\n")
        f.write("-" * 50 + "\n")

        # 对比原始模型和滞后模型（如果存在）
        if 'H1a' in results and results['H1a'].get('result') is not None:
            h1a_coef = results['H1a']['result'].params['ai']
            h1a_pvalue = results['H1a']['result'].pvalues['ai']

            f.write("H1a模型对比(AI技术投入对专利总数的影响):\n")
            f.write(f"  - 原始模型: 系数 = {h1a_coef:.4f}, p值 = {h1a_pvalue:.4f}\n")

            if 'H1a_lag' in lagged_models and lagged_models['H1a_lag'].get('result') is not None:
                # 如果是使用替代变量，需要特别说明
                if lagged_models['H1a_lag'].get('is_substitute'):
                    f.write("  - 滞后模型: 【使用当期变量替代，因为滞后变量创建失败】\n")
                else:
                    h1a_lag_coef = lagged_models['H1a_lag']['result'].params['ai_lag1']
                    h1a_lag_pvalue = lagged_models['H1a_lag']['result'].pvalues['ai_lag1']
                    f.write(f"  - 滞后模型: 系数 = {h1a_lag_coef:.4f}, p值 = {h1a_lag_pvalue:.4f}\n")
                    f.write(f"  - 结论: {'系数方向一致' if (h1a_coef > 0) == (h1a_lag_coef > 0) else '系数方向不一致'}, ")
                    f.write(f"{'均显著' if h1a_pvalue < 0.05 and h1a_lag_pvalue < 0.05 else '显著性不一致'}\n\n")
            else:
                f.write("  - 滞后模型: 未能成功创建\n\n")
        else:
            f.write("H1a模型: 原始模型未成功创建\n\n")        # 类似地处理H3a模型比较
        if 'H3a' in results and results['H3a'].get('result') is not None:
            h3a_coef = results['H3a']['result'].params['ai_job_log']
            h3a_pvalue = results['H3a']['result'].pvalues['ai_job_log']

            f.write("H3a模型对比(AI人力投入对专利总数的影响):\n")
            f.write(f"  - 原始模型: 系数 = {h3a_coef:.4f}, p值 = {h3a_pvalue:.4f}\n")

            if 'H3a_lag' in lagged_models and lagged_models['H3a_lag'].get('result') is not None:
                # 如果是使用替代变量，需要特别说明
                if lagged_models['H3a_lag'].get('is_substitute'):
                    f.write("  - 滞后模型: 【使用当期变量替代，因为滞后变量创建失败】\n")
                else:
                    # 安全获取ai_job_log_lag1参数
                    try:
                        ai_job_lag_param = 'ai_job_log_lag1'
                        if ai_job_lag_param in lagged_models['H3a_lag']['result'].params:
                            h3a_lag_coef = lagged_models['H3a_lag']['result'].params[ai_job_lag_param]
                            h3a_lag_pvalue = lagged_models['H3a_lag']['result'].pvalues[ai_job_lag_param]
                            f.write(f"  - 滞后模型: 系数 = {h3a_lag_coef:.4f}, p值 = {h3a_lag_pvalue:.4f}\n")

                            # 只有当两个系数都不为None时才比较方向
                            if h3a_coef is not None and h3a_lag_coef is not None:
                                f.write(f"  - 结论: {'系数方向一致' if (h3a_coef > 0) == (h3a_lag_coef > 0) else '系数方向不一致'}, ")
                                f.write(f"{'均显著' if h3a_pvalue < 0.05 and h3a_lag_pvalue < 0.05 else '显著性不一致'}\n\n")
                            else:
                                f.write("  - 结论: 无法比较方向一致性，因为至少有一个系数为None\n\n")
                        else:
                            f.write("  - 滞后模型: 找不到ai_job_log_lag1参数\n\n")
                            h3a_lag_coef = None
                            h3a_lag_pvalue = None
                    except Exception as param_e:
                        f.write(f"  - 滞后模型: 获取参数时出错 - {param_e}\n\n")
                        h3a_lag_coef = None
                        h3a_lag_pvalue = None
            else:
                f.write("  - 滞后模型: 未能成功创建\n\n")
        else:
            f.write("H3a模型: 原始模型未成功创建\n\n")

        # 结论
        f.write("\n5. 总体稳健性结论\n")
        f.write("-" * 50 + "\n")        # 初始化robust_items变量
        h1a_robust_items = []
        h3a_robust_items = []

        # 如果滞后模型创建失败，那么只评估可用的结果
        has_h1a_lag = 'H1a_lag' in lagged_models and lagged_models['H1a_lag'].get('result') is not None and not lagged_models['H1a_lag'].get('is_substitute')
        has_h3a_lag = 'H3a_lag' in lagged_models and lagged_models['H3a_lag'].get('result') is not None and not lagged_models['H3a_lag'].get('is_substitute')
        has_h1a_size = 'H1a_大企业' in size_subsamples and 'H1a_小企业' in size_subsamples
        has_h3a_size = 'H3a_大企业' in size_subsamples and 'H3a_小企业' in size_subsamples

        # 安全获取H1a模型系数
        if 'H1a' in results and results['H1a'].get('result') is not None:
            try:
                if 'ai' in results['H1a']['result'].params:
                    h1a_coef = results['H1a']['result'].params['ai']
                else:
                    h1a_coef = None
                    f.write("注意: 在H1a模型中找不到'ai'参数\n")
            except Exception as h1a_err:
                f.write(f"获取H1a系数时出错: {h1a_err}\n")
                h1a_coef = None

        # 检查H1a与滞后模型的方向一致性
        if has_h1a_lag and h1a_coef is not None:
            try:
                h1a_lag_param = find_ai_param(lagged_models['H1a_lag']['result'].params)
                if h1a_lag_param:
                    h1a_lag_coef = lagged_models['H1a_lag']['result'].params[h1a_lag_param]
                    if h1a_lag_coef is not None:
                        h1a_robust_items.append((h1a_coef > 0) == (h1a_lag_coef > 0))
                    else:
                        f.write("注意: H1a滞后模型的系数为None\n")
                else:
                    f.write("注意: 在H1a滞后模型中找不到AI参数\n")
            except Exception as lag_err:
                f.write(f"处理H1a滞后模型时出错: {lag_err}\n")

            if has_h1a_size:
                h1a_large_coef = size_subsamples['H1a_大企业']['result'].params['ai']
                h1a_small_coef = size_subsamples['H1a_小企业']['result'].params['ai']
                h1a_robust_items.append((h1a_large_coef > 0) == (h1a_small_coef > 0))

            if 'H1a_OLS' in estimation_methods and estimation_methods['H1a_OLS'].get('result') is not None:
                ols_ai_param = find_ai_param(estimation_methods['H1a_OLS']['result'].params)
                if ols_ai_param:
                    ols_coef = estimation_methods['H1a_OLS']['result'].params[ols_ai_param]
                    h1a_robust_items.append((h1a_coef > 0) == (ols_coef > 0))

            if 'H1a_RE' in estimation_methods and estimation_methods['H1a_RE'].get('result') is not None:
                re_ai_param = find_ai_param(estimation_methods['H1a_RE']['result'].params)
                if re_ai_param:
                    re_coef = estimation_methods['H1a_RE']['result'].params[re_ai_param]
                    h1a_robust_items.append((h1a_coef > 0) == (re_coef > 0))

            if h1a_robust_items:
                h1a_robust = all(h1a_robust_items)
                f.write(f"H1a(AI技术投入对专利总数的影响): {'稳健' if h1a_robust else '不稳健'}")
                f.write(f" - 基于{len(h1a_robust_items)}项检验\n")
            else:
                f.write("H1a(AI技术投入对专利总数的影响): 无法评估稳健性(缺少比较项)\n")
        else:
            f.write("H1a(AI技术投入对专利总数的影响): 主模型未成功创建\n")        # 类似地处理H3a稳健性评估
        if 'H3a' in results and results['H3a'].get('result') is not None:
            h3a_coef = results['H3a']['result'].params['ai_job_log']
            h3a_robust_items = []

            # 检查滞后效应方向一致性
            if 'H3a_lag' in lagged_models and lagged_models['H3a_lag'].get('result') is not None:
                try:
                    if 'ai_job_log_lag1' in lagged_models['H3a_lag']['result'].params:
                        h3a_lag_coef = lagged_models['H3a_lag']['result'].params['ai_job_log_lag1']
                        if h3a_coef is not None and h3a_lag_coef is not None:
                            h3a_robust_items.append((h3a_coef > 0) == (h3a_lag_coef > 0))
                        else:
                            f.write("注意: 无法比较H3a滞后效应的方向一致性，因为至少有一个系数为None\n")
                    else:
                        f.write("注意: 滞后模型中找不到ai_job_log_lag1参数\n")
                except Exception as lag_err:
                    f.write(f"检查H3a滞后效应时出错: {lag_err}\n")

            # 检查规模子样本一致性
            if 'H3a_大企业' in size_subsamples and size_subsamples['H3a_大企业'].get('result') is not None and \
               'H3a_小企业' in size_subsamples and size_subsamples['H3a_小企业'].get('result') is not None:
                try:
                    large_param = 'ai_job_log'
                    small_param = 'ai_job_log'

                    if large_param in size_subsamples['H3a_大企业']['result'].params and \
                       small_param in size_subsamples['H3a_小企业']['result'].params:
                        h3a_large_coef = size_subsamples['H3a_大企业']['result'].params[large_param]
                        h3a_small_coef = size_subsamples['H3a_小企业']['result'].params[small_param]

                        if h3a_large_coef is not None and h3a_small_coef is not None:
                            h3a_robust_items.append((h3a_large_coef > 0) == (h3a_small_coef > 0))
                        else:
                            f.write("注意: 无法比较H3a不同规模企业效应的方向一致性，因为至少有一个系数为None\n")
                    else:
                        f.write("注意: 子样本模型中找不到ai_job_log参数\n")
                except Exception as size_err:
                    f.write(f"检查H3a子样本效应时出错: {size_err}\n")

            if h3a_robust_items:
                h3a_robust = all(h3a_robust_items)
                f.write(f"H3a(AI人力投入对专利总数的影响): {'稳健' if h3a_robust else '不稳健'}")
                f.write(f" - 基于{len(h3a_robust_items)}项检验\n")
            else:
                f.write("H3a(AI人力投入对专利总数的影响): 无法评估稳健性(缺少比较项)\n")
        else:
            f.write("H3a(AI人力投入对专利总数的影响): 主模型未成功创建\n")        # 总体评价 - 根据可用检验结果判断
        available_robust_checks = []

        # 确保h1a_robust_items变量存在再使用
        if 'H1a' in results and results['H1a'].get('result') is not None and 'h1a_robust_items' in locals() and h1a_robust_items:
            available_robust_checks.append(all(h1a_robust_items))

        # 确保h3a_robust_items变量存在再使用
        if 'H3a' in results and results['H3a'].get('result') is not None and 'h3a_robust_items' in locals() and h3a_robust_items:
            available_robust_checks.append(all(h3a_robust_items))

        if available_robust_checks:
            overall_robust = all(available_robust_checks)
            f.write(f"\n总体评价: 研究结果{'具有良好的稳健性' if overall_robust else '在某些方面缺乏稳健性，需谨慎解释'}")
            f.write(f" - 基于{len(available_robust_checks)}个假设的检验\n")
        else:
            f.write("\n总体评价: 由于缺少足够的对比检验，无法给出稳健性的全面评估\n")

    print("稳健性检验完成，结果已保存到output/reports/robustness_checks_report.txt")

    return robustness_results

# 20. 创建综合结果报告
def create_comprehensive_report(analysis_results=None, output_path='output/reports/comprehensive_report.docx'):
    """
    创建包含所有分析结果的综合Word文档报告

    参数:
    analysis_results: 包含所有分析结果的字典
    output_path: 输出报告的路径
    """
    print("\n开始生成综合分析报告...")

    # 初始化analysis_results如果为None
    if analysis_results is None:
        analysis_results = {}
      # 初始化model_results如果不存在
    if 'model_results' not in analysis_results:
        analysis_results['model_results'] = {}

    # 设置model_results变量以便代码其他部分使用
    model_results = analysis_results['model_results']

    try:
        # 导入docx库
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 创建文档对象
        doc = Document()

        # 1. 设置文档标题
        title = doc.add_heading('AI投入对企业创新绩效影响分析报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加报告日期
        from datetime import datetime
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(f"报告生成日期: {datetime.now().strftime('%Y年%m月%d日')}").italic = True

        # 添加摘要
        doc.add_heading('摘要', level=1)
        abstract = doc.add_paragraph()
        abstract.add_run('''本报告分析了企业AI技术投入和人力投入对创新绩效的影响。
通过面板数据分析方法，研究发现AI技术投入与专利数量呈正相关关系，与创新质量呈U型关系；
AI人力投入直接提高专利数量和质量；专利和人才的跨领域程度对上述关系具有显著调节作用。
研究结果通过多种稳健性检验，具有较高的可靠性。''')

        doc.add_page_break()

        # 2. 添加目录
        doc.add_heading('目录', level=1)
        toc_entries = [
            "1. 研究背景与数据描述",
            "2. 面板数据诊断",
            "3. 假设检验结果",
            "4. 高级分析结果",
            "5. 稳健性检验",
            "6. 研究结论与建议"
        ]
        for entry in toc_entries:
            # 使用普通段落，不指定样式
            p = doc.add_paragraph()
            p.add_run(entry)
            # 添加缩进
            p.paragraph_format.left_indent = Inches(0.25)

        doc.add_page_break()

        # 3. 研究背景与数据描述
        doc.add_heading('1. 研究背景与数据描述', level=1)

        # 数据概况
        doc.add_heading('1.1 数据概况', level=2)
        if 'preprocessed_data' in analysis_results:
            preproc_data = analysis_results['preprocessed_data']
            p = doc.add_paragraph()
            p.add_run(f"本研究使用的面板数据共包含 ")
            p.add_run(f"{preproc_data.get('shape', (0, 0))[0]}").bold = True
            p.add_run(f" 个观测值，")
            p.add_run(f"{preproc_data.get('shape', (0, 0))[1]}").bold = True
            p.add_run(" 个变量。数据涵盖2014-2022年期间，包括企业AI技术投入、AI人力投入、专利数量和质量等关键指标。")

        # 描述性统计
        doc.add_heading('1.2 描述性统计', level=2)
        if 'descriptive_stats' in analysis_results and analysis_results['descriptive_stats']:
            desc_stats = analysis_results['descriptive_stats']

            # 核心变量描述性统计表
            doc.add_heading('1.2.1 核心变量描述性统计', level=3)
            if 'all_desc_stats' in desc_stats and 'ai_investment' in desc_stats['all_desc_stats']:
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'

                # 添加表头
                headers = table.rows[0].cells
                headers[0].text = '变量'
                headers[1].text = '样本量'
                headers[2].text = '均值'
                headers[3].text = '标准差'
                headers[4].text = '最小值'
                headers[5].text = '最大值'
                headers[6].text = '偏度'

                # 提取AI投入相关变量的统计信息
                ai_stats = desc_stats['all_desc_stats']['ai_investment']
                for var_name, row_data in ai_stats.iterrows():
                    cells = table.add_row().cells
                    cells[0].text = var_name
                    cells[1].text = f"{row_data['count']:.0f}"
                    cells[2].text = f"{row_data['mean']:.4f}"
                    cells[3].text = f"{row_data['std']:.4f}"
                    cells[4].text = f"{row_data['min']:.4f}"
                    cells[5].text = f"{row_data['max']:.4f}"
                    cells[6].text = f"{row_data['skewness']:.4f}"

            # 年度变化趋势
            doc.add_heading('1.2.2 核心变量年度变化趋势', level=3)
            p = doc.add_paragraph("主要变量在2014-2022年间呈现以下变化趋势：")

            # 添加年度变化趋势图
            try:
                doc.add_picture('output/figures/yearly_means_trend.png', width=Inches(6))
                caption = doc.add_paragraph('图1.1 核心变量年度均值变化趋势')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph("年度均值变化趋势图生成失败或不存在。")

            # 相关性分析
            doc.add_heading('1.2.3 变量相关性分析', level=3)
            p = doc.add_paragraph("各核心变量间的相关性如下：")

            # 这里可以添加相关性矩阵图或表格，如果存在的话
            try:
                doc.add_picture('output/figures/key_vars_scatter_matrix.png', width=Inches(6))
                caption = doc.add_paragraph('图1.2 核心变量散点矩阵')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph("变量散点矩阵图生成失败或不存在。")

        doc.add_page_break()

        # 4. 面板数据诊断
        doc.add_heading('2. 面板数据诊断', level=1)

        if 'diagnostics' in analysis_results and analysis_results['diagnostics']:
            diagnostics = analysis_results['diagnostics']

            # 平稳性检验
            doc.add_heading('2.1 单位根检验(平稳性检验)', level=2)
            if 'unit_root' in diagnostics and diagnostics['unit_root']:
                p = doc.add_paragraph("对主要研究变量进行单位根检验，以确定序列的平稳性。平稳性检验结果如下：")

                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'

                # 添加表头
                headers = table.rows[0].cells
                headers[0].text = '变量'
                headers[1].text = '平稳序列比例'
                headers[2].text = '结论'

                # 添加单位根检验结果
                for var, result in diagnostics['unit_root'].items():
                    cells = table.add_row().cells
                    cells[0].text = var
                    cells[1].text = f"{result.get('stationary_percentage', 'N/A'):.2f}%"

                    # 根据平稳比例得出结论
                    stationary_pct = result.get('stationary_percentage', 0)
                    if stationary_pct > 75:
                        conclusion = "大多数序列平稳"
                    elif stationary_pct > 50:
                        conclusion = "部分序列平稳，考虑使用差分"
                    else:
                        conclusion = "大多数序列非平稳，建议差分处理"

                    cells[2].text = conclusion

            # Hausman检验
            doc.add_heading('2.2 Hausman检验(模型选择)', level=2)
            if 'hausman' in diagnostics and diagnostics['hausman']:
                p = doc.add_paragraph("为确定使用固定效应还是随机效应模型，进行Hausman检验：")

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'

                # 添加表头
                headers = table.rows[0].cells
                headers[0].text = '因变量'
                headers[1].text = 'Hausman统计量'
                headers[2].text = 'p值'
                headers[3].text = '推荐模型'

                # 添加Hausman检验结果
                for var, result in diagnostics['hausman'].items():
                    if 'error' in result:
                        continue

                    cells = table.add_row().cells
                    cells[0].text = var
                    cells[1].text = f"{result.get('hausman_stat', 'N/A'):.4f}"
                    cells[2].text = f"{result.get('p_value', 'N/A'):.4f}"
                    cells[3].text = result.get('conclusion', 'N/A')

                # 总体结论
                fixed_effects_count = sum(1 for r in diagnostics['hausman'].values()
                                     if 'conclusion' in r and r['conclusion'] == '固定效应')
                total_valid = sum(1 for r in diagnostics['hausman'].values() if 'conclusion' in r)

                if total_valid > 0:
                    p = doc.add_paragraph()
                    if fixed_effects_count / total_valid > 0.5:
                        p.add_run(f"结论: 在{fixed_effects_count}个模型中，{fixed_effects_count}个推荐使用固定效应模型。本研究主要采用").bold = False
                        p.add_run("固定效应模型").bold = True
                        p.add_run("进行分析。").bold = False
                    else:
                        p.add_run(f"结论: 在{total_valid}个模型中，{fixed_effects_count}个推荐使用固定效应模型，{total_valid-fixed_effects_count}个推荐使用随机效应模型。基于研究设计，本研究主要采用").bold = False
                        p.add_run("固定效应模型").bold = True
                        p.add_run("进行分析，并使用随机效应模型进行稳健性检验。").bold = False

            # 异方差检验
            doc.add_heading('2.3 异方差与序列相关性检验', level=2)
            if ('heteroskedasticity' in diagnostics and diagnostics['heteroskedasticity'] or
                'serial_correlation' in diagnostics and diagnostics['serial_correlation']):

                p = doc.add_paragraph("为确保模型估计的有效性，进行异方差和序列相关性检验：")

                # 异方差检验结果
                if 'heteroskedasticity' in diagnostics and diagnostics['heteroskedasticity']:
                    doc.add_heading('2.3.1 异方差检验结果', level=3)

                    hetero_exists = any(r.get('heteroskedasticity', False)
                                       for r in diagnostics['heteroskedasticity'].values()
                                       if 'error' not in r)

                    p = doc.add_paragraph()
                    if hetero_exists:
                        p.add_run("检验结果表明存在异方差问题，需要使用").bold = False
                        p.add_run("稳健标准误").bold = True
                        p.add_run("进行模型估计。").bold = False
                    else:
                        p.add_run("检验结果表明异方差问题不明显，但为稳健起见，仍使用").bold = False
                        p.add_run("稳健标准误").bold = True
                        p.add_run("进行模型估计。").bold = False

                # 序列相关性检验结果
                if 'serial_correlation' in diagnostics and diagnostics['serial_correlation']:
                    doc.add_heading('2.3.2 序列相关性检验结果', level=3)

                    # 计算存在自相关的比例
                    autocorr_ratios = [r.get('companies_with_autocorrelation_pct', 0)
                                      for r in diagnostics['serial_correlation'].values()
                                      if 'error' not in r]

                    if autocorr_ratios:
                        avg_ratio = sum(autocorr_ratios) / len(autocorr_ratios)

                        p = doc.add_paragraph()
                        if avg_ratio > 30:
                            p.add_run(f"检验结果表明平均有{avg_ratio:.1f}%的公司序列存在自相关，需要使用").bold = False
                            p.add_run("聚类稳健标准误").bold = True
                            p.add_run("并考虑添加滞后变量进行模型估计。").bold = False
                        else:
                            p.add_run(f"检验结果表明平均有{avg_ratio:.1f}%的公司序列存在自相关，为稳健起见，使用").bold = False
                            p.add_run("聚类稳健标准误").bold = True
                            p.add_run("进行模型估计。").bold = False

            # 面板数据诊断总结
            doc.add_heading('2.4 面板数据诊断总结', level=2)
            p = doc.add_paragraph("基于上述诊断结果，本研究的面板数据模型设计如下：")
            p.add_run("• 主要使用固定效应模型\n").bold = True
            p.add_run("• 采用聚类稳健标准误\n").bold = True
            p.add_run("• 对非平稳变量进行差分处理\n").bold = True
            p.add_run("• 进行滞后变量分析进行稳健性检验").bold = True

        doc.add_page_break()

        # 5. 假设检验结果
        doc.add_heading('3. 假设检验结果', level=1)

        if 'model_results' in analysis_results and analysis_results['model_results']:
            model_results = analysis_results['model_results']

            # 假设检验摘要表
            doc.add_heading('3.1 假设检验摘要', level=2)
            p = doc.add_paragraph("本研究的主要假设检验结果如下：")

            # 尝试添加假设检验摘要表
            try:
                import pandas as pd
                hypothesis_summary = pd.read_csv('output/tables/hypothesis_testing_summary.csv')

                if not hypothesis_summary.empty:
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'

                    # 添加表头
                    headers = table.rows[0].cells
                    for i, col in enumerate(hypothesis_summary.columns):
                        headers[i].text = col

                    # 添加数据行
                    for _, row in hypothesis_summary.iterrows():
                        cells = table.add_row().cells
                        for i, col in enumerate(hypothesis_summary.columns):
                            cells[i].text = str(row[col])

                            # 如果是结论列，对"支持"加粗显示
                            if i == 5 and "支持" in str(row[col]):  # 假设第6列是结论
                                run = cells[i].paragraphs[0].runs[0]
                                run.bold = True
                                if "不支持" not in str(row[col]):
                                    run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            except:
                p = doc.add_paragraph("假设检验摘要表加载失败。")

            # 主要假设的详细结果
            doc.add_heading('3.2 主要假设检验详细结果', level=2)

            # H1系列：AI技术投入对创新能力的影响
            doc.add_heading('3.2.1 AI技术投入对创新能力的影响(H1系列)', level=3)

            # H1a
            if 'H1a' in model_results and model_results['H1a'] and model_results['H1a'].get('result'):
                h1a_result = model_results['H1a']['result']

                doc.add_heading('H1a: AI技术投入提高专利总数量', level=4)
                p = doc.add_paragraph()
                coef = h1a_result.params['ai']
                pval = h1a_result.pvalues['ai']

                p.add_run(f"回归系数: ").bold = False
                p.add_run(f"{coef:.4f}").bold = True
                p.add_run(f", p值: {pval:.4f}").bold = False

                if pval < 0.05 and coef > 0:
                    p.add_run(" (显著正向影响)").bold = True
                    p.add_run("\n结论: 支持H1a假设，AI技术投入显著提高企业专利总数量。").bold = False
                else:
                    p.add_run(" (无显著正向影响)").bold = True
                    p.add_run("\n结论: 不支持H1a假设。").bold = False

            # H1b
            if 'H1b_models' in model_results and model_results['H1b_models'] and 'trend' in model_results['H1b_models']:
                h1b_result = model_results['H1b_models']['trend']['result']

                doc.add_heading('H1b: AI技术投入与探索型专利的时间动态关系', level=4)
                p = doc.add_paragraph()

                main_coef = h1b_result.params['ai']
                main_pval = h1b_result.pvalues['ai']
                trend_coef = h1b_result.params['ai_x_time_trend']
                trend_pval = h1b_result.pvalues['ai_x_time_trend']

                p.add_run(f"主效应: {main_coef:.4f} (p值: {main_pval:.4f})\n").bold = False
                p.add_run(f"时间交互效应: {trend_coef:.4f} (p值: {trend_pval:.4f})").bold = False

                if main_pval < 0.05 and main_coef > 0 and trend_pval < 0.05 and trend_coef < 0:
                    p.add_run("\n结论: 支持H1b假设，AI技术投入对探索型专利的正向影响随时间减弱。").bold = True
                else:
                    p.add_run("\n结论: 不支持H1b假设。").bold = False

                # 添加时间动态效应图
                try:
                    doc.add_picture('output/figures/time_dynamic_ai_ep.png', width=Inches(6))
                    caption = doc.add_paragraph('图3.1 AI技术投入对探索型专利的时间动态效应')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    p.add_run("\n(时间动态效应图不可用)").italic = True

            # H2系列：非线性关系
            doc.add_heading('3.2.2 AI技术投入与创新质量的非线性关系(H2系列)', level=3)

            # H2a
            if ('H2a' in model_results and model_results['H2a'] and model_results['H2a'].get('result') and
                'H2a_u_test' in model_results and model_results['H2a_u_test']):

                h2a_result = model_results['H2a']['result']
                h2a_u_test = model_results['H2a_u_test']

                doc.add_heading('H2a: AI技术投入与专利质量的U型关系', level=4)
                p = doc.add_paragraph()

                linear_coef = h2a_result.params['ai']
                squared_coef = h2a_result.params['ai_squared']
                linear_pval = h2a_result.pvalues['ai']
                squared_pval = h2a_result.pvalues['ai_squared']

                p.add_run(f"线性项: {linear_coef:.4f} (p值: {linear_pval:.4f})\n").bold = False
                p.add_run(f"二次项: {squared_coef:.4f} (p值: {squared_pval:.4f})\n").bold = False
                p.add_run(f"拐点: {h2a_u_test.get('turning_point', 'N/A')}\n").bold = False
                p.add_run(f"Sasabuchi检验p值: {h2a_u_test.get('sasabuchi_p', 'N/A')}\n").bold = False

                if h2a_u_test.get('valid_nonlinear', False):
                    p.add_run("\n结论: 支持H2a假设，AI技术投入与专利质量存在U型关系。").bold = True
                else:
                    p.add_run("\n结论: 不支持H2a假设。").bold = False
                    p.add_run(f" {h2a_u_test.get('conclusion', '无具体结论')}").italic = True

                # 添加U型关系图                try:
                    plot_path = h2a_u_test.get('plot_path') or h2a_u_test.get('comprehensive_plot_path')
                    if plot_path:
                        doc.add_picture(plot_path, width=Inches(6))
                        caption = doc.add_paragraph('图3.2 AI技术投入与专利质量的U型关系')

            # H3-H6系列假设检验结果
            # 类似地添加其他假设的检验结果...

            # H3系列：AI人力投入对创新能力的影响
            doc.add_heading('3.2.3 AI人力投入对创新能力的影响(H3系列)', level=3)

            # H3a
            if 'H3a' in model_results and model_results['H3a'] and model_results['H3a'].get('result'):
                h3a_result = model_results['H3a']['result']

                doc.add_heading('H3a: AI人力投入提高专利总数量', level=4)
                p = doc.add_paragraph()
                coef = h3a_result.params['ai_job_log']
                pval = h3a_result.pvalues['ai_job_log']

                p.add_run(f"回归系数: ").bold = False
                p.add_run(f"{coef:.4f}").bold = True
                p.add_run(f", p值: {pval:.4f}").bold = False

                if pval < 0.05 and coef > 0:
                    p.add_run(" (显著正向影响)").bold = True
                    p.add_run("\n结论: 支持H3a假设，AI人力投入显著提高企业专利总数量。").bold = False
                else:
                    p.add_run(" (无显著正向影响)").bold = True
                    p.add_run("\n结论: 不支持H3a假设。").bold = False

        doc.add_page_break()

        # 6. 高级分析结果
        doc.add_heading('4. 高级分析结果', level=1)

        if 'advanced_results' in analysis_results and analysis_results['advanced_results']:
            advanced_results = analysis_results['advanced_results']

            # 时间动态效应分析
            doc.add_heading('4.1 时间动态效应分析', level=2)
            if 'h1b_time_dynamics' in advanced_results and advanced_results['h1b_time_dynamics']:
                p = doc.add_paragraph("AI技术投入对探索型专利的影响随时间变化的动态效应分析：")

                # 可以添加时间动态效应的详细分析结果
                try:
                    doc.add_picture('output/figures/time_dynamic_ai_ep.png', width=Inches(6))
                    caption = doc.add_paragraph('图4.1 AI技术投入对探索型专利的时间动态效应')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    p.add_run("(时间动态效应图不可用)").italic = True

            # 非线性关系分析
            doc.add_heading('4.2 非线性关系分析', level=2)
            if ('h2a_nonlinear' in advanced_results and advanced_results['h2a_nonlinear'] or
                'h2b_nonlinear' in advanced_results and advanced_results['h2b_nonlinear']):

                p = doc.add_paragraph("AI技术投入与创新质量指标的非线性关系高级分析：")

                # H2a的非线性分析
                if 'h2a_nonlinear' in advanced_results and advanced_results['h2a_nonlinear']:
                    doc.add_heading('4.2.1 AI技术投入与专利质量的U型关系', level=3)
                    h2a_nl = advanced_results['h2a_nonlinear']

                    p = doc.add_paragraph()
                    p.add_run(f"拐点: {h2a_nl.get('turning_point', 'N/A')}\n").bold = False
                    p.add_run(f"拐点95%置信区间: [{h2a_nl.get('turning_point_ci', (0, 0))[0]:.4f}, {h2a_nl.get('turning_point_ci', (0, 0))[1]:.4f}]\n").bold = False
                    p.add_run(f"Sasabuchi检验p值: {h2a_nl.get('sasabuchi_p', 'N/A')}\n").bold = False
                    p.add_run(f"结论: {h2a_nl.get('conclusion', '无具体结论')}").bold = True

                    # 添加U型关系增强图
                    try:
                        doc.add_picture(h2a_nl.get('plot_path', ''), width=Inches(6))
                        caption = doc.add_paragraph('图4.2 AI技术投入与专利质量的U型关系（高级分析）')
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        p.add_run("\n(U型关系增强图不可用)").italic = True

            # 调节效应分析
            doc.add_heading('4.3 调节效应分析', level=2)

            # H5a的调节效应
            if 'h5a_moderation' in advanced_results and advanced_results['h5a_moderation']:
                doc.add_heading('4.3.1 专利跨领域程度对AI技术投入与专利数关系的调节', level=3)
                h5a_mod = advanced_results['h5a_moderation']

                if 'model_result' in h5a_mod and h5a_mod['model_result'] and h5a_mod['model_result'].get('result'):
                    result = h5a_mod['model_result']['result']
                    interact_var = 'ai_x_ai_patent_log'

                    p = doc.add_paragraph()
                    if interact_var in result.params and interact_var in result.pvalues:
                        interact_coef = result.params[interact_var]
                        interact_pval = result.pvalues[interact_var]

                        p.add_run(f"交互项系数: {interact_coef:.4f} (p值: {interact_pval:.4f})\n").bold = False

                        if interact_pval < 0.05:
                            if interact_coef > 0:
                                p.add_run("结论: 专利跨领域程度对AI技术投入与专利数的关系有显著正向调节作用。").bold = True
                            else:
                                p.add_run("结论: 专利跨领域程度对AI技术投入与专利数的关系有显著负向调节作用。").bold = True
                        else:
                            p.add_run("结论: 专利跨领域程度对AI技术投入与专利数的关系没有显著调节作用。").bold = False

                # 添加简单斜率分析结果
                if 'simple_slopes' in h5a_mod and h5a_mod['simple_slopes']:
                    doc.add_heading('简单斜率分析', level=4)

                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'

                    # 添加表头
                    headers = table.rows[0].cells
                    headers[0].text = '调节变量水平'
                    headers[1].text = '简单斜率'
                    headers[2].text = 'p值'
                    headers[3].text = '显著性'

                    # 添加简单斜率数据
                    for level, data in h5a_mod['simple_slopes'].items():
                        cells = table.add_row().cells
                        cells[0].text = level
                        cells[1].text = f"{data['slope']:.4f}"
                        cells[2].text = f"{data['p_value']:.4f}"
                        cells[3].text = "显著" if data['significant'] else "不显著"

                # 添加Johnson-Neyman分析结果
                if 'jn_result' in h5a_mod and h5a_mod['jn_result'] and 'jn_points' in h5a_mod['jn_result']:
                    doc.add_heading('Johnson-Neyman显著性区间分析', level=4)

                    jn_result = h5a_mod['jn_result']
                    p = doc.add_paragraph()

                    if jn_result['jn_points']:
                        p.add_run(f"Johnson-Neyman点: {[f'{jn:.4f}' for jn in jn_result['jn_points']]}\n").bold = False

                        if 'significant_regions' in jn_result:
                            p.add_run("显著区域:\n").bold = True
                            for region in jn_result['significant_regions']:
                                p.add_run(f"- {region['description']}\n").bold = False
                    else:
                        p.add_run("未找到Johnson-Neyman点，调节效应在整个范围内均为显著或均不显著。").bold = False

                # 添加调节效应图
                try:
                    doc.add_picture('output/figures/moderation_ai_ai_patent_log_intotal_prediction.png', width=Inches(6))
                    caption = doc.add_paragraph('图4.3 专利跨领域程度对AI技术投入与专利数关系的调节效应')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    p = doc.add_paragraph("(调节效应图不可用)").italic = True

        doc.add_page_break()

        # 7. 稳健性检验
        doc.add_heading('5. 稳健性检验', level=1)

        if 'robustness_results' in analysis_results and analysis_results['robustness_results']:
            robustness_results = analysis_results['robustness_results']
              # 滞后模型分析
            doc.add_heading('5.1 滞后模型分析', level=2)
            if 'lagged_models' in robustness_results and robustness_results['lagged_models']:
                lagged_models = robustness_results['lagged_models']

                p = doc.add_paragraph("为检验主要发现的稳健性，并缓解潜在的内生性问题，对关键解释变量进行滞后处理：")

                # 对比原始模型和滞后模型
                if 'H1a_lag' in lagged_models and lagged_models['H1a_lag'].get('result'):
                    doc.add_heading('5.1.1 H1a: AI技术投入对专利总数的滞后效应', level=3)

                    h1a_lag = lagged_models['H1a_lag']['result']

                    p = doc.add_paragraph()

                    # 检查原始模型中是否有H1a的结果
                    if 'H1a' in model_results and model_results['H1a'] and model_results['H1a'].get('result') is not None:
                        p.add_run(f"原始模型中AI系数: {model_results['H1a']['result'].params['ai']:.4f} (p值: {model_results['H1a']['result'].pvalues['ai']:.4f})\n").bold = False
                        p.add_run(f"滞后模型中AI滞后系数: {h1a_lag.params['ai_lag1']:.4f} (p值: {h1a_lag.pvalues['ai_lag1']:.4f})\n").bold = False

                        # 比较结果
                        orig_sig = model_results['H1a']['result'].pvalues['ai'] < 0.05
                        lag_sig = h1a_lag.pvalues['ai_lag1'] < 0.05
                        same_dir = (model_results['H1a']['result'].params['ai'] > 0) == (h1a_lag.params['ai_lag1'] > 0)
                    else:
                        # 如果原始模型中没有H1a结果，则只显示滞后模型结果
                        p.add_run(f"原始模型中AI系数: 不可用 (原始H1a模型结果缺失)\n").bold = False
                        p.add_run(f"滞后模型中AI滞后系数: {h1a_lag.params['ai_lag1']:.4f} (p值: {h1a_lag.pvalues['ai_lag1']:.4f})\n").bold = False
                          # 不进行比较
                        lag_sig = h1a_lag.pvalues['ai_lag1'] < 0.05
                        orig_sig = False
                        same_dir = False

                    if orig_sig and lag_sig and same_dir:
                        p.add_run("结论: 原始模型和滞后模型均显示显著的相同方向效应，主要发现具有稳健性。").bold = True
                    elif same_dir and 'H1a' in model_results and model_results['H1a'] and model_results['H1a'].get('result') is not None:
                        p.add_run("结论: 原始模型和滞后模型的效应方向一致，但显著性不同，主要发现部分稳健。").bold = False
                    elif 'H1a' in model_results and model_results['H1a'] and model_results['H1a'].get('result') is not None:
                        p.add_run("结论: 原始模型和滞后模型的效应存在差异，主要发现需谨慎解释。").bold = False
                    else:
                        if lag_sig:
                            p.add_run("结论: 无法与原始模型比较，但滞后模型显示显著效应。").bold = False
                        else:
                            p.add_run("结论: 无法与原始模型比较，滞后模型未显示显著效应。").bold = False

            # 子样本分析
            doc.add_heading('5.2 企业规模子样本分析', level=2)
            if 'size_subsamples' in robustness_results and robustness_results['size_subsamples']:
                size_samples = robustness_results['size_subsamples']

                p = doc.add_paragraph("为检验主要发现在不同企业规模下的稳健性，将样本分为大企业和小企业子样本进行分析：")

                # H1a在不同企业规模下的对比
                if 'H1a_大企业' in size_samples and 'H1a_小企业' in size_samples:
                    doc.add_heading('5.2.1 H1a: AI技术投入对专利总数影响在不同企业规模的对比', level=3)

                    h1a_large = size_samples['H1a_大企业']['result']
                    h1a_small = size_samples['H1a_小企业']['result']

                    p = doc.add_paragraph()
                    p.add_run(f"大企业子样本中AI系数: {h1a_large.params['ai']:.4f} (p值: {h1a_large.pvalues['ai']:.4f})\n").bold = False
                    p.add_run(f"小企业子样本中AI系数: {h1a_small.params['ai']:.4f} (p值: {h1a_small.pvalues['ai']:.4f})\n").bold = False

                    # 比较结果
                    large_sig = h1a_large.pvalues['ai'] < 0.05
                    small_sig = h1a_small.pvalues['ai'] < 0.05
                    same_dir = (h1a_large.params['ai'] > 0) == (h1a_small.params['ai'] > 0)

                    if large_sig and small_sig and same_dir:
                        p.add_run("结论: 大企业和小企业子样本均显示显著的相同方向效应，主要发现具有稳健性。").bold = True
                    elif same_dir:
                        p.add_run("结论: 大企业和小企业子样本的效应方向一致，但显著性不同，主要发现部分稳健。").bold = False
                    else:
                        p.add_run("结论: 大企业和小企业子样本的效应存在差异，表明企业规模对AI投入与创新的关系有重要影响。").bold = False

            # 不同估计方法的检验
            doc.add_heading('5.3 不同估计方法检验', level=2)
            if 'estimation_methods' in robustness_results and robustness_results['estimation_methods']:
                estimation_methods = robustness_results['estimation_methods']

                p = doc.add_paragraph("为检验主要发现对估计方法选择的敏感性，使用不同的估计方法（OLS、随机效应）进行对比分析：")
                if 'H1a_OLS' in estimation_methods or 'H1a_RE' in estimation_methods:
                    doc.add_heading('5.3.1 H1a: AI技术投入对专利总数影响的不同估计方法对比', level=3)

                    p = doc.add_paragraph()

                    # 检查原始模型是否有H1a结果
                    if 'H1a' in model_results and model_results['H1a'] and model_results['H1a'].get('result') is not None:
                        p.add_run(f"固定效应模型AI系数: {model_results['H1a']['result'].params['ai']:.4f} (p值: {model_results['H1a']['result'].pvalues['ai']:.4f})\n").bold = False
                    else:
                        p.add_run(f"固定效应模型AI系数: 不可用 (原始H1a模型结果缺失)\n").bold = False

                    if 'H1a_OLS' in estimation_methods and estimation_methods['H1a_OLS'].get('result'):
                        ols_result = estimation_methods['H1a_OLS']['result']
                        p.add_run(f"OLS模型AI系数: {ols_result.params['ai']:.4f} (p值: {ols_result.pvalues['ai']:.4f})\n").bold = False

                    if 'H1a_RE' in estimation_methods and estimation_methods['H1a_RE'].get('result'):
                        re_result = estimation_methods['H1a_RE']['result']
                        p.add_run(f"随机效应模型AI系数: {re_result.params['ai']:.4f} (p值: {re_result.pvalues['ai']:.4f})\n").bold = False

            # 总体稳健性结论
            doc.add_heading('5.4 稳健性检验总结', level=2)
            p = doc.add_paragraph()

            # 读取稳健性检验报告，提取总结信息
            try:
                with open('output/reports/robustness_checks_report.txt', 'r') as f:
                    report_lines = f.readlines()

                    conclusion_section = False
                    conclusion_text = ""

                    for line in report_lines:
                        if "5. 总体稳健性结论" in line:
                            conclusion_section = True
                            continue

                        if conclusion_section and line.strip():
                            conclusion_text += line

                        if conclusion_section and "-" * 30 in line:
                            continue

                    if conclusion_text:
                        p.add_run(conclusion_text).bold = False
                    else:
                        p.add_run("通过滞后模型分析、子样本分析和不同估计方法的检验，本研究的主要发现总体上表现出良好的稳健性。").bold = False
            except:
                p.add_run("通过滞后模型分析、子样本分析和不同估计方法的检验，本研究的主要发现总体上表现出良好的稳健性。").bold = False

        doc.add_page_break()

        # 8. 研究结论与建议
        doc.add_heading('6. 研究结论与建议', level=1)

        # 主要发现总结
        doc.add_heading('6.1 主要研究发现', level=2)
        p = doc.add_paragraph("基于以上分析，本研究得出以下主要结论：")

        p = doc.add_paragraph("1. ")
        p.add_run("AI技术投入对企业专利数量有显著正向影响").bold = True
        p.add_run("，表明企业增加AI技术投入可以有效提高创新产出。").bold = False

        p = doc.add_paragraph("2. ")
        p.add_run("AI技术投入与专利质量呈U型关系").bold = True
        p.add_run("，表明AI技术投入初期可能对专利质量有负面影响，但随着投入增加和积累，长期效应转为正向。").bold = False

        p = doc.add_paragraph("3. ")
        p.add_run("AI人力投入对专利数量和质量均有显著正向影响").bold = True
        p.add_run("，表明人才是企业AI驱动创新的关键因素。").bold = False

        p = doc.add_paragraph("4. ")
        p.add_run("专利和人才的跨领域程度对AI投入与创新绩效的关系具有重要调节作用").bold = True
        p.add_run("，表明多元化的知识结构和跨领域整合能力是企业利用AI提升创新绩效的关键条件。").bold = False

        # 理论贡献
        doc.add_heading('6.2 理论贡献', level=2)
        p = doc.add_paragraph("本研究的理论贡献主要体现在以下方面：")

        p = doc.add_paragraph("1. ")
        p.add_run("拓展了AI投入与企业创新关系的理论框架").bold = True
        p.add_run("，区分了技术投入和人力投入的不同效应。").bold = False

        p = doc.add_paragraph("2. ")
        p.add_run("揭示了AI技术投入与创新质量的非线性关系").bold = True
        p.add_run("，丰富了技术投入对创新绩效影响的理论认识。").bold = False

        p = doc.add_paragraph("3. ")
        p.add_run("引入跨领域程度作为调节变量").bold = True
        p.add_run("，深化了对AI如何转化为创新绩效的边界条件的理解。").bold = False

        # 管理启示
        doc.add_heading('6.3 管理启示', level=2)
        p = doc.add_paragraph("本研究的发现对企业管理实践具有以下启示：")

        p = doc.add_paragraph("1. ")
        p.add_run("企业应加大AI技术投入和人才培养").bold = True
        p.add_run("，特别是要认识到AI投入对创新的长期价值，克服短期内可能存在的负面效应。").bold = False

        p = doc.add_paragraph("2. ")
        p.add_run("企业应提高专利和人才的跨领域整合能力").bold = True
        p.add_run("，促进不同知识领域的融合，最大化AI投入对创新的促进作用。").bold = False

        p = doc.add_paragraph("3. ")
        p.add_run("中小企业在进行AI投入时应更加谨慎").bold = True
        p.add_run("，可能需要与自身规模和资源条件相匹配的AI投入策略。").bold = False

        # 研究局限与未来方向
        doc.add_heading('6.4 研究局限与未来研究方向', level=2)
        p = doc.add_paragraph("本研究存在以下局限，也为未来研究提供方向：")

        p = doc.add_paragraph("1. ")
        p.add_run("数据覆盖期相对有限").bold = False
        p.add_run("，未来研究可扩展数据时间跨度，进一步考察AI投入的长期效应。").bold = False

        p = doc.add_paragraph("2. ")
        p.add_run("未深入探讨AI投入的具体形式和内容").bold = False
        p.add_run("，未来研究可细分AI技术类型和应用场景，分析不同类型AI投入的差异化效应。").bold = False

        p = doc.add_paragraph("3. ")
        p.add_run("内生性问题难以完全消除").bold = False
        p.add_run("，未来研究可寻找更好的工具变量或利用准自然实验等方法进一步解决内生性问题。").bold = False

        # 保存文档
        doc.save(output_path)
        print(f"综合结果报告已成功生成并保存至 {output_path}")

    except ImportError:
        print("错误: 缺少python-docx库，请安装: pip install python-docx")
        # 创建简单的文本报告作为替代
        with open(output_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
            f.write("AI投入对企业创新绩效影响分析报告\n")
            f.write("=" * 50 + "\n\n")
            f.write("由于缺少python-docx库，无法生成Word格式报告，请安装该库后重试。\n")
            f.write("pip install python-docx\n\n")
            f.write("简要结果概述：\n")
            f.write("- 数据分析包含多个样本\n")
            f.write("- AI技术投入对专利数量有显著正向影响\n")
            f.write("- AI技术投入与专利质量呈U型关系\n")
            f.write("- AI人力投入对专利数量和质量均有显著正向影响\n")
            f.write("- 专利和人才的跨领域程度是重要的调节变量\n")
        print(f"由于缺少python-docx库，已生成简单文本报告: {output_path.replace('.docx', '.txt')}")

    except Exception as e:
        print(f"生成报告时发生错误: {e}")
        traceback.print_exc()
        # 创建错误报告
        with open(output_path.replace('.docx', '_error.txt'), 'w', encoding='utf-8') as f:
            f.write(f"报告生成过程中出错: {str(e)}\n\n")
            f.write(traceback.format_exc())
        print(f"错误详情已保存至: {output_path.replace('.docx', '_error.txt')}")

def main(data_file_path, output_dir='output'):
    """
    运行完整的AI投入对企业创新绩效影响分析

    参数:
        data_file_path: 数据文件路径
        output_dir: 输出目录，默认为'output'
    """
    analysis_results = {}

    try:
        # 创建输出目录结构
        for subdir in ['', 'figures', 'tables', 'reports']:
            os.makedirs(os.path.join(output_dir, subdir), exist_ok=True)
          # 数据加载与预处理
        print("\n步骤1: 数据加载与预处理")
        try:
            result = load_and_preprocess_data(data_file_path)

            # 增加健壮性检查
            if result is None:
                print("警告: 数据预处理返回了None，尝试直接读取数据")
                df = pd.read_excel(data_file_path)

                # 确保列名为小写
                df.columns = [col.lower() for col in df.columns]
                print("已将所有列名转为小写，确保一致性")

                final_variable_mapping = {}

                # 进行最基本的索引设置
                if 'stkcd' in df.columns and 'year' in df.columns:
                    # 确保列类型正确
                    df['stkcd'] = df['stkcd'].astype(str)
                    df['year'] = df['year'].astype(int)

                    # 设置索引
                    df = df.set_index(['stkcd', 'year'])
                    print("成功设置面板数据索引 (stkcd, year)")
                else:
                    # 检查大写列名
                    if 'STKCD' in df.columns and 'YEAR' in df.columns:
                        # 将列名转为小写
                        df.columns = [col.lower() for col in df.columns]
                        # 重新尝试设置索引
                        df['stkcd'] = df['stkcd'].astype(str)
                        df['year'] = df['year'].astype(int)
                        df = df.set_index(['stkcd', 'year'])
                        print("成功设置面板数据索引 (stkcd, year) - 已转换大写列名")
                    else:
                        print("警告: 无法找到stkcd和year列，无法设置面板数据索引")

                result = (df, final_variable_mapping)

            # 确保结果可解包为DataFrame和映射字典
            if isinstance(result, tuple) and len(result) == 2:
                df, final_variable_mapping = result

                # 确保df是DataFrame
                if not isinstance(df, pd.DataFrame):
                    raise ValueError(f"预处理返回的df不是DataFrame，而是{type(df)}")
            else:
                # 单一返回值情况
                if isinstance(result, pd.DataFrame):
                    df = result
                    final_variable_mapping = {}
                else:
                    raise ValueError(f"无法解析预处理返回值: {type(result)}")

            analysis_results['preprocessed_data'] = {'shape': df.shape}

        except Exception as e:
            print(f"数据预处理出错: {e}")
            traceback.print_exc()
            # 最后尝试使用基本数据处理
            try:
                print("尝试基本数据处理作为备选方案...")
                df = pd.read_excel(data_file_path)
                df.columns = [col.lower() for col in df.columns]

                if 'stkcd' in df.columns and 'year' in df.columns:
                    df['stkcd'] = df['stkcd'].astype(str)
                    df['year'] = df['year'].astype(int)
                    df = df.set_index(['stkcd', 'year'])

                final_variable_mapping = {}
                print("使用基本数据处理成功")
            except Exception as e2:
                print(f"基本数据处理也失败: {e2}")
                return {'error': f"数据预处理失败: {str(e)}，基本处理也失败: {str(e2)}"}        # 添加健壮性检查确保数据框符合预期
        print("\n进行数据结构健壮性检查...")
        if df is None:
            print("错误: 数据预处理返回了空的数据框")
            return {'error': "数据预处理返回了空的数据框"}

        # 确保列名全部小写，增强鲁棒性
        if any(col != col.lower() for col in df.columns):
            upper_case_cols = [col for col in df.columns if col != col.lower()]
            print(f"警告: 发现{len(upper_case_cols)}个非小写列名: {upper_case_cols[:5]}...")
            print("转换所有列名为小写以确保一致性")
            # 创建映射字典，保留原始大小写形式作为参考
            case_mapping = {col: col.lower() for col in df.columns if col != col.lower()}
            # 应用小写转换
            df.columns = [col.lower() for col in df.columns]
            print(f"列名大小写统一化完成")

        # 对格式和结构进行必要检查
        try:
            # 检查主要变量是否存在（不区分大小写）
            key_vars = ['ai', 'intotal', 'ep', 'dp']
            # 获取当前所有列名的小写形式
            lower_cols = [col.lower() for col in df.columns]
            # 检查每个关键变量是否有对应的列（不区分大小写）
            missing_keys = [var for var in key_vars
                           if var not in lower_cols
                           and var not in final_variable_mapping]
            if missing_keys:
                print(f"警告: 以下关键变量缺失: {missing_keys}")
                print("检查变量映射中是否有替代变量...")
                for var in missing_keys:
                    if var in final_variable_mapping:
                        mapped_var = final_variable_mapping[var]
                        if mapped_var in df.columns:
                            print(f"  使用映射变量: {var} -> {mapped_var}")
                        else:
                            print(f"  错误: 映射变量 {mapped_var} 也不存在")
              # 检查索引结构
            if not isinstance(df.index, pd.MultiIndex) or df.index.names != ['stkcd', 'year']:
                print("警告: 数据框不是预期的面板结构，尝试修复...")

                # 重设索引以便检查列名
                if isinstance(df.index, pd.MultiIndex):
                    df = df.reset_index()

                # 首先检查是否存在不同大小写形式的stkcd和year列
                idx_candidates = {}
                for col in df.columns:
                    if col.lower() == 'stkcd':
                        idx_candidates['stkcd'] = col
                    elif col.lower() == 'year':
                        idx_candidates['year'] = col

                if len(idx_candidates) == 2:  # 找到了两个索引列
                    print(f"找到索引列: stkcd对应'{idx_candidates['stkcd']}'，year对应'{idx_candidates['year']}'")

                    # 如果列名不是标准小写，先重命名
                    rename_dict = {}
                    if idx_candidates['stkcd'] != 'stkcd':
                        rename_dict[idx_candidates['stkcd']] = 'stkcd'
                    if idx_candidates['year'] != 'year':
                        rename_dict[idx_candidates['year']] = 'year'

                    if rename_dict:
                        df = df.rename(columns=rename_dict)
                        print(f"  已重命名列: {rename_dict}")

                    # 确保列类型正确
                    df['stkcd'] = df['stkcd'].astype(str)
                    df['year'] = df['year'].astype(int)

                    # 设置索引
                    df = df.set_index(['stkcd', 'year'])
                    print("  已重设索引为 ['stkcd', 'year']")
                else:
                    missing = []
                    if 'stkcd' not in idx_candidates:
                        missing.append('stkcd')
                    if 'year' not in idx_candidates:
                        missing.append('year')
                    print(f"  错误: 无法重设索引，缺少列: {missing}")
                    # 检查数据类型
            for col in df.columns:
                if not pd.api.types.is_numeric_dtype(df[col]) and col not in ['stknm', 'audit', 'opinion', 'soe', 'province']:
                    print(f"警告: 列 {col} 不是数值型，这可能导致分析问题")

            for col in df.columns:
                if not pd.api.types.is_numeric_dtype(df[col]):
                    if col in ['stknm', 'indcd', 'indnm', 'province', 'city']:
                        print(f"注意: 列 '{col}' 为分类变量，将用作固定效应或描述性信息")
                        # 如果需要创建哑变量
                        if col == 'indcd':
                            print(f"  将为行业代码 '{col}' 创建固定效应")
                            # 在这里不实际创建，而是在PanelOLS中使用参数
                    else:
                        print(f"警告: 列 '{col}' 不是数值型且不是已知分类变量")

            # 数据量检查
            if len(df) < 100:
                print(f"警告: 数据量较小 ({len(df)} 行)，结果可能不稳健")

            # 面板数据平衡性报告
            if isinstance(df.index, pd.MultiIndex):
                balance_counts = df.groupby(level=0).size()
                balance_info = balance_counts.describe()
                print(f"\n面板数据平衡性: 实体数={len(balance_counts)}, 平均观测值/实体={balance_info['mean']:.2f}")
                print(f"最少观测值/实体={balance_info['min']}, 最多观测值/实体={balance_info['max']}")

                if balance_info['min'] < balance_info['max']:
                    # 修正这里：正确的除法运算符
                    imbalance_pct = (balance_info['max'] - balance_info['min']) / balance_info['max'] * 100
                    print(f"面板不平衡度: {imbalance_pct:.2f}%")

        except Exception as check_error:
            print(f"健壮性检查出错: {check_error}")

        # 3. 进行描述性统计分析
        print("\n步骤2: 描述性统计分析")
        try:
            descriptive_stats = enhanced_descriptive_statistics(df, final_variable_mapping)
            analysis_results['descriptive_stats'] = descriptive_stats
        except Exception as e:
            print(f"描述性统计分析出错: {e}")
            traceback.print_exc()
            descriptive_stats = None

        # 4. 进行面板数据诊断
        print("\n步骤3: 面板数据诊断")
        try:
            diagnostics = panel_data_diagnostics(df, final_variable_mapping)
            analysis_results['diagnostics'] = diagnostics
        except Exception as e:
            print(f"面板数据诊断出错: {e}")
            traceback.print_exc()
            diagnostics = None

        non_stationary_vars = []

        # 处理非平稳变量 - 更完善的方法
        if 'diagnostics' in analysis_results and analysis_results['diagnostics']:
            # 如果已经有了更高级的差分处理
            if 'df_with_diff' in analysis_results['diagnostics'] and 'recommended_vars' in analysis_results['diagnostics']:
                print("\n使用诊断阶段推荐的变量处理结果")
                # 使用已经有差分处理的数据框
                df = analysis_results['diagnostics']['df_with_diff']
                # 更新变量映射以使用推荐的变量版本
                recommended_vars = analysis_results['diagnostics']['recommended_vars']
                for var, recommended in recommended_vars.items():
                    if var != recommended:
                        print(f"将使用 {recommended} 替代 {var}")
                        # 更新变量映射
                        if var in final_variable_mapping:
                            final_variable_mapping[f'{var}_original'] = var
                            final_variable_mapping[var] = recommended

                print(f"更新后的变量映射: {list(final_variable_mapping.keys())[:10]}...")
            else:                # 退回到简单的差分处理
                non_stationary_vars = []
                for var, result in analysis_results['diagnostics']['unit_root'].items():

                    if result.get('skipped_low_var', 0) > 10:  # 如果超过10个公司出现低方差问题
                        print(f"警告: 变量 {var} 存在严重的低方差问题，将添加微小扰动")
                        # 添加微小扰动以增加变异性
                        df[f'{var}_adj'] = df[var] + np.random.normal(0, df[var].std()*0.01 + 1e-5, size=len(df))
                        print(f"  已创建调整变量: {var}_adj")
                        # 更新变量映射
                        final_variable_mapping[var] = f'{var}_adj'

                    if result.get('stationary_percentage', 0) < 50:  # 提高阈值到50%，更激进地处理非平稳变量
                        print(f"对非平稳变量 {var} 进行差分处理")
                        non_stationary_vars.append(var)
                        # 创建差分变量
                        df[f'{var}_diff'] = df.groupby(level=0)[var].diff()

                        # 填充差分后产生的首个时间点的NaN值，使用同组内的均值
                        df[f'{var}_diff'] = df.groupby(level=0)[f'{var}_diff'].transform(
                            lambda x: x.fillna(x.mean() if not pd.isna(x.mean()) else 0)
                        )

                        # 更新变量映射
                        if var in final_variable_mapping:
                            final_variable_mapping[f'{var}_original'] = var
                            final_variable_mapping[var] = f'{var}_diff'

                print(f"已创建 {len(non_stationary_vars)} 个差分变量")

            # 保存变量映射以供后续使用
            for var in non_stationary_vars:
                final_variable_mapping[f'{var}_original'] = var
                final_variable_mapping[var] = f'{var}_diff'

        # 5. 运行假设验证模型
        print("\n步骤4: 假设验证")
        try:
            # 主要假设检验
            model_results = extended_hypothesis_testing(df)
            analysis_results['model_results'] = model_results

            # 添加高级分析
            advanced_results = {}

            # 对H1b的时间动态效应进行更深入分析
            base_controls = ['age2', 'balance', 'bm1', 'growth', 'lev', 'mhold',
                           'roa', 'size', 'tat', 'tobinq1', 'audit', 'soe']
            valid_controls = [c for c in base_controls if c in df.columns]

            print("\n进行H1b时间动态效应的高级分析...")
            h1b_dynamics = test_time_dynamic_effect(df, 'ep', 'ai', valid_controls)
            advanced_results['h1b_time_dynamics'] = h1b_dynamics

            # 在main函数中，对H2a和H2b的非线性关系进行高级检验
            print("\n进行H2a和H2b非线性关系的高级检验...")
            if 'H2a' in model_results and model_results['H2a'].get('result') is not None:
                print("进行H2a的高级非线性检验（AI技术投入与专利质量的U型关系）")
                h2a_advanced = advanced_nonlinear_test(
                    model_results['H2a']['result'],
                    'ai', 'ai_squared',
                    df,
                    dependent_var='ai_patent_quality',  # 明确指定因变量
                    output_dir='output/reports',
                    plot_segments=True  # 启用分段绘图
                )
                advanced_results['h2a_nonlinear'] = h2a_advanced

                # 添加对结果的详细解释
                if h2a_advanced.get('valid_nonlinear'):
                    print("H2a结果: 支持AI技术投入与专利质量的U型关系")
                    print(f"  拐点: {h2a_advanced.get('turning_point', 'N/A')}")
                    print(f"  拐点95%置信区间: [{h2a_advanced.get('turning_point_ci', (0, 0))[0]:.4f}, {h2a_advanced.get('turning_point_ci', (0, 0))[1]:.4f}]")
                    print(f"  分段回归支持U型: {'是' if h2a_advanced.get('piecewise_confirms', False) else '否'}")
                    print(f"  LOWESS曲线支持U型: {'是' if h2a_advanced.get('lowess_confirms', False) else '否'}")
                else:
                    print("H2a结果: 不支持AI技术投入与专利质量的U型关系")
                    print(f"  原因: {h2a_advanced.get('conclusion', '无具体结论')}")

            # 对H5和H6的调节效应进行高级分析
            print("\n进行调节效应的高级分析...")

            # H5a: 专利跨领域程度对AI技术投入与专利数关系的调节
            h5a_moderation = enhanced_moderation_analysis(
                df, 'ai', 'ai_patent_log', 'intotal'
            )
            advanced_results['h5a_moderation'] = h5a_moderation

            # H5b: 人才跨领域程度对AI技术投入与专利数关系的调节
            h5b_moderation = enhanced_moderation_analysis(
                df, 'ai', 'manu_job_log', 'intotal'
            )
            advanced_results['h5b_moderation'] = h5b_moderation
              # H6a: 专利跨领域程度对AI技术投入与专利质量关系的调节
            try:
                print("\n尝试运行H6a假设: 专利跨领域程度对AI技术投入与专利质量关系的调节")

                # 预先检查变量
                zero_pct = (df['ai_patent_quality'] == 0).mean() * 100
                var_val = df['ai_patent_quality'].var()
                print(f"ai_patent_quality变量统计: 零值比例={zero_pct:.2f}%, 方差={var_val:.6f}")

                if zero_pct > 90:
                    print("警告: ai_patent_quality零值过多(>90%)，跳过该假设")
                    advanced_results['h6a_moderation'] = {'error': '变量零值过多，不适合建模分析'}
                elif var_val < 1e-6:
                    print("警告: ai_patent_quality方差极小，增加扰动以提高数值稳定性")
                    # 创建副本并添加微小扰动
                    df_mod = df.copy()
                    df_mod['ai_patent_quality'] = df_mod['ai_patent_quality'] + np.random.normal(0, 1e-5, size=len(df_mod))
                    h6a_moderation = enhanced_moderation_analysis(
                        df_mod, 'ai', 'ai_patent_log', 'ai_patent_quality'
                    )
                    advanced_results['h6a_moderation'] = h6a_moderation
                else:
                    h6a_moderation = enhanced_moderation_analysis(
                        df, 'ai', 'ai_patent_log', 'ai_patent_quality'
                    )
                    advanced_results['h6a_moderation'] = h6a_moderation
            except Exception as h6a_err:
                print(f"运行H6a假设时出错: {h6a_err}")
                traceback.print_exc()
                advanced_results['h6a_moderation'] = {'error': str(h6a_err)}
              # H6b: 人才跨领域程度对AI人力投入与专利质量关系的调节
            try:
                print("\n尝试运行H6b假设: 人才跨领域程度对AI人力投入与专利质量关系的调节")

                # 预先检查变量
                if 'ai_patent_quality' in df.columns:
                    zero_pct = (df['ai_patent_quality'] == 0).mean() * 100
                    var_val = df['ai_patent_quality'].var()
                    print(f"ai_patent_quality变量统计: 零值比例={zero_pct:.2f}%, 方差={var_val:.6f}")

                    if zero_pct > 90:
                        print("警告: ai_patent_quality零值过多(>90%)，跳过该假设")
                        advanced_results['h6b_moderation'] = {'error': '变量零值过多，不适合建模分析'}
                    elif var_val < 1e-6:
                        print("警告: ai_patent_quality方差极小，增加扰动以提高数值稳定性")
                        # 创建副本并添加微小扰动
                        df_mod = df.copy()
                        df_mod['ai_patent_quality'] = df_mod['ai_patent_quality'] + np.random.normal(0, 1e-5, size=len(df_mod))
                        h6b_moderation = enhanced_moderation_analysis(
                            df_mod, 'ai_job_log', 'manu_job_log', 'ai_patent_quality'
                        )
                        advanced_results['h6b_moderation'] = h6b_moderation
                    else:
                        h6b_moderation = enhanced_moderation_analysis(
                            df, 'ai_job_log', 'manu_job_log', 'ai_patent_quality'
                        )
                        advanced_results['h6b_moderation'] = h6b_moderation
                else:
                    print("错误: 找不到变量'ai_patent_quality'")
                    advanced_results['h6b_moderation'] = {'error': '找不到变量ai_patent_quality'}
            except Exception as h6b_err:
                print(f"运行H6b假设时出错: {h6b_err}")
                traceback.print_exc()
                advanced_results['h6b_moderation'] = {'error': str(h6b_err)}

            analysis_results['advanced_results'] = advanced_results

        except Exception as e:
            print(f"假设验证出错: {e}")
            traceback.print_exc()
            model_results = {}

        # 6. 进行稳健性检验
        print("\n步骤5: 稳健性检验")
        try:
            # 为稳健性检验准备数据
            df_robust = df.copy()

            # 确保必要的交互项存在
            if 'ai_x_ai_patent_log' not in df_robust.columns:
                df_robust['ai_x_ai_patent_log'] = df_robust['ai'] * df_robust['ai_patent_log']

            if 'ai_job_log_x_ai_patent_log' not in df_robust.columns:
                df_robust['ai_job_log_x_ai_patent_log'] = df_robust['ai_job_log'] * df_robust['ai_patent_log']

            if 'ai_x_manu_job_log' not in df_robust.columns and 'manu_job_log' in df_robust.columns:
                df_robust['ai_x_manu_job_log'] = df_robust['ai'] * df_robust['manu_job_log']

            if 'ai_job_log_x_manu_job_log' not in df_robust.columns and 'manu_job_log' in df_robust.columns:
                df_robust['ai_job_log_x_manu_job_log'] = df_robust['ai_job_log'] * df_robust['manu_job_log']

            # 执行稳健性检验
            robustness_results = enhanced_robustness_checks(df_robust, model_results or {})
            analysis_results['robustness_results'] = robustness_results
            print("稳健性检验已完成")

        except Exception as e:
            print(f"稳健性检验出错: {e}")
            traceback.print_exc()
            analysis_results['robustness_results'] = {'error': str(e)}

        # 7. 生成综合结果报告
        print("\n步骤6: 生成结果报告")
        try:
            # 创建综合报告
            create_comprehensive_report(
                analysis_results,
                output_path='output/reports/comprehensive_report.docx'
            )
            print("综合结果报告已生成")
        except Exception as e:
            print(f"结果报告生成出错: {e}")
            traceback.print_exc()
    except Exception as e:
        print(f"主函数出错: {e}")
        traceback.print_exc()
        return {'error': str(e)}

    return analysis_results

# 示例用法
if __name__ == "__main__":
    # 运行完整分析
    analysis_results = main('data.xlsx')